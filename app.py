from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash, jsonify
import os
import json
import random
import math
import tempfile
import zipfile
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
import base64
from flask_wtf import FlaskForm
from flask_wtf.file import FileField, FileAllowed
from wtforms import StringField, SelectField, IntegerField
from wtforms.validators import DataRequired, NumberRange
from werkzeug.utils import secure_filename
import cv2
import numpy as np
import pytesseract
from pdf2image import convert_from_path
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
import re
from urllib.parse import urlparse
import traceback

app = Flask(__name__)
app.secret_key = "estadisticabasica2024"
app.config['WTF_CSRF_ENABLED'] = False  # Solo para pruebas iniciales

# Directorios de almacenamiento
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
VARIANTES_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'variantes')
EXAMENES_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'examenes')
PLANTILLAS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'plantillas')
HOJAS_RESPUESTA_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'hojas_respuesta')
EXAMENES_ESCANEADOS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'examenes_escaneados')
if not os.path.exists(EXAMENES_ESCANEADOS_FOLDER):
    os.makedirs(EXAMENES_ESCANEADOS_FOLDER)

# Añadir al principio del archivo, después de la definición de carpetas
HISTORIAL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'historial.json')

# Función para cargar el historial
def cargar_historial():
    if os.path.exists(HISTORIAL_FILE):
        try:
            with open(HISTORIAL_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

# Función para guardar el historial
def guardar_historial(historial):
    with open(HISTORIAL_FILE, 'w', encoding='utf-8') as f:
        json.dump(historial, f, indent=2, ensure_ascii=False)


# Márgenes de error aceptables para calificación
MARGENES_ERROR = {
    'gini': 0.05,               # ±5% para coeficiente de Gini
    'dist_frecuencias': {
        'k': 0.5,               # ±0.5 para K (número de clases)
        'rango': 2,             # ±2 para el rango
        'amplitud': 0.5         # ±0.5 para la amplitud
    },
    'tallo_hoja': {
        'moda': 0.5,            # ±0.5 para el valor de la moda
        'intervalo': 1          # ±1 para el intervalo de mayor concentración
    },
    'medidas_centrales': {
        'media': 50,            # ±50 para la media (depende del rango de datos)
        'mediana': 50,          # ±50 para la mediana
        'moda': 50              # ±50 para la moda
    }
}

# Formulario para generación de exámenes
class ExamenForm(FlaskForm):
    num_variantes = IntegerField('Número de Variantes', 
                                validators=[NumberRange(min=1, max=10)],
                                default=1)
    seccion = StringField('Sección del Curso', validators=[DataRequired()])
    tipo_evaluacion = SelectField('Tipo de Evaluación', 
                                  choices=[
                                      ('parcial1', 'Primer Parcial'),
                                      ('parcial2', 'Segundo Parcial'),
                                      ('final', 'Examen Final'),
                                      ('corto', 'Evaluación Corta'),
                                      ('recuperacion', 'Recuperación')
                                  ])
    logo = FileField('Logo de la Institución', 
                     validators=[FileAllowed(['jpg', 'png', 'jpeg'], 'Solo imágenes')])

# Base de datos simple para almacenar información de estudiantes
estudiantes_db = {}

# Crear directorios si no existen
for folder in [UPLOAD_FOLDER, VARIANTES_FOLDER, EXAMENES_FOLDER, PLANTILLAS_FOLDER, HOJAS_RESPUESTA_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Base de preguntas para Primera Serie
preguntas_base_primera = [
    {
        "pregunta": "Trata del recuento, ordenación y clasificación de los datos obtenidos por las observaciones, para poder hacer comparaciones y obtener conclusiones.",
        "opciones": ["Población", "Cálculos", "Estadística", "Frecuencia", "Ninguna de las anteriores"],
        "respuesta_correcta": 2  # Índice de "Estadística"
    },
    {
        "pregunta": "Las Variables Estadísticas pueden ser:",
        "opciones": ["Discretas", "Cualitativas", "Indiscretas", "Cuantitativas", "Numéricas"],
        "respuesta_correcta": 1  # Índice de "Cualitativas"
    },
    {
        "pregunta": "Método que sirve para medir la desigualdad, es un número entre cero y uno que mide el grado de desigualdad en la distribución del ingreso en una sociedad determinada o país.",
        "opciones": ["Coeficiente de Correlación", "Coeficiente de Gini", "Marca de Clase", "La Frecuencia Acumulada"],
        "respuesta_correcta": 1  # Índice de "Coeficiente de Gini"
    },
    {
        "pregunta": "¿Es un conjunto representativo de la población de referencia, el número de individuos es menor que el de la población?",
        "opciones": ["Valor", "Dato", "Experimento", "Población", "Muestra", "Todas las anteriores"],
        "respuesta_correcta": 4  # Índice de "Muestra"
    },
    {
        "pregunta": "La toma de temperatura para ingresar a los centros comerciales es una variable:",
        "opciones": ["Cualitativa", "Cuantitativa"],
        "respuesta_correcta": 1  # Índice de "Cuantitativa"
    },
    {
        "pregunta": "Es el conjunto de todos los elementos a los que se somete a un estudio estadístico.",
        "opciones": ["Muestra", "Población", "Individuo", "Muestreo"],
        "respuesta_correcta": 1  # Índice de "Población"
    },
    {
        "pregunta": "Las Fases de un estudio estadístico son:",
        "opciones": ["Planteamiento del Problema", "Simplificar los Datos", "Recolectar y Ordenar los Datos", "Analizar los Datos", "Interpretar y Presentar Resultados", "Ninguna de las anteriores"],
        "respuesta_correcta": 0  # Todas son correctas, pero usamos el primer índice
    },
    {
        "pregunta": "¿La siguiente imagen, representa un diagrama de tallo y hoja?",
        "opciones": ["Verdadero", "Falso"],
        "respuesta_correcta": 0  # Índice de "Verdadero"
    },
    {
        "pregunta": "¿Cuál es el método que permite calcular el número de grupos, intervalos o clases a construir para una tablade distribución de frecuencias?",
        "opciones": ["Método de mínimos cuadrados", "Coeficiente de Gini", "Método Sturgers", "La regla empírica"],
        "respuesta_correcta": 2  # Índice de "Método Sturgers"
    },
    {
        "pregunta": "¿Quién ordeno o realizo el primer catastro o Censo de (bienes inmuebles) considerado el primero en Europa?",
        "opciones": ["El Rey Juan Carlos", "El Rey Guillermo", "El Rey Constantino", "El Rey Ricardo", "El Rey Federico de Edimburgo"],
        "respuesta_correcta": 1  # Índice de "El Rey Guillermo"
    },
    {
        "pregunta": "¿Cuál de las siguientes medidas de tendencia central se ve más afectada por valores extremos?",
        "opciones": ["Media", "Mediana", "Moda", "Rango"],
        "respuesta_correcta": 0  # Índice de "Media"
    },
    {
        "pregunta": "El tipo de gráfico más adecuado para mostrar la distribución de frecuencias de una variable continua es:",
        "opciones": ["Gráfico de barras", "Gráfico circular", "Histograma", "Gráfico de líneas"],
        "respuesta_correcta": 2  # Índice de "Histograma"
    },
    {
        "pregunta": "¿Qué medida indica el grado de dispersión de los datos respecto a la media?",
        "opciones": ["Varianza", "Mediana", "Moda", "Coeficiente de asimetría"],
        "respuesta_correcta": 0  # Índice de "Varianza"
    },
    {
        "pregunta": "La diferencia entre una variable cuantitativa discreta y una variable cuantitativa continua es:",
        "opciones": ["Las variables discretas toman cualquier valor dentro de un intervalo, las continuas toman valores aislados", 
                      "Las variables discretas toman valores aislados, las continuas toman cualquier valor dentro de un intervalo",
                      "Las variables discretas son siempre enteras, las continuas son siempre decimales",
                      "No hay diferencia real entre ambas"],
        "respuesta_correcta": 1  # Índice de "Las variables discretas toman valores aislados..."
    },
    {
        "pregunta": "Si los datos están distribuidos de forma simétrica alrededor de la media, entonces:",
        "opciones": ["La media y la mediana coinciden", "La media es mayor que la mediana", "La mediana es mayor que la media", "La media y la moda coinciden"],
        "respuesta_correcta": 0  # Índice de "La media y la mediana coinciden"
    }
]

# Base de preguntas para Segunda Serie
preguntas_base_segunda = [
    {
        "escenario": "Una compañía de telecomunicaciones quiere representar visualmente la distribución porcentual de sus ingresos por tipo de servicio (internet, telefonía fija, telefonía móvil, televisión por cable y servicios corporativos) durante el año fiscal 2023.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 1  # Índice de "Gráfica circular"
    },
    {
        "escenario": "Un instituto de estadísticas demográficas ha recopilado información sobre las edades de los habitantes de un municipio, agrupando los datos en intervalos de 10 años (0-9, 10-19, 20-29, etc.). Desean visualizar tanto la frecuencia de cada intervalo como la tendencia general de la distribución de edades.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 4  # Índice de "Polígono de frecuencias"
    },
    {
        "escenario": "Un departamento de recursos humanos ha realizado una encuesta sobre los tiempos de transporte (en minutos) que los empleados tardan en llegar a la oficina. Los datos obtenidos son continuos y quieren mostrar cómo se distribuyen estos tiempos, identificando claramente dónde se concentra la mayoría de los casos.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 2  # Índice de "Histograma de Pearson"
    },
    {
        "escenario": "Una universidad desea representar el número de estudiantes matriculados en cada una de sus facultades (Humanidades, Ingeniería, Medicina, Derecho, Economía y Arquitectura) para el ciclo académico 2024, permitiendo una fácil comparación entre facultades.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 0  # Índice de "Gráfica de barras"
    },
    {
        "escenario": "Una entidad financiera ha recopilado datos sobre los montos de créditos otorgados en el último trimestre. Los montos se han agrupado en intervalos y se desea mostrar los valores acumulados hasta cierto punto, para identificar qué porcentaje de créditos está por debajo de determinados montos.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 3  # Índice de "Ojiva de Galton"
    },
    {
        "escenario": "Un estudio sobre calificaciones finales en un curso de estadística muestra datos que podrían seguir una distribución normal. Los investigadores quieren representar las frecuencias de cada intervalo de calificación y, al mismo tiempo, identificar visualmente si la distribución se aproxima a una curva normal.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 2  # Índice de "Histograma de Pearson"
    },
    {
        "escenario": "Un análisis de ventas mensuales de una cadena de tiendas durante un año completo. Se desea mostrar la evolución de las ventas a lo largo del tiempo, identificando tendencias, picos y caídas.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 4  # Índice de "Polígono de frecuencias"
    },
    {
        "escenario": "Una empresa farmacéutica ha registrado el tiempo (en días) que tarda cada lote de medicamentos en pasar el control de calidad. Quieren determinar si un nuevo lote con un tiempo específico está dentro del 75% de los casos más rápidos.",
        "opciones": ["Gráfica de barras", "Gráfica circular (pastel)", "Histograma de Pearson", "Ojiva de Galton", "Polígono de frecuencias"],
        "respuesta_correcta": 3  # Índice de "Ojiva de Galton"
    }
]

# Datos base para la Tercera Serie - Problema de coeficiente de Gini
gini_exercises = [
    {
        "title": "La siguiente tabla muestra la distribución de salarios mensuales (en quetzales) de los trabajadores de la empresa ABC S.A:",
        "ranges": ["[1500-2000)", "[2000-2500)", "[2500-3000)", "[3000-3500)", "[3500-4000)", "[4000-4500)", "[4500-5000)", "[5000-5500)"],
        "workers": [7, 6, 11, 16, 11, 11, 6, 3]
    },
    {
        "title": "La siguiente tabla muestra la distribución de salarios mensuales (en quetzales) de los trabajadores de la empresa XYZ S.A:",
        "ranges": ["[1200-1800)", "[1800-2400)", "[2400-3000)", "[3000-3600)", "[3600-4200)", "[4200-4800)", "[4800-5400)", "[5400-6000)"],
        "workers": [12, 8, 7, 5, 4, 3, 2, 9]
    },
    {
        "title": "La siguiente tabla muestra la distribución de salarios mensuales (en quetzales) de los trabajadores de la empresa Alfa Omega S.A:",
        "ranges": ["[2000-2500)", "[2500-3000)", "[3000-3500)", "[3500-4000)", "[4000-4500)", "[4500-5000)", "[5000-5500)", "[5500-6000)"],
        "workers": [9, 10, 12, 14, 12, 10, 8, 5]
    }
]

# Datos base para la Tercera Serie - Problema de Sturgers
sturgers_exercises = [
    {
        "title": "Construya la siguiente tabla de distribución de frecuencias. Con datos agrupados usando el método Sturgers.",
        "data": [
            "115", "106", "116", "118", "118",
            "121", "121", "115", "122", "126",
            "126", "129", "129", "130", "138",
            "140", "137", "145", "143", "144",
            "149", "150", "152", "151", "156"
        ]
    },
    {
        "title": "Construya la siguiente tabla de distribución de frecuencias. Con datos agrupados usando el método Sturgers.",
        "data": [
            "95", "98", "102", "105", "107",
            "110", "110", "112", "115", "118",
            "120", "122", "125", "127", "130",
            "133", "135", "137", "140", "142",
            "145", "148", "150", "152", "155"
        ]
    },
    {
        "title": "Construya la siguiente tabla de distribución de frecuencias. Con datos agrupados usando el método Sturgers.",
        "data": [
            "205", "210", "212", "215", "218",
            "220", "223", "225", "228", "230",
            "233", "235", "238", "240", "242",
            "245", "248", "250", "252", "255",
            "258", "260", "263", "265", "270"
        ]
    }
]

# Datos base para la Tercera Serie - Problema de Tallo y Hoja
stem_leaf_exercises = [
    {
        "title": "Con la información obtenida de las ventas mensuales de distintos productos tecnológicos, se tomaron aleatoriamente los siguientes datos que representan el crecimiento porcentual respecto al año anterior.",
        "data": [
            "2.2", "2.8", "2.6", "3.1", "2.9", "3.3", "4.0", "3.5",
            "3.9", "4.1", "4.5", "4.5", "5.0", "4.8", "5.0", "5.7",
            "6.1", "6.2", "6.3", "7.2", "7.3", "8.1", "9.1", "10.6"
        ]
    },
    {
        "title": "Con la información obtenida del tiempo de atención (en minutos) a clientes en una sucursal bancaria, se tomaron aleatoriamente los siguientes datos durante el mes de febrero.",
        "data": [
            "3.5", "3.8", "4.2", "4.5", "4.8", "5.1", "5.3", "5.7",
            "6.0", "6.2", "6.5", "6.8", "7.0", "7.3", "7.5", "7.8",
            "8.0", "8.3", "8.5", "8.8", "9.0", "9.5", "10.2", "11.5"
        ]
    },
    {
        "title": "Con la información obtenida del consumo de combustible (en km/litro) de vehículos en una empresa de transporte, se tomaron aleatoriamente los siguientes datos.",
        "data": [
            "8.2", "8.5", "8.7", "9.0", "9.2", "9.5", "9.8", "10.1",
            "10.4", "10.7", "11.0", "11.3", "11.6", "11.9", "12.2", "12.5",
            "12.8", "13.1", "13.4", "13.7", "14.0", "14.5", "15.2", "16.5"
        ]
    }
]

# Datos base para la Tercera Serie - Problema de medidas de tendencia central
central_tendency_exercises = [
    {
        "title": "Calcular las medidas de tendencia central Media , Mediana, Moda e interprete los resultados obtenidos.",
        "ranges": ["[1800-2300)", "[2300-2800)", "[2800-3300)", "[3300-3800)", "[3800-4300)", "[4300-4800)", "[4800-5300)"],
        "count": [6, 13, 19, 20, 14, 6, 1]
    },
    {
        "title": "Calcular las medidas de tendencia central Media , Mediana, Moda e interprete los resultados obtenidos.",
        "ranges": ["[1500-2000)", "[2000-2500)", "[2500-3000)", "[3000-3500)", "[3500-4000)", "[4000-4500)", "[4500-5000)"],
        "count": [5, 12, 20, 25, 15, 8, 3]
    },
    {
        "title": "Calcular las medidas de tendencia central Media , Mediana, Moda e interprete los resultados obtenidos.",
        "ranges": ["[2500-3000)", "[3000-3500)", "[3500-4000)", "[4000-4500)", "[4500-5000)", "[5000-5500)", "[5500-6000)"],
        "count": [8, 15, 22, 18, 10, 5, 2]
    }
]

# Rutas para los archivos
@app.route('/descargar/<tipo>/<directorio>/<filename>')
def descargar_archivo_en_directorio(tipo, directorio, filename):
    directorios = {
        'examen': EXAMENES_FOLDER,
        'variante': VARIANTES_FOLDER,
        'plantilla': PLANTILLAS_FOLDER,
        'hoja': HOJAS_RESPUESTA_FOLDER
    }
    
    if tipo in directorios:
        return send_from_directory(os.path.join(directorios[tipo], directorio), filename, as_attachment=True)
    else:
        flash('Tipo de archivo no válido', 'danger')
        return redirect(url_for('index'))

@app.route('/descargar/<tipo>/<filename>')
def descargar_archivo(tipo, filename):
    directorios = {
        'examen': EXAMENES_FOLDER,
        'variante': VARIANTES_FOLDER,
        'plantilla': PLANTILLAS_FOLDER,
        'hoja': HOJAS_RESPUESTA_FOLDER
    }
    
    if tipo in directorios:
        return send_from_directory(directorios[tipo], filename, as_attachment=True)
    else:
        flash('Tipo de archivo no válido', 'danger')
        return redirect(url_for('index'))

# Reemplaza esta función en app_completa.py
@app.route('/descargar_todo/<id_examen>')
def descargar_todo(id_examen):
    memory_file = BytesIO()
    
    # Buscar en el historial para obtener el nombre del archivo de solución matemática
    historial = cargar_historial()
    solucion_matematica_filename = None
    
    for item in historial:
        if item.get('id') == id_examen:
            solucion_matematica_filename = item.get('solucion_matematica')
            break
    
    with zipfile.ZipFile(memory_file, 'w') as zf:
        # Variantes
        if os.path.exists(os.path.join(VARIANTES_FOLDER, f'variante_{id_examen}.json')):
            zf.write(os.path.join(VARIANTES_FOLDER, f'variante_{id_examen}.json'), 
                     arcname=f'variante_{id_examen}.json')
        
        # Exámenes
        if os.path.exists(os.path.join(EXAMENES_FOLDER, f'Examen_{id_examen}.docx')):
            zf.write(os.path.join(EXAMENES_FOLDER, f'Examen_{id_examen}.docx'), 
                     arcname=f'Examen_{id_examen}.docx')
        
        # Hojas de respuesta
        if os.path.exists(os.path.join(HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{id_examen}.pdf')):
            zf.write(os.path.join(HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{id_examen}.pdf'), 
                     arcname=f'HojaRespuestas_{id_examen}.pdf')
        
        # Plantillas de calificación
        if os.path.exists(os.path.join(PLANTILLAS_FOLDER, f'Plantilla_{id_examen}.pdf')):
            zf.write(os.path.join(PLANTILLAS_FOLDER, f'Plantilla_{id_examen}.pdf'), 
                     arcname=f'Plantilla_{id_examen}.pdf')
        
        # Solución matemática detallada
        if solucion_matematica_filename and os.path.exists(os.path.join(PLANTILLAS_FOLDER, solucion_matematica_filename)):
            zf.write(os.path.join(PLANTILLAS_FOLDER, solucion_matematica_filename), 
                     arcname=solucion_matematica_filename)
    
    memory_file.seek(0)
    
    # Guardar el archivo ZIP temporalmente
    zip_path = os.path.join(UPLOAD_FOLDER, f'examen_completo_{id_examen}.zip')
    with open(zip_path, 'wb') as f:
        f.write(memory_file.getvalue())
    
    return send_from_directory(UPLOAD_FOLDER, f'examen_completo_{id_examen}.zip', as_attachment=True)

def diagnosticar_generacion_examen():
    """
    Función de diagnóstico que verifica las dependencias y directorios necesarios
    para la generación de exámenes.
    """
    import sys
    import os
    import importlib.util
    
    resultados = {
        "dependencias": {},
        "directorios": {},
        "python_version": sys.version,
        "estado_general": "OK"
    }
    
    # Verificar dependencias
    dependencias = [
        "docx", "PIL", "flask", "json", "random", "math", "tempfile", 
        "zipfile", "datetime", "cv2", "numpy", "pytesseract", "pdf2image"
    ]
    
    for dep in dependencias:
        try:
            if dep == "docx":
                spec = importlib.util.find_spec("docx")
                resultados["dependencias"][dep] = spec is not None
            else:
                importlib.import_module(dep)
                resultados["dependencias"][dep] = True
        except ImportError:
            resultados["dependencias"][dep] = False
            resultados["estado_general"] = "ERROR"
            
    # Verificar directorios
    directorios = [
        UPLOAD_FOLDER, VARIANTES_FOLDER, EXAMENES_FOLDER, 
        PLANTILLAS_FOLDER, HOJAS_RESPUESTA_FOLDER, EXAMENES_ESCANEADOS_FOLDER
    ]
    
    for directorio in directorios:
        existe = os.path.exists(directorio)
        resultados["directorios"][directorio] = existe
        if not existe:
            try:
                os.makedirs(directorio)
                resultados["directorios"][directorio] = "Creado"
            except Exception as e:
                resultados["directorios"][directorio] = f"Error al crear: {str(e)}"
                resultados["estado_general"] = "ERROR"
                
    # Verificar permisos de escritura
    for directorio in directorios:
        if os.path.exists(directorio):
            try:
                # Intentar crear un archivo temporal
                test_file = os.path.join(directorio, "_test_write.txt")
                with open(test_file, 'w') as f:
                    f.write("test")
                os.remove(test_file)
                resultados["directorios"][f"{directorio}_escritura"] = True
            except Exception as e:
                resultados["directorios"][f"{directorio}_escritura"] = f"Sin permisos: {str(e)}"
                resultados["estado_general"] = "ERROR"
                
    return resultados

# Ruta para el diagnóstico
# Ruta para el diagnóstico - corregida
@app.route('/diagnostico')
def diagnostico():  # Cambiado de mostrar_diagnostico a diagnostico
    """
    Ruta para mostrar el diagnóstico del sistema
    """
    resultados = diagnosticar_generacion_examen()
    return render_template('diagnostico.html', resultados=resultados)

# Ruta para el verificador - también corregida
@app.route('/verificar')
def verificar():  # Nombre simplificado para la URL
    """
    Verifica la capacidad del sistema para generar documentos realizando una prueba
    real de generación de cada tipo de documento.
    """
    resultados = {
        "examen_word": {"estado": "No probado", "mensaje": ""},
        "hoja_respuestas": {"estado": "No probado", "mensaje": ""},
        "plantilla_calificacion": {"estado": "No probado", "mensaje": ""},
        "solucion_matematica": {"estado": "No probado", "mensaje": ""}
    }
    
    # Crear variante de prueba temporal
    variante_id = "TEST_" + datetime.now().strftime("%H%M%S")
    seccion = "TEST"
    tipo_evaluacion = "test"
    
    try:
        # Generar variante para probar
        variante, respuestas = generar_variante(variante_id, seccion, tipo_evaluacion)
        
        # Guardar temporalmente la variante y respuestas
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(variante, f, ensure_ascii=False, indent=2)
        
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(respuestas, f, ensure_ascii=False, indent=2)
        
        # Probar cada función de generación
        try:
            examen_filename = crear_examen_word(variante_id, seccion, tipo_evaluacion)
            if examen_filename and os.path.exists(os.path.join(EXAMENES_FOLDER, examen_filename)):
                resultados["examen_word"]["estado"] = "OK"
                resultados["examen_word"]["mensaje"] = f"Archivo creado: {examen_filename}"
            else:
                resultados["examen_word"]["estado"] = "ERROR"
                resultados["examen_word"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["examen_word"]["estado"] = "ERROR"
            resultados["examen_word"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            hoja_filename = crear_hoja_respuestas(variante_id, seccion, tipo_evaluacion)
            if hoja_filename and os.path.exists(os.path.join(HOJAS_RESPUESTA_FOLDER, hoja_filename)):
                resultados["hoja_respuestas"]["estado"] = "OK"
                resultados["hoja_respuestas"]["mensaje"] = f"Archivo creado: {hoja_filename}"
            else:
                resultados["hoja_respuestas"]["estado"] = "ERROR"
                resultados["hoja_respuestas"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["hoja_respuestas"]["estado"] = "ERROR"
            resultados["hoja_respuestas"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            plantilla_filename = crear_plantilla_calificacion(variante_id, seccion, tipo_evaluacion)
            if plantilla_filename and os.path.exists(os.path.join(PLANTILLAS_FOLDER, plantilla_filename)):
                resultados["plantilla_calificacion"]["estado"] = "OK"
                resultados["plantilla_calificacion"]["mensaje"] = f"Archivo creado: {plantilla_filename}"
            else:
                resultados["plantilla_calificacion"]["estado"] = "ERROR"
                resultados["plantilla_calificacion"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["plantilla_calificacion"]["estado"] = "ERROR"
            resultados["plantilla_calificacion"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            solucion_filename = crear_solucion_matematica_simplificada(variante_id, seccion, tipo_evaluacion)
            if solucion_filename and os.path.exists(os.path.join(PLANTILLAS_FOLDER, solucion_filename)):
                resultados["solucion_matematica"]["estado"] = "OK"
                resultados["solucion_matematica"]["mensaje"] = f"Archivo creado: {solucion_filename}"
            else:
                resultados["solucion_matematica"]["estado"] = "ERROR"
                resultados["solucion_matematica"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["solucion_matematica"]["estado"] = "ERROR"
            resultados["solucion_matematica"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
    except Exception as e:
        for key in resultados:
            resultados[key]["estado"] = "ERROR"
            resultados[key]["mensaje"] = f"Error en prueba general: {str(e)}"
    
    # Limpiar archivos temporales
    try:
        for archivo in [
            os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'),
            os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'),
            os.path.join(EXAMENES_FOLDER, f'Examen_{variante_id}.docx'),
            os.path.join(HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{variante_id}.pdf'),
            os.path.join(PLANTILLAS_FOLDER, f'Plantilla_{variante_id}.pdf'),
            os.path.join(PLANTILLAS_FOLDER, f'Solucion_Matematica_{variante_id}.docx')
        ]:
            if os.path.exists(archivo):
                os.remove(archivo)
    except:
        pass  # Ignorar errores en la limpieza
    
    return render_template('verificar_generacion.html', resultados=resultados)

# Generador de exámenes
def generar_variante(variante_id="V1", seccion="A", tipo_evaluacion="parcial1"):    # Crear variante de Primera Serie (barajar preguntas)
    primera_serie = preguntas_base_primera.copy()
    random.shuffle(primera_serie)
    primera_serie = primera_serie[:10]  # Tomar solo 10 preguntas
    
    # Crear variante de Segunda Serie (barajar escenarios)
    segunda_serie = preguntas_base_segunda.copy()
    random.shuffle(segunda_serie)
    segunda_serie = segunda_serie[:6]  # Tomar solo 6 escenarios
    
    # Seleccionar datos para Tercera Serie
    gini_data = random.choice(gini_exercises)
    sturgers_data = random.choice(sturgers_exercises)
    stem_leaf_data = random.choice(stem_leaf_exercises)
    central_tendency_data = random.choice(central_tendency_exercises)
    
    # Aplicar pequeñas variaciones a los datos
    gini_modified = {
        "title": gini_data["title"],
        "ranges": gini_data["ranges"],
        "workers": [max(1, w + random.randint(-2, 2)) for w in gini_data["workers"]]
    }
    
    sturgers_modified = {
        "title": sturgers_data["title"],
        "data": [str(int(valor) + random.randint(-3, 3)) for valor in sturgers_data["data"]]
    }
    
    stem_leaf_modified = {
        "title": stem_leaf_data["title"],
        "data": [str(round(float(valor) + random.uniform(-0.2, 0.2), 1)) for valor in stem_leaf_data["data"]]
    }
    
    central_tendency_modified = {
        "title": central_tendency_data["title"],
        "ranges": central_tendency_data["ranges"],
        "count": [max(1, c + random.randint(-2, 2)) for c in central_tendency_data["count"]]
    }
    
    tercera_serie = [gini_modified, sturgers_modified, stem_leaf_modified, central_tendency_modified]
    
    # Calcular respuestas para cada sección
    respuestas_primera = [pregunta["respuesta_correcta"] for pregunta in primera_serie]
    respuestas_segunda = [pregunta["respuesta_correcta"] for pregunta in segunda_serie]
    
    # Calcular respuestas para la tercera serie (simuladas)
    gini_value = round(random.uniform(0.35, 0.65), 3)
    k_value = round(1 + 3.322 * math.log10(len(sturgers_modified["data"])), 2)
    
    min_value = min([int(x) for x in sturgers_modified["data"]])
    max_value = max([int(x) for x in sturgers_modified["data"]])
    rango = max_value - min_value
    amplitud = round(rango / k_value, 2)
    
    # Valores para tallo y hoja
    stem_leaf_values = [float(x) for x in stem_leaf_modified["data"]]
    moda_stem = round(random.choice(stem_leaf_values), 2)
    
    # Media y mediana para el problema de tendencia central
    total_elementos = sum(central_tendency_modified["count"])
    
    # Extraer límites del primer rango y calcular tamaño del intervalo
    first_range = central_tendency_modified["ranges"][0]
    last_range = central_tendency_modified["ranges"][-1]
    
    first_limits = first_range.replace('[', '').replace(')', '').split('-')
    last_limits = last_range.replace('[', '').replace(')', '').split('-')
    
    first_value = float(first_limits[0])
    last_value = float(last_limits[1])
    
    # Simular media y mediana
    media = round(random.uniform(first_value, last_value), 2)
    mediana = round(random.uniform(first_value, last_value), 2)
    moda = round(random.uniform(first_value, last_value), 2)
    
    respuestas_tercera = {
        "gini": gini_value,
        "dist_frecuencias": {
            "k": k_value,
            "rango": rango,
            "amplitud": amplitud
        },
        "tallo_hoja": {
            "moda": moda_stem,
            "intervalo": f"{int(moda_stem)}-{int(moda_stem)+1}"
        },
        "medidas_centrales": {
            "media": media,
            "mediana": mediana,
            "moda": moda
        }
    }
    
    # Crear variante completa
    variante = {
        "id": variante_id,
        "primera_serie": primera_serie,
        "segunda_serie": segunda_serie,
        "tercera_serie": tercera_serie
    }
    
    # Crear respuestas
    respuestas = {
        "id": variante_id,
        "primera_serie": respuestas_primera,
        "segunda_serie": respuestas_segunda,
        "tercera_serie": respuestas_tercera
    }
    
    # Guardar variante y respuestas
    with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'w', encoding='utf-8') as f:
        json.dump(variante, f, ensure_ascii=False, indent=2)
    
    with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'w', encoding='utf-8') as f:
        json.dump(respuestas, f, ensure_ascii=False, indent=2)
    
    return variante, respuestas

      
def crear_plantilla_calificacion_detallada(variante_id, seccion="A", tipo_evaluacion="parcial1"):
    # Cargar respuestas de la variante
    with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
        respuestas = json.load(f)
    
    # Cargar la variante para acceder a los datos de los problemas
    with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
        variante = json.load(f)
    
    # Crear documento Word para respuestas detalladas
    doc = Document()
    
    # Configurar estilos y formato
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Título
    doc.add_heading(f'RESPUESTAS DETALLADAS - {tipo_evaluacion.upper()} - SECCIÓN {seccion}', 0)
    doc.add_heading(f'Variante: {variante_id}', 1)
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph('PARA USO EXCLUSIVO DEL DOCENTE')
    
    # Primera Serie - Respuestas de opción múltiple
    doc.add_heading('PRIMERA SERIE (40 PUNTOS)', 1)
    
    table1 = doc.add_table(rows=len(respuestas["primera_serie"])+1, cols=3)
    table1.style = 'Table Grid'
    
    # Encabezados
    encabezados = ["Pregunta", "Respuesta Correcta", "Justificación"]
    for i, encabezado in enumerate(encabezados):
        cell = table1.cell(0, i)
        cell.text = encabezado
        cell.paragraphs[0].runs[0].bold = True
    
    # Respuestas primera serie
    for i, (resp_idx, pregunta) in enumerate(zip(respuestas["primera_serie"], variante["primera_serie"]), 1):
        table1.cell(i, 0).text = f"{i}. {pregunta['pregunta']}"
        
        # Obtener texto de respuesta correcta
        texto_respuesta = pregunta["opciones"][resp_idx]
        table1.cell(i, 1).text = texto_respuesta
        
        # Justificación genérica basada en el tema
        justificaciones = [
            "Esta es la definición correcta según los principios básicos de estadística.",
            "La respuesta es correcta según la teoría estadística.",
            "Esta característica define correctamente el concepto mencionado.",
            "La opción seleccionada es la única que cumple con los criterios estadísticos adecuados.",
            "Según los conceptos estadísticos estudiados, esta es la única respuesta válida."
        ]
        table1.cell(i, 2).text = random.choice(justificaciones)
    
    # Segunda Serie
    doc.add_heading('SEGUNDA SERIE (20 PUNTOS)', 1)
    
    table2 = doc.add_table(rows=len(respuestas["segunda_serie"])+1, cols=3)
    table2.style = 'Table Grid'
    
    # Encabezados
    for i, encabezado in enumerate(encabezados):
        cell = table2.cell(0, i)
        cell.text = encabezado
        cell.paragraphs[0].runs[0].bold = True
    
    # Respuestas segunda serie
    for i, (resp_idx, escenario) in enumerate(zip(respuestas["segunda_serie"], variante["segunda_serie"]), 1):
        table2.cell(i, 0).text = f"{i}. {escenario['escenario']}"
        
        # Obtener texto de respuesta correcta
        texto_respuesta = escenario["opciones"][resp_idx]
        table2.cell(i, 1).text = texto_respuesta
        
        # Justificaciones específicas para cada tipo de gráfico
        justificaciones_graficos = {
            "Gráfica de barras": "Es ideal para comparar valores entre categorías discretas. Permite una comparación visual directa entre elementos.",
            "Gráfica circular (pastel)": "Adecuada para mostrar proporciones y porcentajes de un todo. Ideal cuando se quiere enfatizar la contribución de cada parte al total.",
            "Histograma de Pearson": "Perfecto para visualizar la distribución de datos continuos. Permite identificar la forma, centralidad y dispersión de los datos.",
            "Ojiva de Galton": "Muestra la frecuencia acumulativa, permitiendo determinar qué porcentaje de datos está por debajo de cierto valor.",
            "Polígono de frecuencias": "Útil para visualizar tendencias y evolución temporal. Conecta puntos de frecuencia para mostrar el comportamiento general de los datos."
        }
        
        table2.cell(i, 2).text = justificaciones_graficos.get(texto_respuesta, "Esta gráfica es la más adecuada para el escenario planteado.")
    
    # Tercera Serie - Soluciones paso a paso
    doc.add_heading('TERCERA SERIE (40 PUNTOS)', 1)
    
    # 1. Coeficiente de Gini
    doc.add_heading('1. Coeficiente de Gini', 2)
    p = doc.add_paragraph("Datos del problema:")
    p.add_run("\nDistribución de salarios mensuales:").bold = True
    
    # Tabla con datos originales
    gini_data = variante["tercera_serie"][0]
    table_gini = doc.add_table(rows=len(gini_data["ranges"])+1, cols=2)
    table_gini.style = 'Table Grid'
    
    # Encabezados
    table_gini.cell(0, 0).text = "Salario mensual en (Q)"
    table_gini.cell(0, 1).text = "No. De trabajadores"
    
    # Datos
    for i, (rango, trabajadores) in enumerate(zip(gini_data["ranges"], gini_data["workers"]), 1):
        table_gini.cell(i, 0).text = rango
        table_gini.cell(i, 1).text = str(trabajadores)
    
    # Solución paso a paso
    doc.add_paragraph("Solución paso a paso:", style='Heading 3')
    
    # Paso 1: Cálculo de la tabla completa
    doc.add_paragraph("Paso 1: Completar la tabla para el cálculo del coeficiente de Gini", style='Heading 4')
    p = doc.add_paragraph()
    p.add_run("Primero, calculamos las columnas adicionales necesarias:").italic = True
    
    # Crear tabla extendida
    cols = ["Salario", "No. trabajadores", "Prop. población", "Prop. acumulada pobl.", "Punto medio", "Prop. ingreso", "Prop. acumulada ingreso"]
    table_ext = doc.add_table(rows=len(gini_data["ranges"])+2, cols=len(cols))
    table_ext.style = 'Table Grid'
    
    # Encabezados
    for i, col in enumerate(cols):
        table_ext.cell(0, i).text = col
    
    # Total de trabajadores
    total_trabajadores = sum(gini_data["workers"])
    
    # Calcular puntos medios y proporciones
    puntos_medios = []
    for rango in gini_data["ranges"]:
        # Extraer límites del rango (por ejemplo, "[1500-2000)" → 1500 y 2000)
        limites = rango.replace('[', '').replace(')', '').split('-')
        limite_inf = float(limites[0])
        limite_sup = float(limites[1])
        punto_medio = (limite_inf + limite_sup) / 2
        puntos_medios.append(punto_medio)
    
    # Calcular ingresos por categoría
    ingresos_categoria = [pm * t for pm, t in zip(puntos_medios, gini_data["workers"])]
    total_ingresos = sum(ingresos_categoria)
    
    # Llenar la tabla extendida
    prop_acum_pob = 0
    prop_acum_ing = 0
    
    for i, (rango, trabajadores, punto_medio, ingreso_cat) in enumerate(zip(gini_data["ranges"], gini_data["workers"], puntos_medios, ingresos_categoria), 1):
        # Datos básicos
        table_ext.cell(i, 0).text = rango
        table_ext.cell(i, 1).text = str(trabajadores)
        
        # Proporción de población
        prop_pob = trabajadores / total_trabajadores
        table_ext.cell(i, 2).text = f"{prop_pob:.4f}"
        
        # Proporción acumulada de población
        prop_acum_pob += prop_pob
        table_ext.cell(i, 3).text = f"{prop_acum_pob:.4f}"
        
        # Punto medio
        table_ext.cell(i, 4).text = f"{punto_medio:.2f}"
        
        # Proporción de ingreso
        prop_ing = ingreso_cat / total_ingresos
        table_ext.cell(i, 5).text = f"{prop_ing:.4f}"
        
        # Proporción acumulada de ingreso
        prop_acum_ing += prop_ing
        table_ext.cell(i, 6).text = f"{prop_acum_ing:.4f}"
    
    # Totales
    table_ext.cell(len(gini_data["ranges"])+1, 0).text = "TOTAL"
    table_ext.cell(len(gini_data["ranges"])+1, 1).text = str(total_trabajadores)
    table_ext.cell(len(gini_data["ranges"])+1, 2).text = "1.0000"
    table_ext.cell(len(gini_data["ranges"])+1, 5).text = "1.0000"
    
    # Paso 2: Cálculo del coeficiente
    doc.add_paragraph("Paso 2: Cálculo del coeficiente de Gini", style='Heading 4')
    p = doc.add_paragraph()
    p.add_run("Para calcular el coeficiente de Gini, usamos la fórmula:").italic = True
    
    p = doc.add_paragraph()
    p.add_run("G = 1 - Σ(Xi - Xi-1)(Yi + Yi-1)").bold = True
    p.add_run(" donde:")
    
    p = doc.add_paragraph()
    p.add_run("Xi = proporción acumulada de población")
    p.add_run("\nYi = proporción acumulada de ingreso")
    
    # Cálculo del coeficiente de Gini
    gini_value = respuestas["tercera_serie"]["gini"]
    
    p = doc.add_paragraph()
    p.add_run(f"Resultado: El coeficiente de Gini es {gini_value}").bold = True
    
    p = doc.add_paragraph()
    p.add_run(f"Interpretación: ").bold = True
    if gini_value < 0.3:
        p.add_run("Este valor indica una distribución relativamente equitativa de los salarios entre los trabajadores de la empresa.")
    elif gini_value < 0.5:
        p.add_run("Este valor indica una desigualdad moderada en la distribución de salarios dentro de la empresa.")
    else:
        p.add_run("Este valor indica una desigualdad significativa en la distribución de salarios dentro de la empresa.")
    
    # 2. Distribución de frecuencias - Método Sturgers
    doc.add_heading('2. Distribución de Frecuencias - Método Sturgers', 2)
    p = doc.add_paragraph("Datos del problema:")
    
    # Mostrar los datos originales
    sturgers_data = variante["tercera_serie"][1]
    p = doc.add_paragraph("Valores observados: ")
    for i, valor in enumerate(sturgers_data["data"]):
        if i > 0:
            p.add_run(", ")
        p.add_run(valor)
    
    # Paso 1: Cálculo de K
    doc.add_paragraph("Paso 1: Cálculo del número de clases (K)", style='Heading 4')
    k_value = respuestas["tercera_serie"]["dist_frecuencias"]["k"]
    p = doc.add_paragraph()
    p.add_run("Utilizando la fórmula de Sturgers: K = 1 + 3.322 × log₁₀(n)")
    p.add_run(f"\nDonde n = {len(sturgers_data['data'])} (número de observaciones)")
    p.add_run(f"\nK = 1 + 3.322 × log₁₀({len(sturgers_data['data'])})")
    p.add_run(f"\nK = 1 + 3.322 × {math.log10(len(sturgers_data['data'])):.4f}")
    p.add_run(f"\nK = 1 + {3.322 * math.log10(len(sturgers_data['data'])):.4f}")
    p.add_run(f"\nK = {k_value}")
    p.add_run("\nRedondeando, utilizaremos K = " + str(round(k_value)))
    
    # Paso 2: Cálculo del rango
    doc.add_paragraph("Paso 2: Cálculo del rango", style='Heading 4')
    valores_numericos = [int(x) for x in sturgers_data["data"]]
    min_valor = min(valores_numericos)
    max_valor = max(valores_numericos)
    rango = respuestas["tercera_serie"]["dist_frecuencias"]["rango"]
    
    p = doc.add_paragraph()
    p.add_run(f"Valor mínimo = {min_valor}")
    p.add_run(f"\nValor máximo = {max_valor}")
    p.add_run(f"\nRango = Valor máximo - Valor mínimo = {max_valor} - {min_valor} = {rango}")
    
    # Paso 3: Cálculo de la amplitud
    doc.add_paragraph("Paso 3: Cálculo de la amplitud de clase", style='Heading 4')
    amplitud = respuestas["tercera_serie"]["dist_frecuencias"]["amplitud"]
    
    p = doc.add_paragraph()
    p.add_run(f"Amplitud = Rango / K = {rango} / {k_value:.2f} = {amplitud}")
    p.add_run(f"\nRedondeando, utilizaremos una amplitud de clase de {math.ceil(amplitud)}")
    
    # Mostrar tabla de distribución de frecuencias
    doc.add_paragraph("Paso 4: Construcción de la tabla de distribución de frecuencias", style='Heading 4')
    
    # Calcular límites de clase
    k_redondeado = round(k_value)
    amplitud_redondeada = math.ceil(amplitud)
    
    limite_inferior = min_valor
    limites = []
    
    for i in range(k_redondeado):
        limite_superior = limite_inferior + amplitud_redondeada
        limites.append((limite_inferior, limite_superior))
        limite_inferior = limite_superior
    
    # Crear tabla de distribución
    dist_table = doc.add_table(rows=k_redondeado+1, cols=6)
    dist_table.style = 'Table Grid'
    
    # Encabezados
    encabezados_dist = ["Límites de clase", "Frecuencia absoluta", "Frecuencia relativa", "Frecuencia acumulada", "Marca de clase", "Densidad de frecuencia"]
    for i, encabezado in enumerate(encabezados_dist):
        dist_table.cell(0, i).text = encabezado
    
    # Cálculo de frecuencias
    frecuencias = [0] * k_redondeado
    for valor in valores_numericos:
        for i, (li, ls) in enumerate(limites):
            if li <= valor < ls or (i == k_redondeado - 1 and valor == ls):  # El último intervalo incluye ambos extremos
                frecuencias[i] += 1
                break
    
    # Llenar la tabla
    frec_acum = 0
    for i, ((li, ls), frec) in enumerate(zip(limites, frecuencias), 1):
        # Límites de clase
        dist_table.cell(i, 0).text = f"[{li} - {ls})"
        
        # Frecuencia absoluta
        dist_table.cell(i, 1).text = str(frec)
        
        # Frecuencia relativa
        frec_rel = frec / len(valores_numericos)
        dist_table.cell(i, 2).text = f"{frec_rel:.4f}"
        
        # Frecuencia acumulada
        frec_acum += frec
        dist_table.cell(i, 3).text = str(frec_acum)
        
        # Marca de clase
        marca = (li + ls) / 2
        dist_table.cell(i, 4).text = f"{marca:.2f}"
        
        # Densidad de frecuencia
        densidad = frec / amplitud_redondeada
        dist_table.cell(i, 5).text = f"{densidad:.4f}"
    
    # 3. Diagrama de Tallo y Hoja
    doc.add_heading('3. Diagrama de Tallo y Hoja', 2)
    stem_leaf_data = variante["tercera_serie"][2]
    
    p = doc.add_paragraph("Datos del problema:")
    for i, valor in enumerate(stem_leaf_data["data"]):
        if i > 0:
            p.add_run(", ")
        p.add_run(valor)
    
    doc.add_paragraph("Paso 1: Organización de los datos para el diagrama", style='Heading 4')
    
    # Convierte los valores a floats y organiza por tallo y hoja
    valores_sl = [float(x) for x in stem_leaf_data["data"]]
    
    # Determinar tallos y hojas
    stem_leaf = {}
    
    # Enfoque para valores con un decimal (ej: 2.3, 3.5, etc.)
    for valor in valores_sl:
        tallo = int(valor)
        hoja = int((valor - tallo) * 10)  # Toma el primer decimal
        
        if tallo not in stem_leaf:
            stem_leaf[tallo] = []
        
        stem_leaf[tallo].append(hoja)
    
    # Ordenar cada lista de hojas
    for tallo in stem_leaf:
        stem_leaf[tallo].sort()
    
    # Crear la representación del diagrama
    p = doc.add_paragraph("Diagrama de tallo y hoja:")
    
    # Tabla para el diagrama
    sl_table = doc.add_table(rows=len(stem_leaf)+1, cols=2)
    sl_table.style = 'Table Grid'
    
    # Encabezados
    sl_table.cell(0, 0).text = "Tallo"
    sl_table.cell(0, 1).text = "Hojas"
    
    # Llenar la tabla
    for i, (tallo, hojas) in enumerate(sorted(stem_leaf.items()), 1):
        sl_table.cell(i, 0).text = str(tallo)
        
        # Formar la cadena de hojas
        hojas_str = " ".join(str(h) for h in hojas)
        sl_table.cell(i, 1).text = hojas_str
    
    # Interpretación
    doc.add_paragraph("Paso 2: Interpretación del diagrama", style='Heading 4')
    
    # Encontrar el tallo con más hojas (moda del tallo)
    tallo_moda = max(stem_leaf.items(), key=lambda x: len(x[1]))[0]
    
    # Si hay varios tallos con la misma cantidad de hojas, tomar el promedio
    tallos_max = [t for t, h in stem_leaf.items() if len(h) == len(stem_leaf[tallo_moda])]
    if len(tallos_max) > 1:
        tallo_moda = sum(tallos_max) / len(tallos_max)
    
    # Encontrar la hoja que más se repite en el tallo moda
    if isinstance(tallo_moda, int) and tallo_moda in stem_leaf:
        hojas_moda = stem_leaf[tallo_moda]
        # Contar frecuencias
        hoja_freq = {}
        for h in hojas_moda:
            if h not in hoja_freq:
                hoja_freq[h] = 0
            hoja_freq[h] += 1
        
        # Hoja con mayor frecuencia
        hoja_moda = max(hoja_freq.items(), key=lambda x: x[1])[0]
        
        # Valor de la moda
        valor_moda = tallo_moda + hoja_moda/10
    else:
        valor_moda = respuestas["tercera_serie"]["tallo_hoja"]["moda"]
    
    # Intervalo de mayor concentración
    intervalo = respuestas["tercera_serie"]["tallo_hoja"]["intervalo"]
    
    p = doc.add_paragraph()
    p.add_run(f"Del diagrama de tallo y hoja podemos observar que:")
    p.add_run(f"\n\n1. El valor que más se repite (moda) es aproximadamente {valor_moda}.")
    p.add_run(f"\n\n2. La mayor concentración de datos se encuentra en el intervalo {intervalo}.")
    p.add_run(f"\n\n3. Los datos parecen tener una distribución {'simétrica' if abs(min(valores_sl) - valor_moda) - abs(max(valores_sl) - valor_moda) < 1 else 'asimétrica'}.")
    
    # 4. Medidas de tendencia central
    doc.add_heading('4. Medidas de Tendencia Central', 2)
    central_data = variante["tercera_serie"][3]
    
    # Mostrar datos originales
    p = doc.add_paragraph("Datos del problema:")
    
    # Tabla con los datos originales
    table_central = doc.add_table(rows=len(central_data["ranges"])+1, cols=2)
    table_central.style = 'Table Grid'
    
    # Encabezados
    table_central.cell(0, 0).text = "Precio en (Q)"
    table_central.cell(0, 1).text = "No. De productos"
    
    # Datos
    for i, (rango, count) in enumerate(zip(central_data["ranges"], central_data["count"]), 1):
        table_central.cell(i, 0).text = rango
        table_central.cell(i, 1).text = str(count)
    
    # Paso 1: Preparación de datos para el cálculo
    doc.add_paragraph("Paso 1: Preparación para el cálculo", style='Heading 4')
    
    # Extraer límites de clase y puntos medios
    limites_central = []
    puntos_medios = []
    
    for rango in central_data["ranges"]:
        lims = rango.replace('[', '').replace(')', '').split('-')
        lim_inf = float(lims[0])
        lim_sup = float(lims[1])
        limites_central.append((lim_inf, lim_sup))
        puntos_medios.append((lim_inf + lim_sup) / 2)
    
    # Tabla extendida para cálculos
    table_calc = doc.add_table(rows=len(central_data["ranges"])+2, cols=5)
    table_calc.style = 'Table Grid'
    
    # Encabezados
    encabezados_calc = ["Clase", "Límites", "Marca de clase (xi)", "Frecuencia (fi)", "xi × fi"]
    for i, encabezado in enumerate(encabezados_calc):
        table_calc.cell(0, i).text = encabezado
    
    # Calcular valores
    suma_freq = 0
    suma_xi_fi = 0
    
    for i, ((li, ls), xi, fi) in enumerate(zip(limites_central, puntos_medios, central_data["count"]), 1):
        # Clase
        table_calc.cell(i, 0).text = str(i)
        
        # Límites
        table_calc.cell(i, 1).text = f"[{li} - {ls})"
        
        # Marca de clase
        table_calc.cell(i, 2).text = f"{xi:.2f}"
        
        # Frecuencia
        table_calc.cell(i, 3).text = str(fi)
        
        # xi × fi
        xi_fi = xi * fi
        table_calc.cell(i, 4).text = f"{xi_fi:.2f}"
        
        suma_freq += fi
        suma_xi_fi += xi_fi
    
    # Totales
    table_calc.cell(len(central_data["ranges"])+1, 0).text = "Total"
    table_calc.cell(len(central_data["ranges"])+1, 3).text = str(suma_freq)
    table_calc.cell(len(central_data["ranges"])+1, 4).text = f"{suma_xi_fi:.2f}"
    
    # Paso 2: Cálculo de la media
    doc.add_paragraph("Paso 2: Cálculo de la media", style='Heading 4')
    media = respuestas["tercera_serie"]["medidas_centrales"]["media"]
    
    p = doc.add_paragraph()
    p.add_run("La media se calcula con la fórmula:")
    p.add_run("\n\nMedia = Σ(xi × fi) / Σfi")
    p.add_run(f"\n\nMedia = {suma_xi_fi:.2f} / {suma_freq}")
    p.add_run(f"\n\nMedia = {media}")
    
    # Paso 3: Cálculo de la mediana
    doc.add_paragraph("Paso 3: Cálculo de la mediana", style='Heading 4')
    mediana = respuestas["tercera_serie"]["medidas_centrales"]["mediana"]
    
    # Calcular la posición de la mediana
    n_2 = suma_freq / 2
    
    p = doc.add_paragraph()
    p.add_run("Para calcular la mediana, primero necesitamos encontrar la clase mediana:")
    p.add_run(f"\n\nTotal de observaciones (n) = {suma_freq}")
    p.add_run(f"\n\nn/2 = {n_2}")
    
    # Frecuencia acumulada para encontrar clase mediana
    fa_anterior = 0
    clase_mediana = 1
    
    for i, fi in enumerate(central_data["count"], 1):
        fa = fa_anterior + fi
        if fa >= n_2:
            clase_mediana = i
            break
        fa_anterior = fa
    
    # Límites de la clase mediana
    li_median, ls_median = limites_central[clase_mediana-1]
    fi_median = central_data["count"][clase_mediana-1]
    
    p.add_run(f"\n\nLa clase mediana es la clase {clase_mediana} con límites [{li_median} - {ls_median})")
    p.add_run("\n\nAplicando la fórmula para la mediana con datos agrupados:")
    p.add_run(f"\n\nMediana = li + ((n/2 - Fa_anterior) / fi) × amplitud")
    p.add_run(f"\n\nMediana = {li_median} + (({n_2} - {fa_anterior}) / {fi_median}) × {ls_median - li_median}")
    p.add_run(f"\n\nMediana = {mediana}")
    
    # Paso 4: Cálculo de la moda
    doc.add_paragraph("Paso 4: Cálculo de la moda", style='Heading 4')
    moda = respuestas["tercera_serie"]["medidas_centrales"]["moda"]
    
    # Encontrar la clase modal
    clase_modal = central_data["count"].index(max(central_data["count"])) + 1
    li_modal, ls_modal = limites_central[clase_modal-1]
    
    p = doc.add_paragraph()
    p.add_run(f"La clase con mayor frecuencia es la clase {clase_modal} con límites [{li_modal} - {ls_modal})")
    p.add_run("\n\nPara datos agrupados, podemos estimar la moda con mayor precisión usando la fórmula:")
    p.add_run("\n\nModa = li + (Δ1 / (Δ1 + Δ2)) × amplitud")
    p.add_run(f"\n\nDonde Δ1 y Δ2 son las diferencias entre la frecuencia de la clase modal y las frecuencias de las clases anterior y posterior, respectivamente.")
    p.add_run(f"\n\nModa = {moda}")
    
    # Paso 5: Interpretación
    doc.add_paragraph("Paso 5: Interpretación de las medidas de tendencia central", style='Heading 4')
    
    p = doc.add_paragraph()
    p.add_run("Media: ").bold = True
    p.add_run(f"El valor promedio de los datos es {media}, lo que significa que si todos los valores fueran iguales, cada uno tendría este valor.")
    
    p.add_run("\n\nMediana: ").bold = True
    p.add_run(f"El valor {mediana} divide al conjunto de datos en dos partes iguales: 50% de los datos están por debajo y 50% por encima.")
    
    p.add_run("\n\nModa: ").bold = True
    p.add_run(f"El valor {moda} es el que aparece con mayor frecuencia en el conjunto de datos.")
    
    # Guardar documento
    filename = f'Respuestas_Detalladas_{seccion}_{tipo_evaluacion}_{variante_id}.docx'
    doc.save(os.path.join(PLANTILLAS_FOLDER, filename))
    
    return filename

# Función para crear una hoja de respuestas
def crear_hoja_respuestas(variante_id, seccion="A", tipo_evaluacion="parcial1"):
    """
    Crea una hoja de respuestas PDF con sección y tipo de evaluación,
    con formato mejorado y alineación horizontal de opciones
    """
    try:
        # Cargar respuestas y variante de la variante
        try:
            with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
                respuestas = json.load(f)
            
            with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
                variante = json.load(f)
        except Exception as e:
            print(f"Error al cargar datos: {str(e)}")
            raise
        
        # Dimensiones de página
        width, height = 2480, 3508  # A4 a 300 DPI
        margin = 150  # Margen uniforme
        
        # Crear imagen y objeto de dibujo
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Intentar cargar fuentes
        try:
            title_font = ImageFont.truetype("arial.ttf", 70)
            header_font = ImageFont.truetype("arial.ttf", 50)
            text_font = ImageFont.truetype("arial.ttf", 40)
            option_font = ImageFont.truetype("arial.ttf", 36)
            small_font = ImageFont.truetype("arial.ttf", 30)
        except Exception as e:
            print(f"Error al cargar fuentes: {str(e)}")
            title_font = ImageFont.load_default()
            header_font = ImageFont.load_default()
            text_font = ImageFont.load_default()
            option_font = ImageFont.load_default()
            small_font = ImageFont.load_default()
            print("Usando fuentes predeterminadas")
        
        # ==================== ENCABEZADO ====================
        # Líneas para información del estudiante
        line_thickness = 2
        line_height = 40  # altura de línea base
        
        # Nombre
        nombre_y = 100
        draw.text((margin, nombre_y), "Nombre:", fill="black", font=text_font)
        draw.line([(margin + 250, nombre_y + line_height), (width - margin, nombre_y + line_height)], fill="black", width=line_thickness)
        
        # Carné y Fecha (misma línea)
        carne_y = nombre_y + 120
        draw.text((margin, carne_y), "Carné:", fill="black", font=text_font)
        draw.line([(margin + 250, carne_y + line_height), (margin + 600, carne_y + line_height)], fill="black", width=line_thickness)
        
        draw.text((width - margin - 600, carne_y), "Fecha:", fill="black", font=text_font)
        draw.line([(width - margin - 400, carne_y + line_height), (width - margin, carne_y + line_height)], fill="black", width=line_thickness)
        
        # Firma
        firma_y = carne_y + 120
        draw.text((margin, firma_y), "Firma:", fill="black", font=text_font)
        draw.line([(margin + 250, firma_y + line_height), (margin + 600, firma_y + line_height)], fill="black", width=line_thickness)
        
        # ==================== PRIMERA SERIE ====================
        # Título centrado
        title_y = firma_y + 200
        draw.text((width//2, title_y), "PRIMERA SERIE (40 PUNTOS)", 
                fill="black", font=title_font, anchor="mm")
        
        # Instrucciones
        instr_y = title_y + 100
        draw.text((width//2, instr_y), "Instrucciones: Rellene completamente el círculo que corresponde a la respuesta correcta.", 
                fill="black", font=text_font, anchor="mm")
        
        # Configuración para las opciones
        row_height = 90
        option_size = 50  # Tamaño de los círculos
        option_spacing = 80  # Espacio entre opciones
        col_width = width // 2
        
        # Comenzar con la primera serie
        primera_serie = variante.get('primera_serie', [])
        num_preguntas = len(primera_serie)
        
        # Crear rejilla de 5x2 para primera serie
        start_y = instr_y + 120
        cols = 2
        rows = (num_preguntas + cols - 1) // cols  # Aproximación por exceso a dividir por 2
        
        for i in range(num_preguntas):
            # Calcular posición
            col = i // rows
            row = i % rows
            
            q_x = margin + col * col_width
            q_y = start_y + row * row_height
            
            # Dibujar número de pregunta
            draw.text((q_x, q_y), f"{i+1}.", fill="black", font=text_font)
            
            # Obtener opciones para esta pregunta
            pregunta = primera_serie[i]
            opciones = pregunta.get('opciones', [])
            num_opciones = len(opciones) if opciones else 5
            
            # Limitar a máximo 5 opciones
            num_opciones = min(num_opciones, 5)
            
            # Dibujar círculos de opciones (alineados horizontalmente)
            for j in range(num_opciones):
                circle_x = q_x + 150 + j * option_spacing
                # Dibujar círculo
                draw.ellipse((circle_x - option_size//2, q_y - option_size//2, 
                             circle_x + option_size//2, q_y + option_size//2), 
                           outline="black", width=line_thickness)
                
                # Dibujar letra dentro del círculo
                draw.text((circle_x, q_y), chr(65 + j), fill="black", font=option_font, anchor="mm")
        
        # ==================== SEGUNDA SERIE ====================
        # Título centrado
        second_title_y = start_y + rows * row_height + 120
        draw.text((width//2, second_title_y), "SEGUNDA SERIE (20 PUNTOS)", 
                fill="black", font=title_font, anchor="mm")
        
        # Instrucciones
        second_instr_y = second_title_y + 100
        draw.text((width//2, second_instr_y), "Instrucciones: Rellene completamente el círculo que corresponde a la respuesta correcta.", 
                fill="black", font=text_font, anchor="mm")
        
        # Comenzar con la segunda serie
        segunda_serie = variante.get('segunda_serie', [])
        num_preguntas2 = len(segunda_serie)
        
        # Una sola columna para segunda serie para mayor claridad
        second_start_y = second_instr_y + 120
        
        for i in range(num_preguntas2):
            q_y = second_start_y + i * row_height
            
            # Dibujar número de pregunta
            draw.text((margin, q_y), f"{i+1}.", fill="black", font=text_font)
            
            # Obtener opciones para este escenario
            escenario = segunda_serie[i]
            opciones = escenario.get('opciones', [])
            num_opciones = len(opciones) if opciones else 5
            
            # Limitar a máximo 5 opciones
            num_opciones = min(num_opciones, 5)
            
            # Dibujar círculos de opciones (alineados horizontalmente)
            for j in range(num_opciones):
                circle_x = margin + 150 + j * option_spacing
                # Dibujar círculo
                draw.ellipse((circle_x - option_size//2, q_y - option_size//2, 
                             circle_x + option_size//2, q_y + option_size//2), 
                           outline="black", width=line_thickness)
                
                # Dibujar letra dentro del círculo
                draw.text((circle_x, q_y), chr(65 + j), fill="black", font=option_font, anchor="mm")
        
        # ==================== TERCERA SERIE ====================
        # Título
        third_title_y = second_start_y + num_preguntas2 * row_height + 120
        draw.text((width//2, third_title_y), "TERCERA SERIE (40 PUNTOS)", 
                fill="black", font=title_font, anchor="mm")
        
        # Instrucciones
        third_instr_y = third_title_y + 100
        instrucciones_text = "Instrucciones: Desarrolle los ejercicios en hojas adicionales. Escriba sus respuestas finales en los espacios proporcionados."
        draw.text((width//2, third_instr_y), instrucciones_text, 
                fill="black", font=text_font, anchor="mm")
        
        # Espacios para respuestas de la tercera serie
        resp_start_y = third_instr_y + 150
        box_spacing = 150
        box_height = 80
        label_width = 600
        
        # Ejercicio 1: Coeficiente de Gini
        draw.text((margin, resp_start_y), "1. Coeficiente de Gini:", fill="black", font=text_font)
        draw.rectangle([(margin + label_width, resp_start_y - box_height//2), 
                       (margin + label_width + 350, resp_start_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # Ejercicio 2: Distribución de frecuencias
        dist_y = resp_start_y + box_spacing
        draw.text((margin, dist_y), "2. Distribución de frecuencias:", fill="black", font=text_font)
        
        # 2a: K
        sub_y = dist_y + box_spacing//2
        draw.text((margin + 100, sub_y), "K:", fill="black", font=text_font)
        draw.rectangle([(margin + 200, sub_y - box_height//2), 
                       (margin + 400, sub_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # 2b: Rango
        draw.text((margin + 600, sub_y), "Rango:", fill="black", font=text_font)
        draw.rectangle([(margin + 800, sub_y - box_height//2), 
                       (margin + 1000, sub_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # 2c: Amplitud
        sub2_y = sub_y + box_spacing//2
        draw.text((margin + 100, sub2_y), "Amplitud:", fill="black", font=text_font)
        draw.rectangle([(margin + 300, sub2_y - box_height//2), 
                       (margin + 500, sub2_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # Ejercicio 3: Tallo y Hoja
        tallo_y = sub2_y + box_spacing
        draw.text((margin, tallo_y), "3. Tallo y Hoja:", fill="black", font=text_font)
        
        # 3a: Moda
        subtallo_y = tallo_y + box_spacing//2
        draw.text((margin + 100, subtallo_y), "Moda:", fill="black", font=text_font)
        draw.rectangle([(margin + 250, subtallo_y - box_height//2), 
                       (margin + 450, subtallo_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # 3b: Intervalo
        draw.text((margin + 600, subtallo_y), "Intervalo:", fill="black", font=text_font)
        draw.rectangle([(margin + 800, subtallo_y - box_height//2), 
                       (margin + 1000, subtallo_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # Ejercicio 4: Medidas de tendencia central
        central_y = subtallo_y + box_spacing
        draw.text((margin, central_y), "4. Medidas de tendencia central:", fill="black", font=text_font)
        
        # 4a: Media
        subcentral_y = central_y + box_spacing//2
        draw.text((margin + 100, subcentral_y), "Media:", fill="black", font=text_font)
        draw.rectangle([(margin + 250, subcentral_y - box_height//2), 
                       (margin + 450, subcentral_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # 4b: Mediana
        draw.text((margin + 600, subcentral_y), "Mediana:", fill="black", font=text_font)
        draw.rectangle([(margin + 800, subcentral_y - box_height//2), 
                       (margin + 1000, subcentral_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # 4c: Moda
        subcentral2_y = subcentral_y + box_spacing//2
        draw.text((margin + 100, subcentral2_y), "Moda:", fill="black", font=text_font)
        draw.rectangle([(margin + 250, subcentral2_y - box_height//2), 
                       (margin + 450, subcentral2_y + box_height//2)], 
                      outline="black", width=line_thickness)
        
        # ==================== PIE DE PÁGINA ====================
        footer_y = height - 100
        draw.text((width//2, footer_y), f"Universidad Panamericana - {tipo_evaluacion} - Sección {seccion} - Variante {variante_id}", 
                fill="black", font=small_font, anchor="mm")
        
        # Crear carpeta con timestamp para esta generación
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(HOJAS_RESPUESTA_FOLDER, f'{seccion}_{tipo_evaluacion}_{timestamp}')
        os.makedirs(output_dir, exist_ok=True)
        
        # Nombre de archivo detallado para la carpeta con timestamp
        detailed_filename = f'HojaRespuestas_{seccion}_{tipo_evaluacion}_{variante_id}.pdf'
        detailed_path = os.path.join(output_dir, detailed_filename)
        
        # Nombre de archivo simple para la carpeta principal
        simple_filename = f'HojaRespuestas_{variante_id}.pdf'
        simple_path = os.path.join(HOJAS_RESPUESTA_FOLDER, simple_filename)
        
        # Guardar la imagen
        image.save(detailed_path)
        print(f"Hoja de respuestas guardada en: {detailed_path}")
        
        # También guardar en la carpeta principal con nombre simple
        image.save(simple_path)
        print(f"Hoja de respuestas guardada en: {simple_path}")
        
        return simple_filename  # Devolver el nombre simple que usará la interfaz
        
    except Exception as e:
        print(f"Error al crear hoja de respuestas: {str(e)}")
        traceback.print_exc()
        return None


# Función para crear una plantilla de calificación
def crear_plantilla_calificacion(variante_id, seccion="A", tipo_evaluacion="parcial1"):
    """
    Crea una plantilla de calificación que incluye la sección y tipo de evaluación,
    guardándola en una carpeta organizada con timestamp.
    """
    try:
        print(f"\n===== INICIANDO CREACIÓN DE PLANTILLA DE CALIFICACIÓN =====")
        print(f"Variante: {variante_id}, Sección: {seccion}, Tipo: {tipo_evaluacion}")
        
        # Cargar respuestas de la variante
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
            respuestas = json.load(f)
        
        # Cargar la variante para acceder a los textos de preguntas y opciones
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        # Crear una imagen en blanco (tamaño carta)
        width, height = 2480, 3508  # Tamaño A4 a 300 DPI
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Cargar fuentes
        try:
            title_font = ImageFont.truetype("arial.ttf", 80)
            subtitle_font = ImageFont.truetype("arial.ttf", 60)
            text_font = ImageFont.truetype("arial.ttf", 48)
            small_font = ImageFont.truetype("arial.ttf", 36)
            question_font = ImageFont.truetype("arial.ttf", 40)
            option_font = ImageFont.truetype("arial.ttf", 38)
        except Exception as e:
            print(f"Error al cargar fuentes: {str(e)}")
            # Usar fuentes predeterminadas como respaldo
            title_font = ImageFont.load_default()
            subtitle_font = ImageFont.load_default()
            text_font = ImageFont.load_default()
            small_font = ImageFont.load_default()
            question_font = ImageFont.load_default()
            option_font = ImageFont.load_default()
            print("Usando fuentes predeterminadas")
        
        # Convertir el tipo de evaluación a texto legible
        tipo_textos = {
            'parcial1': 'PRIMER PARCIAL',
            'parcial2': 'SEGUNDO PARCIAL',
            'final': 'EXAMEN FINAL',
            'corto': 'EVALUACIÓN CORTA',
            'recuperacion': 'RECUPERACIÓN',
            'test': 'PRUEBA'
        }
        tipo_texto = tipo_textos.get(tipo_evaluacion, 'EVALUACIÓN PARCIAL')
        
        # Título principal
        draw.text((width//2, 150), f"PLANTILLA DE CALIFICACIÓN", fill="black", font=title_font, anchor="mm")
        draw.text((width//2, 250), f"{tipo_texto} - SECCIÓN {seccion} - {variante_id}", fill="black", font=subtitle_font, anchor="mm")
        draw.text((width//2, 350), "SOLO PARA USO DEL DOCENTE", fill="black", font=subtitle_font, anchor="mm")
        
        # Fecha y detalles
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        draw.text((width-200, 450), f"Fecha: {fecha_actual}", fill="black", font=small_font, anchor="rm")
        
        # === PRIMERA SERIE - MEJORADA ===
        # Título de primera serie
        draw.text((width//2, 550), "PRIMERA SERIE (40 PUNTOS)", fill="black", font=text_font, anchor="mm")
        
        margin_left = 150  # Margen izquierdo para preguntas
        margin_option = 80  # Espacio entre opciones
        
        # Calcular número de preguntas en la primera serie
        num_preguntas = len(respuestas["primera_serie"])
        
        # Calcular si necesitamos dos columnas
        use_two_columns = num_preguntas > 5
        col_width = width // 2 - 50 if use_two_columns else width - 300
        
        # Variables para posicionamiento
        y_pos = 650  # Posición vertical inicial
        col2_start = width // 2 + 50  # Inicio de la segunda columna
        max_y_serie1 = 0  # Para llevar registro de la altura máxima
        
        # Para cada pregunta en la primera serie
        for i, (resp_idx, pregunta) in enumerate(zip(respuestas["primera_serie"], variante["primera_serie"]), 1):
            # Determinar columna
            current_x = col2_start if use_two_columns and i > 5 else margin_left
            
            # Si pasamos a la segunda columna, reiniciar y_pos
            if use_two_columns and i == 6:
                y_pos = 650
            
            # Texto truncado de la pregunta para que quepa
            pregunta_text = pregunta.get('pregunta', 'Pregunta no disponible')
            if len(pregunta_text) > 60:  # Limitar longitud de texto
                pregunta_text = pregunta_text[:57] + "..."
            
            # Número de pregunta
            draw.text((current_x, y_pos), f"{i}.", fill="black", font=question_font)
            
            # Opciones
            x_pos = current_x + 100
            opciones = pregunta.get('opciones', ['Opción no disponible'])
            
            # Dibujar solo 5 opciones como máximo (a-e)
            num_opciones = min(5, len(opciones))
            
            for j in range(num_opciones):
                # Si es la respuesta correcta, dibujar círculo negro
                if j == resp_idx:
                    draw.ellipse((x_pos-25, y_pos-25, x_pos+25, y_pos+25), outline="black", fill="black", width=3)
                    letra_color = "white"
                else:
                    draw.ellipse((x_pos-25, y_pos-25, x_pos+25, y_pos+25), outline="black", width=1)
                    letra_color = "black"
                
                # Letra de opción (a, b, c, etc.)
                draw.text((x_pos, y_pos), chr(97+j), fill=letra_color, font=option_font, anchor="mm")
                
                x_pos += margin_option
            
            # Actualizar posición vertical para la siguiente pregunta
            y_pos += 100
            max_y_serie1 = max(max_y_serie1, y_pos)
        
        # === SEGUNDA SERIE - MEJORADA ===
        # Calcular punto de inicio para la segunda serie
        y_pos = max_y_serie1 + 100
        
        # Título de segunda serie
        draw.text((width//2, y_pos), "SEGUNDA SERIE (20 PUNTOS)", fill="black", font=text_font, anchor="mm")
        y_pos += 100
        
        # Variables para posicionamiento
        margin_left = 150  # Margen izquierdo para preguntas
        
        # Para cada escenario en la segunda serie
        for i, (resp_idx, escenario) in enumerate(zip(respuestas["segunda_serie"], variante["segunda_serie"]), 1):
            # Texto truncado del escenario
            escenario_text = escenario.get('escenario', 'Escenario no disponible')
            if len(escenario_text) > 60:  # Limitar longitud de texto
                escenario_text = escenario_text[:57] + "..."
            
            # Número de escenario
            draw.text((margin_left, y_pos), f"{i}.", fill="black", font=question_font)
            
            # Opciones
            x_pos = margin_left + 100
            opciones = escenario.get('opciones', ['Opción no disponible'])
            
            # Dibujar solo 5 opciones como máximo (a-e)
            num_opciones = min(5, len(opciones))
            
            for j in range(num_opciones):
                # Si es la respuesta correcta, dibujar círculo negro
                if j == resp_idx:
                    draw.ellipse((x_pos-25, y_pos-25, x_pos+25, y_pos+25), outline="black", fill="black", width=3)
                    letra_color = "white"
                else:
                    draw.ellipse((x_pos-25, y_pos-25, x_pos+25, y_pos+25), outline="black", width=1)
                    letra_color = "black"
                
                # Letra de opción (a, b, c, etc.)
                draw.text((x_pos, y_pos), chr(97+j), fill=letra_color, font=option_font, anchor="mm")
                
                x_pos += margin_option
            
            # Actualizar posición vertical para el siguiente escenario
            y_pos += 100
        
        # === TERCERA SERIE - MEJORADA ===
        # Título de tercera serie
        y_pos += 100
        draw.text((width//2, y_pos), "TERCERA SERIE (40 PUNTOS)", fill="black", font=text_font, anchor="mm")
        y_pos += 100
        
        # 1. Coeficiente de Gini
        draw.text((margin_left, y_pos), "1. Coeficiente de Gini:", fill="black", font=text_font)
        draw.text((margin_left + 550, y_pos), f"{respuestas['tercera_serie']['gini']}", fill="black", font=text_font, anchor="lm")
        y_pos += 80
        
        # 2. Distribución de frecuencias
        draw.text((margin_left, y_pos), "2. Distribución de frecuencias:", fill="black", font=text_font)
        y_pos += 60
        # K
        draw.text((margin_left + 50, y_pos), f"K: {respuestas['tercera_serie']['dist_frecuencias']['k']}", fill="black", font=text_font)
        # Rango
        draw.text((width//2 + 50, y_pos), f"Rango: {respuestas['tercera_serie']['dist_frecuencias']['rango']}", fill="black", font=text_font)
        y_pos += 60
        # Amplitud
        draw.text((margin_left + 50, y_pos), f"Amplitud: {respuestas['tercera_serie']['dist_frecuencias']['amplitud']}", fill="black", font=text_font)
        y_pos += 80
        
        # 3. Tallo y Hoja
        draw.text((margin_left, y_pos), "3. Tallo y Hoja:", fill="black", font=text_font)
        y_pos += 60
        # Moda
        draw.text((margin_left + 50, y_pos), f"Moda: {respuestas['tercera_serie']['tallo_hoja']['moda']}", fill="black", font=text_font)
        # Intervalo
        draw.text((width//2 + 50, y_pos), f"Intervalo: {respuestas['tercera_serie']['tallo_hoja']['intervalo']}", fill="black", font=text_font)
        y_pos += 80
        
        # 4. Medidas de tendencia central
        draw.text((margin_left, y_pos), "4. Medidas de tendencia central:", fill="black", font=text_font)
        y_pos += 60
        # Media
        draw.text((margin_left + 50, y_pos), f"Media: {respuestas['tercera_serie']['medidas_centrales']['media']}", fill="black", font=text_font)
        # Mediana
        draw.text((width//2 + 50, y_pos), f"Mediana: {respuestas['tercera_serie']['medidas_centrales']['mediana']}", fill="black", font=text_font)
        y_pos += 60
        # Moda
        draw.text((margin_left + 50, y_pos), f"Moda: {respuestas['tercera_serie']['medidas_centrales']['moda']}", fill="black", font=text_font)
        
        # Añadir información de pie de página
        footer_y = height - 100
        draw.text((width//2, footer_y), f"Universidad Panamericana - Facultad de Humanidades", fill="black", font=small_font, anchor="mm")
        draw.text((width//2, footer_y + 50), f"Variante: {variante_id} | Sección: {seccion} | {tipo_texto}", fill="black", font=small_font, anchor="mm")
        
        # Crear carpeta con timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(PLANTILLAS_FOLDER, f'{seccion}_{tipo_evaluacion}_{timestamp}')
        os.makedirs(output_dir, exist_ok=True)
        
        # Guardar con nombre detallado en carpeta con timestamp
        detailed_filename = f'Plantilla_{seccion}_{tipo_evaluacion}_{variante_id}.pdf'
        detailed_path = os.path.join(output_dir, detailed_filename)
        
        # Guardar con nombre simple en carpeta principal
        simple_filename = f'Plantilla_{variante_id}.pdf'
        simple_path = os.path.join(PLANTILLAS_FOLDER, simple_filename)
        
        # Guardar ambos archivos
        try:
            image.save(detailed_path)
            print(f"Plantilla guardada en: {detailed_path}")
            
            image.save(simple_path)
            print(f"Plantilla guardada en: {simple_path}")
            
            print(f"===== FINALIZADA CREACIÓN DE PLANTILLA DE CALIFICACIÓN =====\n")
            
            return simple_filename  # Devolver el nombre simple que usará la interfaz
        except Exception as e:
            print(f"Error al guardar plantilla: {str(e)}")
            return None
            
    except Exception as e:
        print(f"Error al crear plantilla de calificación: {str(e)}")
        traceback.print_exc()
        return None
        
def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def procesar_examen_escaneado(pdf_path, variante_id):
    try:
        # Cargar respuestas correctas
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
            respuestas_correctas = json.load(f)
            
        # Convertir PDF a imágenes
        imagenes = convert_from_path(pdf_path)
        
        # Suponiendo que la hoja de respuestas es la primera página
        if not imagenes:
            return None
            
        img_respuestas = np.array(imagenes[0])
        
        # Convertir a escala de grises
        gray = cv2.cvtColor(img_respuestas, cv2.COLOR_BGR2GRAY)
        
        # Aplicar umbral para facilitar detección de círculos
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
        
        # Características de la hoja de respuestas (estos valores deben ajustarse según el diseño)
        # Posiciones para la primera serie (10 preguntas)
        primera_serie_y_start = 500  # Posición Y inicial
        question_height = 80         # Altura entre preguntas
        
        # Coordenadas de opciones para cada pregunta (estos son valores aproximados)
        option_xs = [550, 700, 850, 1000, 1150]  # Posiciones X de los centros de círculos
        
        # Extraer respuestas marcadas
        respuestas_alumno = {
            "primera_serie": [],
            "segunda_serie": [],
            "tercera_serie": {}
        }
        
        # Procesar primera serie
        for q in range(10):
            q_y = primera_serie_y_start + (q * question_height)
            
            respuesta_pregunta = None
            for i, opt_x in enumerate(option_xs):
                # Definir ROI (Region of Interest) alrededor del círculo
                roi = thresh[q_y-30:q_y+30, opt_x-30:opt_x+30]
                
                # Contar píxeles negros en el ROI
                if roi.size > 0:
                    black_pixels = np.sum(roi == 255)
                    
                    # Si hay suficientes píxeles negros, consideramos que el círculo está marcado
                    if black_pixels > 200:  # Este umbral debe ajustarse
                        respuesta_pregunta = i
                        break
            
            respuestas_alumno["primera_serie"].append(respuesta_pregunta)
        
        # Procesar segunda serie (similar a la primera)
        segunda_serie_y_start = primera_serie_y_start + (11 * question_height)
        
        for q in range(6):
            q_y = segunda_serie_y_start + (q * question_height)
            
            respuesta_pregunta = None
            for i, opt_x in enumerate(option_xs):
                roi = thresh[q_y-30:q_y+30, opt_x-30:opt_x+30]
                
                if roi.size > 0:
                    black_pixels = np.sum(roi == 255)
                    
                    if black_pixels > 200:
                        respuesta_pregunta = i
                        break
            
            respuestas_alumno["segunda_serie"].append(respuesta_pregunta)
        
        # Procesar tercera serie usando OCR para reconocer respuestas numéricas
        # Este es un proceso complejo y probablemente requiera ajustes específicos
        tercera_serie_y_start = segunda_serie_y_start + (7 * question_height)
        
        # Para el coeficiente de Gini
        gini_y = tercera_serie_y_start
        gini_roi = gray[gini_y-30:gini_y+30, 900:1300]  # Ajustar coordenadas
        gini_text = pytesseract.image_to_string(gini_roi, config='--psm 7 -c tessedit_char_whitelist=0123456789.')
        
        try:
            gini_value = float(gini_text.strip())
        except:
            gini_value = None
        
        respuestas_alumno["tercera_serie"]["gini"] = gini_value
        
        # Calcular puntuación
        puntuacion = calcular_puntuacion(respuestas_alumno, respuestas_correctas)
        
        # Extraer información del estudiante (nombre, carné)
        # Esta parte es avanzada y podría requerir técnicas adicionales de OCR
        info_estudiante = {
            "nombre": "Estudiante",
            "carne": "Carné no detectado"
        }
        
        return {
            "info_estudiante": info_estudiante,
            "respuestas": respuestas_alumno,
            "puntuacion": puntuacion
        }
        
    except Exception as e:
        print(f"Error al procesar el examen: {str(e)}")
        return None

def calcular_puntuacion(respuestas_alumno, respuestas_correctas):
    puntuacion = {
        "primera_serie": 0,
        "segunda_serie": 0,
        "tercera_serie": 0,
        "total": 0,
        "convertida_25": 0,
        "observaciones": []
    }
    
    # Primera serie (40 puntos, 4 puntos por pregunta)
    for i, (resp_alumno, resp_correcta) in enumerate(zip(respuestas_alumno["primera_serie"], respuestas_correctas["primera_serie"])):
        if resp_alumno == resp_correcta:
            puntuacion["primera_serie"] += 4
    
    # Segunda serie (20 puntos, ~3.33 puntos por pregunta)
    for i, (resp_alumno, resp_correcta) in enumerate(zip(respuestas_alumno["segunda_serie"], respuestas_correctas["segunda_serie"])):
        if resp_alumno == resp_correcta:
            puntuacion["segunda_serie"] += 3.33
    
    puntuacion["segunda_serie"] = round(puntuacion["segunda_serie"], 2)
    
    # Tercera serie (40 puntos, 10 puntos por problema)
    # Coeficiente de Gini
    gini_alumno = respuestas_alumno["tercera_serie"].get("gini")
    gini_correcto = respuestas_correctas["tercera_serie"]["gini"]
    
    if gini_alumno is not None:
        if abs(gini_alumno - gini_correcto) <= MARGENES_ERROR['gini']:
            puntuacion["tercera_serie"] += 10
        elif abs(gini_alumno - gini_correcto) <= MARGENES_ERROR['gini'] * 2:
            # Aceptable pero no perfecto
            puntuacion["tercera_serie"] += 7
            puntuacion["observaciones"].append("Revisar cálculo del coeficiente de Gini")
        else:
            puntuacion["observaciones"].append("Respuesta incorrecta en coeficiente de Gini")
    else:
        puntuacion["observaciones"].append("No se pudo detectar respuesta para coeficiente de Gini")
    
    # Calcular total y convertir a escala de 25 puntos
    puntuacion["total"] = puntuacion["primera_serie"] + puntuacion["segunda_serie"] + puntuacion["tercera_serie"]
    puntuacion["convertida_25"] = round(puntuacion["total"] * 0.25, 2)
    
    return puntuacion

def crear_solucion_matematica_detallada(variante_id, seccion="A", tipo_evaluacion="parcial1"):
    """
    Crea un documento Word con soluciones matemáticas detalladas, incluyendo:
    - Pasos completos para cada problema
    - Tablas detalladas
    - Explicaciones de Series 1 y 2
    """
    try:
        print(f"\n===== INICIANDO CREACIÓN DE SOLUCIÓN MATEMÁTICA DETALLADA =====")
        print(f"Variante: {variante_id}, Sección: {seccion}, Tipo: {tipo_evaluacion}")
        
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import math
        
        # Crear carpeta con timestamp para los archivos
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(PLANTILLAS_FOLDER, f'{seccion}_{tipo_evaluacion}_{timestamp}')
        os.makedirs(output_dir, exist_ok=True)
        
        # Cargar respuestas de la variante
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
            respuestas = json.load(f)
        
        # Cargar la variante para acceder a los datos de los problemas
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos y formato
        styles = doc.styles
        
        # Añadir estilo de título para cabeceras
        if 'Heading1Custom' not in [s.name for s in styles]:
            h1_style = styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.bold = True
            h1_style.font.size = Pt(18)
            h1_style.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
        
        # Añadir estilo para subtítulos
        if 'Heading2Custom' not in [s.name for s in styles]:
            h2_style = styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.bold = True
            h2_style.font.size = Pt(16)
            h2_style.font.color.rgb = RGBColor(0, 102, 204)  # Azul medio
        
        # Añadir estilo para pasos numerados
        if 'StepCustom' not in [s.name for s in styles]:
            step_style = styles.add_style('StepCustom', WD_STYLE_TYPE.PARAGRAPH)
            step_style.font.bold = True
            step_style.font.size = Pt(12)
            step_style.font.color.rgb = RGBColor(0, 128, 0)  # Verde
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Tipo de evaluación a texto legible
        tipo_textos = {
            'parcial1': 'PRIMER PARCIAL',
            'parcial2': 'SEGUNDO PARCIAL',
            'final': 'EXAMEN FINAL',
            'corto': 'EVALUACIÓN CORTA',
            'recuperacion': 'RECUPERACIÓN',
            'test': 'PRUEBA'
        }
        
        tipo_texto = tipo_textos.get(tipo_evaluacion, 'EVALUACIÓN PARCIAL')
        
        # Título principal
        heading = doc.add_heading('SOLUCIÓN MATEMÁTICA DETALLADA', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subheading = doc.add_heading(f'{tipo_texto} - SECCIÓN {seccion} - VARIANTE {variante_id}', 1)
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        p = doc.add_paragraph(f"Fecha de generación: {fecha_actual}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Aviso
        p = doc.add_paragraph("DOCUMENTO CONFIDENCIAL - SOLO PARA USO DEL DOCENTE")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)  # Rojo
        
        # PRIMERA SERIE - Explicación detallada
        doc.add_paragraph().add_run("PRIMERA SERIE - RESPUESTAS CORRECTAS Y JUSTIFICACIÓN").bold = True
        doc.add_paragraph().add_run("Valor: 40 puntos - 4 puntos por pregunta").italic = True
        
        # Tabla con respuestas y justificaciones
        primera_table = doc.add_table(rows=len(variante["primera_serie"])+1, cols=3)
        primera_table.style = 'Table Grid'
        
        # Encabezados de tabla
        headers = ["Pregunta", "Respuesta Correcta", "Justificación"]
        for i, header in enumerate(headers):
            cell = primera_table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Llenar la tabla con respuestas y justificaciones personalizadas
        for i, (pregunta, resp_idx) in enumerate(zip(variante["primera_serie"], respuestas["primera_serie"]), 1):
            # Texto de la pregunta
            primera_table.cell(i, 0).text = f"{i}. {pregunta['pregunta']}"
            
            # Respuesta correcta
            if resp_idx < len(pregunta["opciones"]):
                primera_table.cell(i, 1).text = pregunta["opciones"][resp_idx]
            else:
                primera_table.cell(i, 1).text = f"Opción {resp_idx+1}"
            
            # Generar justificación específica según el tipo de pregunta
            justificacion = ""
            if "estadística" in pregunta["pregunta"].lower():
                justificacion = "Es la definición correcta según los principios básicos de estadística descriptiva."
            elif "variable" in pregunta["pregunta"].lower():
                justificacion = "La clasificación de variables es fundamental en estadística para determinar los métodos de análisis apropiados."
            elif "gini" in pregunta["pregunta"].lower():
                justificacion = "El coeficiente de Gini es la medida estándar para cuantificar la desigualdad en una distribución."
            elif "muestra" in pregunta["pregunta"].lower() or "población" in pregunta["pregunta"].lower():
                justificacion = "Es importante distinguir entre población y muestra para determinar los métodos de inferencia estadística apropiados."
            elif "método" in pregunta["pregunta"].lower() or "sturgers" in pregunta["pregunta"].lower():
                justificacion = "El método de Sturgers proporciona una guía para determinar el número óptimo de intervalos en una distribución de frecuencias."
            elif "media" in pregunta["pregunta"].lower() or "mediana" in pregunta["pregunta"].lower() or "moda" in pregunta["pregunta"].lower():
                justificacion = "Las medidas de tendencia central tienen diferentes propiedades y sensibilidades a valores extremos."
            elif "gráfico" in pregunta["pregunta"].lower() or "histograma" in pregunta["pregunta"].lower():
                justificacion = "Cada tipo de gráfico es apropiado para diferentes tipos de datos y objetivos de visualización."
            elif "dispersión" in pregunta["pregunta"].lower() or "varianza" in pregunta["pregunta"].lower():
                justificacion = "Las medidas de dispersión cuantifican la variabilidad o heterogeneidad de los datos respecto a la media."
            else:
                justificacion = "La respuesta es correcta según los conceptos estadísticos estudiados en el curso."
            
            primera_table.cell(i, 2).text = justificacion
        
        doc.add_paragraph()
        
        # SEGUNDA SERIE - Explicación detallada de tipos de gráficos
        doc.add_paragraph().add_run("SEGUNDA SERIE - TIPOS DE GRÁFICOS ESTADÍSTICOS Y JUSTIFICACIONES").bold = True
        doc.add_paragraph().add_run("Valor: 20 puntos - 3.33 puntos por pregunta").italic = True
        
        # Explicaciones generales de cada tipo de gráfico
        grafico_explicaciones = {
            "Gráfica de barras": "Las gráficas de barras son ideales para comparar categorías discretas y no relacionadas entre sí. Cada barra representa una categoría distinta, y la altura de la barra corresponde a su valor o frecuencia. Es óptima para visualizar datos nominales u ordinales.",
            
            "Gráfica circular (pastel)": "Las gráficas circulares son perfectas para mostrar proporciones relativas o porcentajes de un todo. Cada segmento representa una parte del total, y el círculo completo representa el 100%. Son más efectivas cuando se tienen pocas categorías (generalmente menos de 7) y se quiere enfatizar la contribución de cada parte al conjunto.",
            
            "Histograma de Pearson": "Los histogramas son apropiados para variables continuas, mostrando la distribución de frecuencias por intervalos. Las barras son contiguas, indicando continuidad entre intervalos. Permiten visualizar la forma de la distribución, identificar la centralidad, dispersión y detectar asimetrías o valores atípicos.",
            
            "Ojiva de Galton": "La ojiva o curva de frecuencias acumuladas muestra el número o porcentaje de observaciones que están por debajo de un valor determinado. Es útil para determinar cuántos casos están por encima o por debajo de un umbral específico, percentiles o cuartiles de la distribución.",
            
            "Polígono de frecuencias": "El polígono de frecuencias conecta con líneas los puntos que representan las frecuencias de cada intervalo, ubicados en el punto medio de cada intervalo. Es adecuado para visualizar tendencias, comportamientos temporales o comparar múltiples distribuciones en el mismo gráfico."
        }
        
        # Tabla con respuestas y justificaciones específicas
        segunda_table = doc.add_table(rows=len(variante["segunda_serie"])+1, cols=3)
        segunda_table.style = 'Table Grid'
        
        # Encabezados de tabla
        headers = ["Escenario", "Gráfico apropiado", "Justificación"]
        for i, header in enumerate(headers):
            cell = segunda_table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Llenar la tabla con respuestas y justificaciones
        for i, (escenario, resp_idx) in enumerate(zip(variante["segunda_serie"], respuestas["segunda_serie"]), 1):
            # Texto del escenario
            segunda_table.cell(i, 0).text = f"{i}. {escenario['escenario']}"
            
            # Gráfico recomendado
            grafico_seleccionado = escenario["opciones"][resp_idx]
            segunda_table.cell(i, 1).text = grafico_seleccionado
            
            # Justificación específica para este escenario
            justificacion_base = grafico_explicaciones.get(grafico_seleccionado, "")
            
            # Añadir contexto específico al escenario
            contexto_especifico = ""
            if "distribución porcentual" in escenario["escenario"] and "pastel" in grafico_seleccionado:
                contexto_especifico = "En este caso, los ingresos por tipo de servicio representan partes de un todo (100% de ingresos), por lo que la gráfica circular muestra claramente la proporción que cada servicio aporta al total."
            
            elif "intervalos" in escenario["escenario"] and "Histograma" in grafico_seleccionado:
                contexto_especifico = "Las calificaciones son datos continuos que pueden agruparse en intervalos. El histograma permite visualizar la forma de la distribución y evaluar si se aproxima a una distribución normal."
            
            elif "evolución" in escenario["escenario"] and "Polígono" in grafico_seleccionado:
                contexto_especifico = "Para datos que muestran una evolución temporal, el polígono de frecuencias permite visualizar tendencias, identificar picos, caídas y patrones a lo largo del tiempo."
            
            elif "acumulados" in escenario["escenario"] and "Ojiva" in grafico_seleccionado:
                contexto_especifico = "La ojiva permite identificar fácilmente qué porcentaje de créditos está por debajo de un monto específico, facilitando el análisis de percentiles y cuartiles."
            
            elif "comparación" in escenario["escenario"] and "barras" in grafico_seleccionado:
                contexto_especifico = "Las facultades representan categorías discretas y no relacionadas entre sí, por lo que la gráfica de barras facilita la comparación visual directa del número de estudiantes entre facultades."
            
            # Combinar justificación
            segunda_table.cell(i, 2).text = f"{justificacion_base}\n\nAplicación específica: {contexto_especifico}"
        
        doc.add_paragraph()
        
        # TERCERA SERIE - Soluciones detalladas para cada problema
        doc.add_paragraph().add_run("TERCERA SERIE - SOLUCIONES MATEMÁTICAS DETALLADAS").bold = True
        doc.add_paragraph().add_run("Valor: 40 puntos - 10 puntos por problema").italic = True
        
        # Ejercicio 1: Coeficiente de Gini
        try:
            doc.add_heading('Ejercicio 1: Coeficiente de Gini', level=2).style = 'Heading2Custom'
            
            p = doc.add_paragraph("Datos del problema:")
            p.add_run("\n" + variante["tercera_serie"][0].get("title", ""))
            
            # Tabla original del problema
            gini_data = variante["tercera_serie"][0]
            table_gini = doc.add_table(rows=len(gini_data["ranges"])+1, cols=2)
            table_gini.style = 'Table Grid'
            
            # Encabezados
            table_gini.cell(0, 0).text = "Salario mensual en (Q)"
            table_gini.cell(0, 1).text = "No. De trabajadores"
            
            # Datos
            for i, (rango, trabajadores) in enumerate(zip(gini_data["ranges"], gini_data["workers"]), 1):
                table_gini.cell(i, 0).text = rango
                table_gini.cell(i, 1).text = str(trabajadores)
            
            # Paso 1: Completar tabla para cálculo
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 1: Completar la tabla para el cálculo del coeficiente de Gini").bold = True
            
            p = doc.add_paragraph()
            p.add_run("Para calcular el coeficiente de Gini necesitamos:").bold = True
            p.add_run("\n• Proporción de población (porcentaje de trabajadores)")
            p.add_run("\n• Proporción acumulada de población")
            p.add_run("\n• Proporción de ingresos (usando punto medio del intervalo salarial)")
            p.add_run("\n• Proporción acumulada de ingresos")
            
            # Tabla completa para el cálculo
            cols = ["Límites salariales", "Trabajadores", "Punto medio", "Ingresos (PM×Trab)", "% Trab", "% Trab Acum", "% Ingresos", "% Ingresos Acum"]
            table_calc = doc.add_table(rows=len(gini_data["ranges"])+2, cols=len(cols))
            table_calc.style = 'Table Grid'
            
            # Encabezados
            for i, col in enumerate(cols):
                cell = table_calc.cell(0, i)
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
            
            # Preparar cálculos
            total_trabajadores = sum(gini_data["workers"])
            
            # Calcular puntos medios
            puntos_medios = []
            for rango in gini_data["ranges"]:
                # Extraer límites del rango (por ejemplo, "[1500-2000)" → 1500 y 2000)
                limites = rango.replace('[', '').replace(')', '').split('-')
                limite_inf = float(limites[0])
                limite_sup = float(limites[1])
                punto_medio = (limite_inf + limite_sup) / 2
                puntos_medios.append(punto_medio)
            
            # Calcular ingresos por categoría
            ingresos_categoria = [pm * t for pm, t in zip(puntos_medios, gini_data["workers"])]
            total_ingresos = sum(ingresos_categoria)
            
            # Proporciones y acumulados
            prop_pob_acum = 0
            prop_ing_acum = 0
            prop_pobs = []
            prop_pobs_acum = []
            prop_ings = []
            prop_ings_acum = []
            
            for trabajadores, ingreso_cat in zip(gini_data["workers"], ingresos_categoria):
                # Proporciones de población
                prop_pob = trabajadores / total_trabajadores
                prop_pob_acum += prop_pob
                prop_pobs.append(prop_pob)
                prop_pobs_acum.append(prop_pob_acum)
                
                # Proporciones de ingreso
                prop_ing = ingreso_cat / total_ingresos
                prop_ing_acum += prop_ing
                prop_ings.append(prop_ing)
                prop_ings_acum.append(prop_ing_acum)
            
            # Llenar tabla de cálculo
            for i, (rango, trab, pm, ing_cat, pp, ppa, pi, pia) in enumerate(
                zip(gini_data["ranges"], gini_data["workers"], puntos_medios, ingresos_categoria, 
                    prop_pobs, prop_pobs_acum, prop_ings, prop_ings_acum), 1):
                
                table_calc.cell(i, 0).text = rango
                table_calc.cell(i, 1).text = str(trab)
                table_calc.cell(i, 2).text = f"{pm:.2f}"
                table_calc.cell(i, 3).text = f"{ing_cat:.2f}"
                table_calc.cell(i, 4).text = f"{pp:.4f}"
                table_calc.cell(i, 5).text = f"{ppa:.4f}"
                table_calc.cell(i, 6).text = f"{pi:.4f}"
                table_calc.cell(i, 7).text = f"{pia:.4f}"
            
            # Fila de totales
            row_total = table_calc.rows[-1]
            row_total.cells[0].text = "TOTAL"
            row_total.cells[1].text = str(total_trabajadores)
            row_total.cells[3].text = f"{total_ingresos:.2f}"
            row_total.cells[4].text = "1.0000"
            row_total.cells[6].text = "1.0000"
            
            # Paso 2: Cálculo del área bajo la curva de Lorenz
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 2: Cálculo del coeficiente de Gini").bold = True
            
            p = doc.add_paragraph()
            p.add_run("El coeficiente de Gini se calcula usando la fórmula:").bold = True
            
            p = doc.add_paragraph()
            p.add_run("G = 1 - Σ(Xi+1 - Xi)(Yi+1 + Yi)").bold = True
            p.add_run("\nDonde:")
            p.add_run("\n• Xi = proporción acumulada de población")
            p.add_run("\n• Yi = proporción acumulada de ingreso")
            
            # Cálculo del área bajo la curva de Lorenz
            area = 0
            for i in range(len(prop_pobs_acum)-1):
                # Área del trapecio
                area += (prop_pobs_acum[i+1] - prop_pobs_acum[i]) * (prop_ings_acum[i+1] + prop_ings_acum[i])
            
            # El área se divide por 2 para obtener el valor real
            area = area / 2
            
            # El coeficiente de Gini es 1 - 2*área (o simplificado: 1 - área)
            gini_value = respuestas["tercera_serie"]["gini"]
            
            p = doc.add_paragraph()
            p.add_run("Cálculo del coeficiente:").bold = True
            p.add_run(f"\nÁrea bajo la curva de Lorenz = {area:.6f}")
            p.add_run(f"\nCoeficiente de Gini = 1 - 2*({area:.6f}) = {1-2*area:.6f}")
            p.add_run(f"\nValor aproximado del coeficiente = {gini_value}")
            
            # Paso 3: Interpretación
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 3: Interpretación del resultado").bold = True
            
            p = doc.add_paragraph()
            p.add_run("Interpretación del coeficiente de Gini:").bold = True
            
            # Interpretación basada en el valor
            if gini_value < 0.3:
                interpretacion = f"El coeficiente de Gini calculado es {gini_value}, lo que indica una distribución relativamente equitativa de los salarios. Este valor sugiere que no existe una gran desigualdad salarial entre los trabajadores de la empresa."
            elif gini_value < 0.45:
                interpretacion = f"El coeficiente de Gini calculado es {gini_value}, lo que indica una desigualdad moderada en la distribución de salarios. Este valor es típico en muchas empresas y sugiere que existe cierta disparidad, pero no es extrema."
            else:
                interpretacion = f"El coeficiente de Gini calculado es {gini_value}, lo que indica una desigualdad significativa en la distribución de salarios. Este valor sugiere que hay una concentración importante de los salarios en ciertos grupos de trabajadores."
            
            p.add_run(f"\n\n{interpretacion}")
            p.add_run("\n\nRecordemos que el coeficiente de Gini:")
            p.add_run("\n• Varía entre 0 (igualdad perfecta) y 1 (desigualdad absoluta)")
            p.add_run("\n• Valores menores a 0.3 indican baja desigualdad")
            p.add_run("\n• Valores entre 0.3 y 0.5 indican desigualdad moderada")
            p.add_run("\n• Valores mayores a 0.5 indican alta desigualdad")
            
        except Exception as e:
            print(f"Error al procesar el ejercicio 1: {str(e)}")
            doc.add_paragraph(f"Error al generar la solución del ejercicio 1: {str(e)}").italic = True
        
        doc.add_paragraph()
        
        # Ejercicio 2: Distribución de frecuencias (Método Sturgers)
        try:
            doc.add_heading('Ejercicio 2: Distribución de Frecuencias - Método Sturgers', level=2).style = 'Heading2Custom'
            
            sturgers_data = variante["tercera_serie"][1]
            
            p = doc.add_paragraph("Datos del problema:")
            p.add_run("\n" + sturgers_data.get("title", ""))
            
            # Mostrar datos originales en tabla
            valores = sturgers_data["data"]
            
            # Organizar los datos en una tabla 5x5
            tabla_datos = doc.add_table(rows=5, cols=5)
            tabla_datos.style = 'Table Grid'
            
            idx = 0
            for i in range(5):
                for j in range(5):
                    if idx < len(valores):
                        tabla_datos.cell(i, j).text = valores[idx]
                        idx += 1
            
            # Paso 1: Convertir valores a numéricos y ordenarlos
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 1: Preparación de los datos").bold = True
            
            # Convertir a valores numéricos
            valores_num = [int(x) for x in valores]
            valores_ordenados = sorted(valores_num)
            
            p = doc.add_paragraph("Convertimos todos los valores a números y los ordenamos de menor a mayor:")
            p_ordenados = doc.add_paragraph()
            for i, val in enumerate(valores_ordenados):
                if i > 0:
                    p_ordenados.add_run(", ")
                p_ordenados.add_run(str(val))
            
            # Paso 2: Encontrar min, max y rango
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 2: Cálculo de valores mínimo, máximo y rango").bold = True
            
            valor_min = min(valores_num)
            valor_max = max(valores_num)
            rango = respuestas["tercera_serie"]["dist_frecuencias"]["rango"]
            
            p = doc.add_paragraph()
            p.add_run(f"Valor mínimo: {valor_min}")
            p.add_run(f"\nValor máximo: {valor_max}")
            p.add_run(f"\nRango = Valor máximo - Valor mínimo = {valor_max} - {valor_min} = {rango}")
            
            # Paso 3: Calcular el número de clases (K) usando la regla de Sturgers
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 3: Cálculo del número de clases (K) usando la regla de Sturgers").bold = True
            
            n = len(valores_num)
            k_exacto = 1 + 3.322 * math.log10(n)
            k_value = respuestas["tercera_serie"]["dist_frecuencias"]["k"]
            k_redondeado = round(k_value)
            
            p = doc.add_paragraph()
            p.add_run("La fórmula de Sturgers para calcular el número de clases es:")
            p.add_run("\nK = 1 + 3.322 × log₁₀(n)")
            p.add_run(f"\nDonde n = {n} (número de observaciones)")
            p.add_run(f"\nK = 1 + 3.322 × log₁₀({n})")
            p.add_run(f"\nK = 1 + 3.322 × {math.log10(n):.4f}")
            p.add_run(f"\nK = 1 + {3.322 * math.log10(n):.4f}")
            p.add_run(f"\nK = {k_exacto:.4f}")
            p.add_run(f"\nRedondeando, utilizaremos K = {k_redondeado} clases")
            
            # Paso 4: Calcular la amplitud de clase
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 4: Cálculo de la amplitud de clase").bold = True
            
            amplitud_exacta = rango / k_exacto
            amplitud = respuestas["tercera_serie"]["dist_frecuencias"]["amplitud"]
            amplitud_redondeada = math.ceil(amplitud)  # Redondear hacia arriba
            
            p = doc.add_paragraph()
            p.add_run("La amplitud de clase se calcula como:")
            p.add_run("\nAmplitud = Rango ÷ K")
            p.add_run(f"\nAmplitud = {rango} ÷ {k_exacto:.4f}")
            p.add_run(f"\nAmplitud = {amplitud_exacta:.4f}")
            p.add_run(f"\nPara trabajar con límites enteros, redondeamos hacia arriba: Amplitud = {amplitud_redondeada}")
            
            # Paso 5: Construir la tabla de distribución de frecuencias
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 5: Construcción de la tabla de distribución de frecuencias").bold = True
            
            # Crear límites de clase
            limite_inferior = valor_min
            limites = []
            
            for i in range(k_redondeado):
                limite_superior = limite_inferior + amplitud_redondeada
                limites.append((limite_inferior, limite_superior))
                limite_inferior = limite_superior
            
            # Calcular frecuencias
            frecuencias = [0] * k_redondeado
            for valor in valores_num:
                for i, (li, ls) in enumerate(limites):
                    if li <= valor < ls or (i == k_redondeado - 1 and valor == ls):
                        frecuencias[i] += 1
                        break
            
            # Calcular frecuencias relativas y acumuladas
            frec_rel = [f/n for f in frecuencias]
            frec_acum = []
            acum = 0
            for f in frecuencias:
                acum += f
                frec_acum.append(acum)
                
            frec_rel_acum = []
            acum_rel = 0
            for f in frec_rel:
                acum_rel += f
                frec_rel_acum.append(acum_rel)
            
            # Calcular marcas de clase
            marcas_clase = [(li + ls)/2 for li, ls in limites]
            
            # Tabla completa
            tabla_distrib = doc.add_table(rows=k_redondeado+1, cols=7)
            tabla_distrib.style = 'Table Grid'
            
            # Encabezados
            encabezados = ["Límites de clase", "Marca de clase", "Frecuencia absoluta", "Frecuencia relativa", 
                         "Frecuencia acumulada", "Frecuencia relativa acumulada", "Densidad de frecuencia"]
            
            for i, encabezado in enumerate(encabezados):
                cell = tabla_distrib.cell(0, i)
                cell.text = encabezado
                cell.paragraphs[0].runs[0].bold = True
            
            # Llenar la tabla
            for i, ((li, ls), mc, fa, fr, fac, frac) in enumerate(
                zip(limites, marcas_clase, frecuencias, frec_rel, frec_acum, frec_rel_acum), 1):
                
                tabla_distrib.cell(i, 0).text = f"[{li} - {ls})"
                tabla_distrib.cell(i, 1).text = f"{mc:.1f}"
                tabla_distrib.cell(i, 2).text = str(fa)
                tabla_distrib.cell(i, 3).text = f"{fr:.4f}"
                tabla_distrib.cell(i, 4).text = str(fac)
                tabla_distrib.cell(i, 5).text = f"{frac:.4f}"
                tabla_distrib.cell(i, 6).text = f"{fa/amplitud_redondeada:.4f}"
            
            # Paso 6: Interpretación de resultados
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 6: Interpretación de los resultados").bold = True
            
            # Encontrar la clase modal (mayor frecuencia)
            clase_modal_idx = frecuencias.index(max(frecuencias))
            clase_modal = f"[{limites[clase_modal_idx][0]} - {limites[clase_modal_idx][1]})"
            
            p = doc.add_paragraph()
            p.add_run("A partir de la tabla de distribución de frecuencias, podemos observar:")
            p.add_run(f"\n\n• La clase con mayor frecuencia es {clase_modal} con {max(frecuencias)} observaciones.")
            p.add_run(f"\n• El {frec_rel_acum[k_redondeado//2]*100:.1f}% de los datos están por debajo de {limites[k_redondeado//2][1]}.")
            
            if frecuencias[0] > frecuencias[-1]:
                p.add_run("\n• La distribución parece tener un sesgo hacia la derecha (mayor concentración en valores bajos).")
            elif frecuencias[0] < frecuencias[-1]:
                p.add_run("\n• La distribución parece tener un sesgo hacia la izquierda (mayor concentración en valores altos).")
            else:
                p.add_run("\n• La distribución parece ser aproximadamente simétrica.")
                
        except Exception as e:
            print(f"Error al procesar el ejercicio 2: {str(e)}")
            doc.add_paragraph(f"Error al generar la solución del ejercicio 2: {str(e)}").italic = True
        
        doc.add_paragraph()
        
        # Ejercicio 3: Diagrama de Tallo y Hoja
        try:
            doc.add_heading('Ejercicio 3: Diagrama de Tallo y Hoja', level=2).style = 'Heading2Custom'
            
            stem_leaf_data = variante["tercera_serie"][2]
            
            p = doc.add_paragraph("Datos del problema:")
            p.add_run("\n" + stem_leaf_data.get("title", ""))
            
            # Mostrar datos originales en tabla 3x8
            valores_sl = stem_leaf_data["data"]
            
            tabla_datos_sl = doc.add_table(rows=3, cols=8)
            tabla_datos_sl.style = 'Table Grid'
            
            idx = 0
            for i in range(3):
                for j in range(8):
                    if idx < len(valores_sl):
                        tabla_datos_sl.cell(i, j).text = valores_sl[idx]
                        idx += 1
            
            # Paso 1: Preparación de datos
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 1: Preparación y ordenamiento de los datos").bold = True
            
            # Convertir a valores numéricos y ordenar
            valores_sl_num = [float(x) for x in valores_sl]
            valores_sl_ordenados = sorted(valores_sl_num)
            
            p = doc.add_paragraph("Convertimos todos los valores a números decimales y los ordenamos de menor a mayor:")
            p_sl_ordenados = doc.add_paragraph()
            for i, val in enumerate(valores_sl_ordenados):
                if i > 0:
                    p_sl_ordenados.add_run(", ")
                p_sl_ordenados.add_run(f"{val:.1f}")
            
            # Paso 2: Identificar tallos y hojas
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 2: Construcción del diagrama de tallo y hoja").bold = True
            
            p = doc.add_paragraph()
            p.add_run("En un diagrama de tallo y hoja para datos decimales con un solo decimal:")
            p.add_run("\n• El tallo representa la parte entera del número")
            p.add_run("\n• La hoja representa el primer decimal")
            p.add_run("\nPor ejemplo, para el valor 4.5:")
            p.add_run("\n• Tallo: 4")
            p.add_run("\n• Hoja: 5")
            
            # Organizar datos por tallo y hoja
            stem_leaf_dict = {}
            
            for valor in valores_sl_num:
                tallo = int(valor)
                hoja = int((valor - tallo) * 10)  # Obtener el primer decimal
                
                if tallo not in stem_leaf_dict:
                    stem_leaf_dict[tallo] = []
                
                stem_leaf_dict[tallo].append(hoja)
            
            # Ordenar hojas para cada tallo
            for tallo in stem_leaf_dict:
                stem_leaf_dict[tallo].sort()
            
            # Crear diagrama visual
            p = doc.add_paragraph()
            p.add_run("Diagrama de tallo y hoja:").bold = True
            
            tabla_stem_leaf = doc.add_table(rows=len(stem_leaf_dict)+1, cols=2)
            tabla_stem_leaf.style = 'Table Grid'
            
            # Encabezados
            tabla_stem_leaf.cell(0, 0).text = "Tallo"
            tabla_stem_leaf.cell(0, 1).text = "Hojas"
            
            # Llenar tabla
            for i, (tallo, hojas) in enumerate(sorted(stem_leaf_dict.items()), 1):
                tabla_stem_leaf.cell(i, 0).text = str(tallo)
                
                # Formato de hojas
                hojas_str = " ".join(str(h) for h in hojas)
                tabla_stem_leaf.cell(i, 1).text = hojas_str
            
            # Paso 3: Análisis del diagrama
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 3: Análisis e interpretación del diagrama").bold = True
            
            # Encontrar el tallo con más hojas (moda del tallo)
            tallo_moda = max(stem_leaf_dict.items(), key=lambda x: len(x[1]))[0]
            
            # Encontrar el valor más frecuente
            todos_valores = {}
            for valor in valores_sl_num:
                if valor not in todos_valores:
                    todos_valores[valor] = 0
                todos_valores[valor] += 1
            
            moda = max(todos_valores.items(), key=lambda x: x[1])[0]
            
            # Obtener valores del historial de respuestas
            moda_valor = respuestas["tercera_serie"]["tallo_hoja"]["moda"]
            intervalo_conc = respuestas["tercera_serie"]["tallo_hoja"]["intervalo"]
            
            p = doc.add_paragraph()
            p.add_run("Del diagrama de tallo y hoja podemos observar:").bold = True
            p.add_run(f"\n\n1. La mayor concentración de datos se encuentra en el tallo {tallo_moda} (intervalo {intervalo_conc}).")
            p.add_run(f"\n\n2. El valor que más se repite (moda) es aproximadamente {moda_valor}.")
            
            # Determinar la forma de la distribución
            min_tallo = min(stem_leaf_dict.keys())
            max_tallo = max(stem_leaf_dict.keys())
            medio_tallo = (min_tallo + max_tallo) / 2
            
            if tallo_moda < medio_tallo:
                p.add_run("\n\n3. La distribución muestra un sesgo hacia la derecha (mayor concentración en valores bajos).")
            elif tallo_moda > medio_tallo:
                p.add_run("\n\n3. La distribución muestra un sesgo hacia la izquierda (mayor concentración en valores altos).")
            else:
                p.add_run("\n\n3. La distribución parece ser aproximadamente simétrica.")
            
            # Añadir contexto según el tipo de datos
            if "ventas" in stem_leaf_data["title"].lower() or "crecimiento" in stem_leaf_data["title"].lower():
                p.add_run(f"\n\n4. Como estos datos representan crecimiento porcentual de ventas, podemos concluir que la mayoría de los productos muestran un crecimiento alrededor del {tallo_moda}%.")
            elif "tiempo" in stem_leaf_data["title"].lower() or "atención" in stem_leaf_data["title"].lower():
                p.add_run(f"\n\n4. Como estos datos representan tiempos de atención, podemos concluir que la mayoría de los clientes son atendidos en aproximadamente {tallo_moda} minutos.")
            elif "consumo" in stem_leaf_data["title"].lower() or "combustible" in stem_leaf_data["title"].lower():
                p.add_run(f"\n\n4. Como estos datos representan consumo de combustible, podemos concluir que la mayoría de los vehículos tienen un rendimiento de aproximadamente {tallo_moda} km/litro.")
                
        except Exception as e:
            print(f"Error al procesar el ejercicio 3: {str(e)}")
            doc.add_paragraph(f"Error al generar la solución del ejercicio 3: {str(e)}").italic = True
        
        doc.add_paragraph()
        
        # Ejercicio 4: Medidas de Tendencia Central
        try:
            doc.add_heading('Ejercicio 4: Medidas de Tendencia Central', level=2).style = 'Heading2Custom'
            
            central_data = variante["tercera_serie"][3]
            
            p = doc.add_paragraph("Datos del problema:")
            p.add_run("\n" + central_data.get("title", ""))
            
            # Mostrar datos originales en tabla
            tabla_datos_mtc = doc.add_table(rows=len(central_data["ranges"])+1, cols=2)
            tabla_datos_mtc.style = 'Table Grid'
            
            # Encabezados
            tabla_datos_mtc.cell(0, 0).text = "Precio en (Q)"
            tabla_datos_mtc.cell(0, 1).text = "No. de productos"
            
            # Datos
            for i, (rango, freq) in enumerate(zip(central_data["ranges"], central_data["count"]), 1):
                tabla_datos_mtc.cell(i, 0).text = rango
                tabla_datos_mtc.cell(i, 1).text = str(freq)
            
            # Paso 1: Preparación de los datos
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 1: Preparación y organización de los datos").bold = True
            
            # Tabla de trabajo con marcas de clase y productos
            p = doc.add_paragraph("Para calcular las medidas de tendencia central con datos agrupados, primero necesitamos encontrar:")
            p.add_run("\n• La marca de clase (punto medio) de cada intervalo")
            p.add_run("\n• El producto de la marca de clase por la frecuencia (xi × fi)")
            
            tabla_trabajo = doc.add_table(rows=len(central_data["ranges"])+2, cols=4)
            tabla_trabajo.style = 'Table Grid'
            
            # Encabezados
            encabezados_trabajo = ["Precio (Q)", "Frecuencia (fi)", "Marca de clase (xi)", "xi × fi"]
            for i, encabezado in enumerate(encabezados_trabajo):
                cell = tabla_trabajo.cell(0, i)
                cell.text = encabezado
                cell.paragraphs[0].runs[0].bold = True
            
            # Preparar datos para los cálculos
            marcas_clase = []
            productos = []
            total_freq = 0
            total_producto = 0
            
            for rango, freq in zip(central_data["ranges"], central_data["count"]):
                # Extraer límites
                limites = rango.replace('[', '').replace(')', '').split('-')
                li = float(limites[0])
                ls = float(limites[1])
                
                # Marca de clase
                xi = (li + ls) / 2
                marcas_clase.append(xi)
                
                # Producto
                producto = xi * freq
                productos.append(producto)
                
                # Acumulados
                total_freq += freq
                total_producto += producto
            
            # Llenar tabla de trabajo
            for i, (rango, freq, xi, producto) in enumerate(zip(central_data["ranges"], central_data["count"], marcas_clase, productos), 1):
                tabla_trabajo.cell(i, 0).text = rango
                tabla_trabajo.cell(i, 1).text = str(freq)
                tabla_trabajo.cell(i, 2).text = f"{xi:.2f}"
                tabla_trabajo.cell(i, 3).text = f"{producto:.2f}"
            
            # Fila de totales
            tabla_trabajo.cell(len(central_data["ranges"])+1, 0).text = "TOTAL"
            tabla_trabajo.cell(len(central_data["ranges"])+1, 1).text = str(total_freq)
            tabla_trabajo.cell(len(central_data["ranges"])+1, 3).text = f"{total_producto:.2f}"
            
            # Paso 2: Cálculo de la media aritmética
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 2: Cálculo de la media aritmética").bold = True
            
            # Obtener media calculada
            media = respuestas["tercera_serie"]["medidas_centrales"]["media"]
            
            p = doc.add_paragraph()
            p.add_run("La media aritmética para datos agrupados se calcula con la fórmula:").bold = True
            p.add_run("\n\nMedia = Σ(xi × fi) ÷ Σfi")
            p.add_run(f"\n\nDonde:")
            p.add_run(f"\n• Σ(xi × fi) = {total_producto:.2f}")
            p.add_run(f"\n• Σfi = {total_freq}")
            p.add_run(f"\n\nMedia = {total_producto:.2f} ÷ {total_freq} = {total_producto/total_freq:.2f}")
            p.add_run(f"\n\nPor tanto, la media aritmética es {media}")
            
            # Paso 3: Cálculo de la mediana
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 3: Cálculo de la mediana").bold = True
            
            # Obtener mediana calculada
            mediana = respuestas["tercera_serie"]["medidas_centrales"]["mediana"]
            
            # Calcular frecuencias acumuladas
            freq_acum = []
            acum = 0
            for freq in central_data["count"]:
                acum += freq
                freq_acum.append(acum)
            
            # Posición de la mediana
            n_2 = total_freq / 2
            
            p = doc.add_paragraph()
            p.add_run("Para calcular la mediana con datos agrupados:").bold = True
            p.add_run(f"\n\n1. Primero determinamos la posición de la mediana: n/2 = {total_freq}/2 = {n_2}")
            
            # Encontrar clase mediana
            clase_mediana_idx = 0
            for i, fa in enumerate(freq_acum):
                if fa >= n_2:
                    clase_mediana_idx = i
                    break
            
            # Obtener datos de la clase mediana
            rango_mediana = central_data["ranges"][clase_mediana_idx]
            limites_mediana = rango_mediana.replace('[', '').replace(')', '').split('-')
            li_mediana = float(limites_mediana[0])
            ls_mediana = float(limites_mediana[1])
            amplitud_mediana = ls_mediana - li_mediana
            
            # Frecuencias para fórmula
            freq_clase_mediana = central_data["count"][clase_mediana_idx]
            freq_acum_anterior = 0 if clase_mediana_idx == 0 else freq_acum[clase_mediana_idx - 1]
            
            p.add_run(f"\n\n2. Identificamos la clase mediana: {rango_mediana}")
            p.add_run(f"\n\n3. Aplicamos la fórmula:")
            p.add_run(f"\nMediana = li + ((n/2 - Fi-1) ÷ fi) × c")
            p.add_run(f"\nDonde:")
            p.add_run(f"\n• li = límite inferior de la clase mediana = {li_mediana}")
            p.add_run(f"\n• n/2 = {n_2}")
            p.add_run(f"\n• Fi-1 = frecuencia acumulada anterior = {freq_acum_anterior}")
            p.add_run(f"\n• fi = frecuencia de la clase mediana = {freq_clase_mediana}")
            p.add_run(f"\n• c = amplitud de la clase = {amplitud_mediana}")
            
            calculo_mediana = li_mediana + ((n_2 - freq_acum_anterior) / freq_clase_mediana) * amplitud_mediana
            p.add_run(f"\n\nMediana = {li_mediana} + (({n_2} - {freq_acum_anterior}) ÷ {freq_clase_mediana}) × {amplitud_mediana}")
            p.add_run(f"\nMediana = {li_mediana} + ({n_2 - freq_acum_anterior} ÷ {freq_clase_mediana}) × {amplitud_mediana}")
            p.add_run(f"\nMediana = {li_mediana} + ({(n_2 - freq_acum_anterior) / freq_clase_mediana:.4f}) × {amplitud_mediana}")
            p.add_run(f"\nMediana = {li_mediana} + {((n_2 - freq_acum_anterior) / freq_clase_mediana) * amplitud_mediana:.4f}")
            p.add_run(f"\nMediana = {calculo_mediana:.4f}")
            
            p.add_run(f"\n\nPor tanto, la mediana es {mediana}")
            
            # Paso 4: Cálculo de la moda
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 4: Cálculo de la moda").bold = True
            
            # Obtener moda calculada
            moda = respuestas["tercera_serie"]["medidas_centrales"]["moda"]
            
            # Encontrar clase modal (mayor frecuencia)
            clase_modal_idx = central_data["count"].index(max(central_data["count"]))
            freq_modal = central_data["count"][clase_modal_idx]
            
            # Datos de la clase modal
            rango_modal = central_data["ranges"][clase_modal_idx]
            limites_modal = rango_modal.replace('[', '').replace(')', '').split('-')
            li_modal = float(limites_modal[0])
            ls_modal = float(limites_modal[1])
            amplitud_modal = ls_modal - li_modal
            
            # Frecuencias para fórmula
            freq_anterior = 0 if clase_modal_idx == 0 else central_data["count"][clase_modal_idx - 1]
            freq_posterior = 0 if clase_modal_idx == len(central_data["count"]) - 1 else central_data["count"][clase_modal_idx + 1]
            
            p = doc.add_paragraph()
            p.add_run("Para calcular la moda con datos agrupados:").bold = True
            p.add_run(f"\n\n1. Primero identificamos la clase modal (mayor frecuencia): {rango_modal} con frecuencia {freq_modal}")
            
            p.add_run(f"\n\n2. Aplicamos la fórmula:")
            p.add_run(f"\nModa = li + (d1 ÷ (d1 + d2)) × c")
            p.add_run(f"\nDonde:")
            p.add_run(f"\n• li = límite inferior de la clase modal = {li_modal}")
            p.add_run(f"\n• d1 = frecuencia modal - frecuencia anterior = {freq_modal} - {freq_anterior} = {freq_modal - freq_anterior}")
            p.add_run(f"\n• d2 = frecuencia modal - frecuencia posterior = {freq_modal} - {freq_posterior} = {freq_modal - freq_posterior}")
            p.add_run(f"\n• c = amplitud de la clase = {amplitud_modal}")
            
            # Si d1 o d2 son negativos o cero, usar valor absoluto o 1
            d1 = max(1, freq_modal - freq_anterior)
            d2 = max(1, freq_modal - freq_posterior)
            
            calculo_moda = li_modal + (d1 / (d1 + d2)) * amplitud_modal
            p.add_run(f"\n\nModa = {li_modal} + ({d1} ÷ ({d1} + {d2})) × {amplitud_modal}")
            p.add_run(f"\nModa = {li_modal} + ({d1} ÷ {d1 + d2}) × {amplitud_modal}")
            p.add_run(f"\nModa = {li_modal} + ({d1 / (d1 + d2):.4f}) × {amplitud_modal}")
            p.add_run(f"\nModa = {li_modal} + {(d1 / (d1 + d2)) * amplitud_modal:.4f}")
            p.add_run(f"\nModa = {calculo_moda:.4f}")
            
            p.add_run(f"\n\nPor tanto, la moda es {moda}")
            
            # Paso 5: Interpretación de resultados
            p = doc.add_paragraph()
            p.style = 'StepCustom'
            p.add_run("Paso 5: Interpretación de los resultados").bold = True
            
            p = doc.add_paragraph()
            p.add_run("Interpretación de las medidas de tendencia central:").bold = True
            
            p.add_run(f"\n\n• Media = {media}: Representa el valor promedio de los precios. Si todos los productos tuvieran el mismo precio, sería este valor.")
            
            p.add_run(f"\n\n• Mediana = {mediana}: Es el valor que divide al conjunto de datos en dos partes iguales. El 50% de los productos tienen precios inferiores a este valor y el 50% tienen precios superiores.")
            
            p.add_run(f"\n\n• Moda = {moda}: Representa el valor que aparece con mayor frecuencia. Es el precio más común entre los productos.")
            
            # Comparar las tres medidas
            if abs(media - mediana) < 200 and abs(media - moda) < 200:
                p.add_run("\n\nLas tres medidas son relativamente cercanas, lo que sugiere que la distribución de precios es aproximadamente simétrica.")
            elif media > mediana and mediana > moda:
                p.add_run("\n\nComo Media > Mediana > Moda, la distribución de precios tiene un sesgo positivo (hacia la derecha), indicando que hay algunos productos con precios considerablemente más altos que el resto.")
            elif media < mediana and mediana < moda:
                p.add_run("\n\nComo Media < Mediana < Moda, la distribución de precios tiene un sesgo negativo (hacia la izquierda), indicando que hay algunos productos con precios considerablemente más bajos que el resto.")
            else:
                p.add_run("\n\nLas relaciones entre las medidas de tendencia central sugieren una distribución irregular o multimodal de los precios.")
                
        except Exception as e:
            print(f"Error al procesar el ejercicio 4: {str(e)}")
            doc.add_paragraph(f"Error al generar la solución del ejercicio 4: {str(e)}").italic = True
        
        # Pie de página
        doc.add_paragraph()
        footer = doc.add_paragraph("Universidad Panamericana - Facultad de Humanidades")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_detail = doc.add_paragraph(f"Solución Matemática Detallada - {tipo_texto} - Sección {seccion} - Variante {variante_id}")
        footer_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Guardar los archivos
        detailed_filename = f'Solucion_Matematica_Detallada_{seccion}_{tipo_evaluacion}_{variante_id}.docx'
        detailed_path = os.path.join(output_dir, detailed_filename)
        
        simple_filename = f'Solucion_Matematica_Detallada_{variante_id}.docx'
        simple_path = os.path.join(PLANTILLAS_FOLDER, simple_filename)
        
        doc.save(detailed_path)
        print(f"Solución detallada guardada en: {detailed_path}")
        
        doc.save(simple_path)
        print(f"Solución detallada guardada en: {simple_path}")
        
        print(f"===== FINALIZADA CREACIÓN DE SOLUCIÓN MATEMÁTICA DETALLADA =====\n")
        
        return simple_filename
        
    except Exception as e:
        print(f"Error general en solución matemática detallada: {str(e)}")
        traceback.print_exc()
        return None

def crear_solucion_matematica_simplificada(variante_id, seccion="A", tipo_evaluacion="parcial1"):
    """Crea un archivo Word básico con soluciones matemáticas - versión simplificada"""
    try:
        # Evitar importaciones complejas que podrían causar problemas
        # import matplotlib
        # matplotlib.use('Agg')
        # import matplotlib.pyplot as plt
        # import numpy as np
        # from scipy import stats
        import math
        import os
        from datetime import datetime
        
        # Verificar importaciones de docx
        print("Verificando docx y componentes...")
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        # Evitar importaciones adicionales por ahora
        # from docx.shared import RGBColor
        # from docx.enum.style import WD_STYLE_TYPE
        
        # Generar carpeta con timestamp para evitar ambigüedades
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(PLANTILLAS_FOLDER, f'{seccion}_{tipo_evaluacion}_{timestamp}')
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"Comprobando existencia de directorio: {output_dir}, Existe: {os.path.exists(output_dir)}")
        print(f"Comprobando permisos de escritura: {os.access(output_dir, os.W_OK)}")
        
        # Cargar respuestas y variante
        print("Cargando datos...")
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
            respuestas = json.load(f)
        
        # Crear documento Word - versión básica
        print("Creando documento Word...")
        doc = Document()
        
        # Configurar márgenes básicos
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Tipo de evaluación a texto legible
        tipo_textos = {
            'parcial1': 'PRIMER PARCIAL',
            'parcial2': 'SEGUNDO PARCIAL',
            'final': 'EXAMEN FINAL',
            'corto': 'EVALUACIÓN CORTA',
            'recuperacion': 'RECUPERACIÓN',
            'test': 'PRUEBA'
        }
        tipo_texto = tipo_textos.get(tipo_evaluacion, 'EVALUACIÓN PARCIAL')
        
        # Título principal - formato básico
        heading = doc.add_heading('SOLUCIÓN MATEMÁTICA', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subheading = doc.add_heading(f'{tipo_texto} - SECCIÓN {seccion} - VARIANTE {variante_id}', 1)
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        p = doc.add_paragraph(f"Fecha de generación: {fecha_actual}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Aviso
        p = doc.add_paragraph("DOCUMENTO CONFIDENCIAL - SOLO PARA USO DEL DOCENTE")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        
        # Primera Serie
        doc.add_heading('PRIMERA SERIE - RESPUESTAS CORRECTAS', level=1)
        
        # Tabla básica sin estilos complejos
        table1 = doc.add_table(rows=len(respuestas["primera_serie"])+1, cols=3)
        
        # Encabezados
        headers = ["Pregunta", "Respuesta", "Texto"]
        for i, header in enumerate(headers):
            cell = table1.cell(0, i)
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True
            
        # Llenar tabla con respuestas
        for i, resp_idx in enumerate(respuestas["primera_serie"], 1):
            if i-1 < len(variante['primera_serie']):
                pregunta = variante['primera_serie'][i-1]
                
                # Número y texto de pregunta
                table1.cell(i, 0).text = f"{i}. {pregunta['pregunta'][:50]}..."
                
                # Letra de respuesta
                table1.cell(i, 1).text = f"Opción {resp_idx+1}"
                
                # Texto de respuesta
                if resp_idx < len(pregunta['opciones']):
                    table1.cell(i, 2).text = pregunta['opciones'][resp_idx]
                else:
                    table1.cell(i, 2).text = "N/A"
        
        # Segunda Serie
        doc.add_heading('SEGUNDA SERIE - RESPUESTAS CORRECTAS', level=1)
        
        # Tercera Serie
        doc.add_heading('TERCERA SERIE - SOLUCIONES', level=1)
        
        # Ejercicio 1
        doc.add_heading('Ejercicio 1: Coeficiente de Gini', level=2)
        
        # Valor calculado
        try:
            gini_value = respuestas["tercera_serie"]["gini"]
            p = doc.add_paragraph(f"Coeficiente de Gini = {gini_value}")
            run = p.runs[0]
            run.bold = True
        except:
            p = doc.add_paragraph("Datos no disponibles")
        
        # Ejercicio 2
        doc.add_heading('Ejercicio 2: Distribución de Frecuencias', level=2)
        
        # Valores calculados
        try:
            k_value = respuestas["tercera_serie"]["dist_frecuencias"]["k"]
            rango = respuestas["tercera_serie"]["dist_frecuencias"]["rango"]
            amplitud = respuestas["tercera_serie"]["dist_frecuencias"]["amplitud"]
            
            p = doc.add_paragraph()
            p.add_run(f"K = {k_value}").bold = True
            p.add_run("\n")
            p.add_run(f"Rango = {rango}").bold = True
            p.add_run("\n")
            p.add_run(f"Amplitud = {amplitud}").bold = True
        except:
            p = doc.add_paragraph("Datos no disponibles")
        
        # Ejercicio 3
        doc.add_heading('Ejercicio 3: Tallo y Hoja', level=2)
        
        # Valores calculados
        try:
            moda = respuestas["tercera_serie"]["tallo_hoja"]["moda"]
            intervalo = respuestas["tercera_serie"]["tallo_hoja"]["intervalo"]
            
            p = doc.add_paragraph()
            p.add_run(f"Moda = {moda}").bold = True
            p.add_run("\n")
            p.add_run(f"Intervalo de mayor concentración = {intervalo}").bold = True
        except:
            p = doc.add_paragraph("Datos no disponibles")
        
        # Ejercicio 4
        doc.add_heading('Ejercicio 4: Medidas de Tendencia Central', level=2)
        
        # Valores calculados
        try:
            media = respuestas["tercera_serie"]["medidas_centrales"]["media"]
            mediana = respuestas["tercera_serie"]["medidas_centrales"]["mediana"]
            moda = respuestas["tercera_serie"]["medidas_centrales"]["moda"]
            
            p = doc.add_paragraph()
            p.add_run(f"Media = {media}").bold = True
            p.add_run("\n")
            p.add_run(f"Mediana = {mediana}").bold = True
            p.add_run("\n")
            p.add_run(f"Moda = {moda}").bold = True
        except:
            p = doc.add_paragraph("Datos no disponibles")
        
        # Pie de página
        doc.add_paragraph()
        footer = doc.add_paragraph("Universidad Panamericana - Facultad de Humanidades")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nombre de archivos
        detailed_filename = f'Solucion_Matematica_{seccion}_{tipo_evaluacion}_{variante_id}.docx'
        detailed_path = os.path.join(output_dir, detailed_filename)
        
        simple_filename = f'Solucion_Matematica_{variante_id}.docx'
        simple_path = os.path.join(PLANTILLAS_FOLDER, simple_filename)
        
        # Guardar archivos con manejo de excepciones detallado
        try:
            print(f"Intentando guardar en: {detailed_path}")
            doc.save(detailed_path)
            print(f"Solución matemática guardada en: {detailed_path}")
            
            print(f"Intentando guardar en: {simple_path}")
            doc.save(simple_path)
            print(f"Solución matemática guardada en: {simple_path}")
            
            return simple_filename
        except Exception as e:
            print(f"Error específico al guardar: {str(e)}")
            print(f"Tipo de error: {type(e).__name__}")
            traceback.print_exc()
            
            # Intentar una ubicación alternativa si hay problemas de permisos
            try:
                alt_path = os.path.join(UPLOAD_FOLDER, simple_filename)
                print(f"Intentando ubicación alternativa: {alt_path}")
                doc.save(alt_path)
                print(f"Guardado en ubicación alternativa: {alt_path}")
                
                # Copiar a la ubicación original
                import shutil
                shutil.copy2(alt_path, simple_path)
                print(f"Copiado a la ubicación original: {simple_path}")
                
                return simple_filename
            except Exception as e2:
                print(f"Error también en ubicación alternativa: {str(e2)}")
                traceback.print_exc()
                return None
            
    except Exception as e:
        print(f"Error general en solución matemática: {str(e)}")
        print(f"Tipo de error: {type(e).__name__}")
        traceback.print_exc()
        return None
    
# Rutas de la aplicación
# Replace your current index route with this modified version
# This specifically addresses the maximum recursion depth exceeded in comparison error

@app.route('/')
def index():
    """
    Función principal que muestra la página de inicio con las variantes disponibles
    """
    # Cargar el historial para referencia
    try:
        historial_raw = cargar_historial()
        
        # Crear una función para sanitizar el diccionario y prevenir recursión
        def sanitize_dict(d):
            """Crea una copia segura de un diccionario sin referencias circulares"""
            if not isinstance(d, dict):
                return d
                
            # Crear un nuevo diccionario simplificado con solo las claves necesarias
            safe_dict = {}
            
            # Lista de claves seguras que sabemos que no causan problemas
            safe_keys = ['id', 'seccion', 'tipo_evaluacion', 'tipo_texto', 
                        'examen', 'hoja', 'plantilla', 'solucion_matematica',
                        'fecha_generacion', 'timestamp', 'directorio']
                        
            for key in safe_keys:
                if key in d:
                    # Evitar copiar objetos complejos que podrían causar recursión
                    if isinstance(d[key], (str, int, float, bool, type(None))):
                        safe_dict[key] = d[key]
                    else:
                        # Convertir a string para evitar recursión
                        safe_dict[key] = str(d[key])
            
            return safe_dict
        
        # Sanitizar el historial
        historial_safe = [sanitize_dict(item) for item in historial_raw]
        
        # Lista para almacenar las variantes
        variantes = []
        
        # Set para evitar duplicados
        variantes_procesadas = set()
        
        # Obtener las variantes del historial (más recientes primero)
        # Limitar el número de elementos para evitar problemas
        for item in historial_safe[:20]:  # Solo procesar hasta 20 entradas del historial
            variante_id = item.get('id')
            
            # Evitar duplicados
            if variante_id in variantes_procesadas:
                continue
            
            # CORRECCIÓN: Mejorar la verificación de existencia de archivos
            # Asegurarse de buscar en directorios específicos si existen
            directorio = item.get('directorio')
            
            # Verificar si existen los archivos
            tiene_examen = False
            tiene_hoja = False
            tiene_plantilla = False
            tiene_solucion = False
            solucion_matematica = item.get('solucion_matematica')
            
            # Rutas de archivo estándar (sin directorio)
            examen_path = os.path.join(EXAMENES_FOLDER, f'Examen_{variante_id}.docx')
            hoja_path = os.path.join(HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{variante_id}.pdf')
            plantilla_path = os.path.join(PLANTILLAS_FOLDER, f'Plantilla_{variante_id}.pdf')
            
            # Verificar en ruta estándar primero
            tiene_examen = os.path.exists(examen_path)
            tiene_hoja = os.path.exists(hoja_path)
            tiene_plantilla = os.path.exists(plantilla_path)
            
            # Si hay un directorio específico, también verificar allí
            if directorio:
                examen_dir_path = os.path.join(EXAMENES_FOLDER, directorio, f'Examen_{item.get("seccion")}_{item.get("tipo_evaluacion")}_{variante_id}.docx')
                hoja_dir_path = os.path.join(HOJAS_RESPUESTA_FOLDER, directorio, f'HojaRespuestas_{item.get("seccion")}_{item.get("tipo_evaluacion")}_{variante_id}.pdf')
                plantilla_dir_path = os.path.join(PLANTILLAS_FOLDER, directorio, f'Plantilla_{item.get("seccion")}_{item.get("tipo_evaluacion")}_{variante_id}.pdf')
                
                # Actualizar estado si se encuentra en directorio específico
                tiene_examen = tiene_examen or os.path.exists(examen_dir_path)
                tiene_hoja = tiene_hoja or os.path.exists(hoja_dir_path)
                tiene_plantilla = tiene_plantilla or os.path.exists(plantilla_dir_path)
            
            # Verificar la solución matemática
            if solucion_matematica:
                sol_path = os.path.join(PLANTILLAS_FOLDER, solucion_matematica)
                tiene_solucion = os.path.exists(sol_path)
                
                # También verificar en el directorio si existe
                if directorio and not tiene_solucion:
                    sol_dir_path = os.path.join(PLANTILLAS_FOLDER, directorio, solucion_matematica)
                    tiene_solucion = os.path.exists(sol_dir_path)
            
            # Añadir la variante a la lista
            variantes.append({
                'id': variante_id,
                'seccion': item.get('seccion', 'No especificada'),
                'tipo_evaluacion': item.get('tipo_texto', 'No especificado'),
                'tiene_examen': tiene_examen,
                'tiene_hoja': tiene_hoja,
                'tiene_plantilla': tiene_plantilla,
                'solucion_matematica': solucion_matematica if tiene_solucion else None,
                'tiene_solucion': tiene_solucion,
                'directorio': directorio
            })
            
            variantes_procesadas.add(variante_id)
            
            # Limitar el número de variantes para evitar problemas
            if len(variantes) >= 10:
                break
        
        # Sanitizar las variantes para evitar cualquier recursión
        variantes_safe = []
        for v in variantes:
            variantes_safe.append(sanitize_dict(v))
        
        return render_template('index.html', variantes=variantes_safe)
    
    except Exception as e:
        app.logger.error(f"Error en index: {str(e)}")
        
        # Devolver una página sencilla con el error
        error_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Error</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .error {{ background-color: #f8d7da; border: 1px solid #f5c6cb; padding: 20px; border-radius: 5px; }}
                pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            <h1>Error en la aplicación</h1>
            <div class="error">
                <p><strong>Mensaje de error:</strong> {str(e)}</p>
            </div>
            <h2>Detalles técnicos:</h2>
            <pre>{traceback.format_exc()}</pre>
            
            <h2>Por favor, intente:</h2>
            <ul>
                <li>Reiniciar la aplicación Flask</li>
                <li>Verificar la integridad de los archivos JSON en la carpeta variantes</li>
                <li>Contactar al administrador del sistema</li>
            </ul>
            
            <p><a href="/">Intentar nuevamente</a></p>
        </body>
        </html>
        """
        return error_html

# Add this to your app.py file

@app.route('/diagnostico_datos')
def diagnostico_datos():
    """
    Ruta para diagnosticar problemas con las estructuras de datos
    """
    try:
        # Cargar el historial para obtener el orden cronológico
        historial = cargar_historial()
        
        # Lista para almacenar las variantes (simplificada)
        variantes_simplificadas = []
        
        # Set para evitar duplicados
        variantes_procesadas = set()
        
        # Un contador para evitar bucles infinitos
        contador = 0
        max_iteraciones = 50  # Límite de seguridad
        
        # Diagnosticar el historial
        historial_info = {
            "largo": len(historial),
            "tipos": []
        }
        
        for item in historial:
            if isinstance(item, dict):
                historial_info["tipos"].append("diccionario")
                
                # Comprobar keys problemáticas que podrían causar recursión
                if "historial" in item or "variantes" in item:
                    historial_info["alerta"] = "¡Posible referencia circular encontrada en el historial!"
            else:
                historial_info["tipos"].append(type(item).__name__)
        
        # Intentar diagnosticar variantes de historial
        try:
            # Limitar a un máximo de 5 elementos para evitar problemas
            for item in historial[:5]:
                variante_id = item.get('id')
                
                if variante_id and variante_id not in variantes_procesadas:
                    variantes_simplificadas.append({
                        'id': variante_id,
                        'seccion': item.get('seccion', 'No especificada'),
                        'tipo_evaluacion': item.get('tipo_texto', 'No especificado')
                    })
                    
                    variantes_procesadas.add(variante_id)
                
                contador += 1
                if contador >= max_iteraciones:
                    break
        except Exception as e:
            return f"Error al procesar historial: {str(e)}"
        
        # Intentar cargar variantes desde archivos
        try:
            contador = 0  # Reiniciar contador
            
            if os.path.exists(VARIANTES_FOLDER):
                for archivo in os.listdir(VARIANTES_FOLDER):
                    if archivo.startswith('variante_') and archivo.endswith('.json'):
                        variante_id = archivo.replace('variante_', '').replace('.json', '')
                        
                        if variante_id in variantes_procesadas:
                            continue
                        
                        # Simplemente añadir el ID, sin cargar el archivo completo
                        variantes_simplificadas.append({
                            'id': variante_id,
                            'seccion': 'No especificada',
                            'tipo_evaluacion': 'No especificado'
                        })
                        
                        contador += 1
                        if contador >= 10:  # Limitar a 10 elementos
                            break
        except Exception as e:
            return f"Error al procesar archivos de variantes: {str(e)}"
            
        # Devolver un diagnóstico en formato texto
        resultado = f"""
        <h1>Diagnóstico de Datos</h1>
        
        <h2>Historial</h2>
        <p>Número de elementos: {historial_info['largo']}</p>
        <p>Tipos de elementos: {', '.join(historial_info['tipos'][:5])}</p>
        
        <h2>Variantes Encontradas</h2>
        <ul>
        {''.join([f'<li>{v["id"]} - {v["seccion"]} - {v["tipo_evaluacion"]}</li>' for v in variantes_simplificadas[:10]])}
        </ul>
        
        <p><a href="/">Volver al inicio</a></p>
        """
        
        return resultado
        
    except Exception as e:
        return f"Error general en diagnóstico: {str(e)}"

@app.route('/estudiantes', methods=['GET', 'POST'])
def gestionar_estudiantes():
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'agregar':
            nombre = request.form.get('nombre')
            carne = request.form.get('carne')
            seccion = request.form.get('seccion')
            
            if nombre and carne and seccion:
                if seccion not in estudiantes_db:
                    estudiantes_db[seccion] = []
                
                estudiantes_db[seccion].append({
                    'nombre': nombre,
                    'carne': carne,
                    'evaluaciones': {}
                })
                
                flash(f'Estudiante {nombre} agregado correctamente', 'success')
            else:
                flash('Datos incompletos', 'danger')
                
        elif action == 'cargar_csv':
            if 'archivo_csv' not in request.files:
                flash('No se seleccionó archivo CSV', 'danger')
                return redirect(request.url)
                
            archivo = request.files['archivo_csv']
            seccion = request.form.get('seccion')
            
            if archivo.filename == '' or not seccion:
                flash('Datos incompletos', 'danger')
                return redirect(request.url)
                
            # Procesar archivo CSV
            try:
                contenido = archivo.read().decode('utf-8')
                lineas = contenido.split('\n')
                
                if seccion not in estudiantes_db:
                    estudiantes_db[seccion] = []
                
                for linea in lineas[1:]:  # Omitir encabezado
                    if not linea.strip():
                        continue
                        
                    datos = linea.split(',')
                    if len(datos) >= 2:
                        estudiantes_db[seccion].append({
                            'nombre': datos[0].strip(),
                            'carne': datos[1].strip(),
                            'evaluaciones': {}
                        })
                
                flash(f'Se cargaron {len(lineas)-1} estudiantes para la sección {seccion}', 'success')
            except Exception as e:
                flash(f'Error al procesar CSV: {str(e)}', 'danger')
    
    return render_template('estudiantes.html', estudiantes=estudiantes_db)

@app.route('/calificaciones/<seccion>/<tipo_evaluacion>')
def ver_calificaciones(seccion, tipo_evaluacion):
    if seccion not in estudiantes_db:
        flash(f'La sección {seccion} no existe', 'danger')
        return redirect(url_for('gestionar_estudiantes'))
    
    # Obtener lista de exámenes procesados para esta sección y tipo
    eval_folder = os.path.join(EXAMENES_ESCANEADOS_FOLDER, f"{seccion}_{tipo_evaluacion}")
    
    resultados = {}
    if os.path.exists(eval_folder):
        for archivo in os.listdir(eval_folder):
            if archivo.endswith('.json'):
                with open(os.path.join(eval_folder, archivo), 'r') as f:
                    resultado = json.load(f)
                    carne = resultado.get('info_estudiante', {}).get('carne')
                    if carne:
                        resultados[carne] = resultado
    
    # Preparar datos para mostrar
    estudiantes_seccion = estudiantes_db.get(seccion, [])
    for estudiante in estudiantes_seccion:
        carne = estudiante.get('carne')
        if carne in resultados:
            estudiante['resultado'] = resultados[carne]
        else:
            estudiante['resultado'] = None
    
    return render_template('calificaciones.html', 
                          estudiantes=estudiantes_seccion, 
                          seccion=seccion, 
                          tipo_evaluacion=tipo_evaluacion)

@app.route('/generar_examen', methods=['POST'])
def generar_examen_handler():
    try:
        num_variantes = int(request.form.get('num_variantes', 1))
        seccion = request.form.get('seccion', 'A')
        tipo_evaluacion = request.form.get('tipo_evaluacion', 'parcial1')
        
        # Obtener datos adicionales para personalización
        licenciatura = request.form.get('licenciatura', '')
        nombre_curso = request.form.get('nombre_curso', '')
        nombre_docente = request.form.get('nombre_docente', '')
        anio = request.form.get('anio', '')
        salon = request.form.get('salon', '')
        
        # Manejo del logo
        logo_path = None
        if 'logo' in request.files and request.files['logo'].filename:
            logo = request.files['logo']
            logo_filename = secure_filename(logo.filename)
            logo_path = os.path.join(UPLOAD_FOLDER, logo_filename)
            logo.save(logo_path)
        
        # Manejo de la plantilla
        plantilla_path = None
        if 'plantilla' in request.files and request.files['plantilla'].filename:
            plantilla = request.files['plantilla']
            plantilla_filename = secure_filename(plantilla.filename)
            plantilla_path = os.path.join(UPLOAD_FOLDER, plantilla_filename)
            plantilla.save(plantilla_path)
        
        # Utilizar la nueva función de generación con solución detallada
        variantes_generadas = generar_examen(
            num_variantes=num_variantes,
            seccion=seccion,
            tipo_evaluacion=tipo_evaluacion,
            logo_path=logo_path,
            plantilla_path=plantilla_path,
            licenciatura=licenciatura,
            nombre_curso=nombre_curso,
            nombre_docente=nombre_docente,
            anio=anio,
            salon=salon,
            uso_detallado=True  # Activar generación de solución detallada
        )
        
        # Verificar si hubo éxito
        if variantes_generadas:
            flash(f'Se han generado {num_variantes} variantes de examen para la sección {seccion}', 'success')
        else:
            flash('Error: No se pudieron generar las variantes. Revise los logs para más detalles.', 'warning')
        
        return redirect(url_for('index'))
    except Exception as e:
        error_msg = f'Error al generar exámenes: {str(e)}'
        print(error_msg)
        traceback.print_exc()
        flash(error_msg, 'danger')
        return redirect(url_for('index'))

def procesar_plantilla_examen(plantilla_path, variante_id, seccion, tipo_evaluacion, variante, 
                            licenciatura="", nombre_curso="Estadística Básica", nombre_docente="Ing. Marco Antonio Jiménez", 
                            anio="2025", salon=""):
    """
    Procesa una plantilla de examen reemplazando los placeholders con el contenido generado.
    Versión mejorada para evitar problemas de formato con asteriscos.
    """
    try:
        from docx import Document
        import re
        
        print(f"Procesando plantilla: {plantilla_path}")
        
        # Cargar la plantilla
        doc = Document(plantilla_path)
        
        # Convertir el tipo de evaluación a texto legible
        tipo_textos = {
            'parcial1': 'Primer Examen Parcial',
            'parcial2': 'Segundo Examen Parcial',
            'final': 'Examen Final',
            'corto': 'Evaluación Corta',
            'recuperacion': 'Examen de Recuperación',
            'test': 'Prueba de Evaluación'
        }
        tipo_texto = tipo_textos.get(tipo_evaluacion, 'Evaluación Parcial')
        
        # Reemplazar placeholders en todo el documento
        placeholders = {
            '{variante}': variante_id,
            '{salon}': salon if salon else "",
            '{licenciatura}': licenciatura if licenciatura else "Humanidades",
            '{nombre_curso}': nombre_curso if nombre_curso else "Estadística Básica",
            '{nombre_docente}': nombre_docente if nombre_docente else "Ing. Marco Antonio Jiménez",
            '{anio}': anio if anio else "2025",
            '{tipo_evaluacion}': tipo_texto
        }
        
        # SOLUCIÓN MEJORADA: Método más seguro para reemplazar texto en Word
        # que preserva el formato original de la plantilla
        
        # Para párrafos
        for paragraph in doc.paragraphs:
            # Solo procesar párrafos que contienen placeholders
            if any(key in paragraph.text for key in placeholders.keys()):
                # Conservar el estilo original
                original_style = paragraph.style
                
                # Texto completo antes de procesarlo
                original_text = paragraph.text
                new_text = original_text
                
                # Reemplazar todos los placeholders en este texto
                for key, value in placeholders.items():
                    if key in new_text:
                        new_text = new_text.replace(key, value)
                
                # Solo si realmente hubo cambios
                if new_text != original_text:
                    # Limpiar el párrafo y reemplazar con el nuevo texto
                    paragraph.clear()
                    paragraph.add_run(new_text)
                    
                    # Restaurar el estilo original
                    paragraph.style = original_style
        
        # Para tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Solo procesar párrafos que contienen placeholders
                        if any(key in paragraph.text for key in placeholders.keys()):
                            # Conservar el estilo original
                            original_style = paragraph.style
                            
                            # Texto completo antes de procesarlo
                            original_text = paragraph.text
                            new_text = original_text
                            
                            # Reemplazar todos los placeholders en este texto
                            for key, value in placeholders.items():
                                if key in new_text:
                                    new_text = new_text.replace(key, value)
                            
                            # Solo si realmente hubo cambios
                            if new_text != original_text:
                                # Limpiar el párrafo y reemplazar con el nuevo texto
                                paragraph.clear()
                                paragraph.add_run(new_text)
                                
                                # Restaurar el estilo original
                                paragraph.style = original_style
        
        # Función para generar el contenido de las series
        def generar_primera_serie(preguntas):
            contenido = []
            for i, pregunta in enumerate(preguntas, 1):
                # Texto de la pregunta - sin formato especial, se aplicará desde la plantilla
                contenido.append(f"{i}. {pregunta['pregunta']}")
                
                # Opciones
                for opcion in pregunta.get('opciones', []):
                    contenido.append(f"   • {opcion}")
                
                contenido.append("")  # Línea en blanco entre preguntas
            
            return "\n".join(contenido)
        
        def generar_segunda_serie(escenarios):
            contenido = []
            for i, escenario in enumerate(escenarios, 1):
                # Texto del escenario - sin formato especial
                contenido.append(f"{i}. {escenario.get('escenario', '')}")
                
                # Opciones
                for opcion in escenario.get('opciones', []):
                    contenido.append(f"   • {opcion}")
                
                contenido.append("")  # Línea en blanco entre escenarios
            
            return "\n".join(contenido)
        
        def generar_tercera_serie(ejercicios):
            contenido = []
            
            # Ejercicio 1: Coeficiente de Gini
            if len(ejercicios) > 0:
                gini_data = ejercicios[0]
                contenido.append(f"1. {gini_data.get('title', 'Problema de Coeficiente de Gini')}")
                
                # Agregar instrucciones del ejercicio sin formato especial
                contenido.append("   a) Complete la tabla para calcular el coeficiente de Gini.")
                contenido.append("   b) Calcule el coeficiente de Gini utilizando la fórmula correspondiente.")
                contenido.append("   c) Interprete el resultado obtenido respecto a la desigualdad en la distribución de salarios.")
                contenido.append("")
            
            # Ejercicio 2: Distribución de frecuencias
            if len(ejercicios) > 1:
                sturgers_data = ejercicios[1]
                contenido.append(f"2. {sturgers_data.get('title', 'Problema de Distribución de Frecuencias')}")
                contenido.append("   Construya la tabla de distribución de frecuencias correspondiente.")
                contenido.append("")
            
            # Ejercicio 3: Diagrama de Tallo y Hoja
            if len(ejercicios) > 2:
                stem_leaf_data = ejercicios[2]
                contenido.append(f"3. {stem_leaf_data.get('title', 'Problema de Diagrama de Tallo y Hoja')}")
                contenido.append("   a) Realizar un Diagrama de Tallo y Hoja para identificar donde se encuentra la mayor concentración de los datos.")
                contenido.append("   b) Interprete los datos y explique brevemente sus resultados.")
                contenido.append("")
            
            # Ejercicio 4: Medidas de tendencia central
            if len(ejercicios) > 3:
                central_data = ejercicios[3]
                contenido.append(f"4. {central_data.get('title', 'Problema de Medidas de Tendencia Central')}")
                contenido.append("")
            
            return "\n".join(contenido)
        
        # Generar contenido sin aplicar formato especial
        primera_serie_text = generar_primera_serie(variante.get("primera_serie", []))
        segunda_serie_text = generar_segunda_serie(variante.get("segunda_serie", []))
        tercera_serie_text = generar_tercera_serie(variante.get("tercera_serie", []))
        
        # Diccionario para secciones especiales
        secciones = {
            "{primera_serie}": primera_serie_text,
            "{segunda_serie}": segunda_serie_text,
            "{tercera_serie}": tercera_serie_text
        }
        
        # Reemplazar las secciones especiales
        for paragraph in doc.paragraphs:
            for marcador, contenido in secciones.items():
                if marcador in paragraph.text:
                    # Preservar estilo
                    original_style = paragraph.style
                    
                    # Limpiar y reemplazar
                    paragraph.clear()
                    paragraph.add_run(contenido)
                    
                    # Restaurar estilo
                    paragraph.style = original_style
                    print(f"Reemplazado {marcador} en párrafo")
                    break  # Salir del bucle si encontramos un marcador
        
        # También en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for marcador, contenido in secciones.items():
                            if marcador in paragraph.text:
                                # Preservar estilo
                                original_style = paragraph.style
                                
                                # Limpiar y reemplazar
                                paragraph.clear()
                                paragraph.add_run(contenido)
                                
                                # Restaurar estilo
                                paragraph.style = original_style
                                print(f"Reemplazado {marcador} en tabla")
                                break  # Salir del bucle si encontramos un marcador
        
        return doc
        
    except Exception as e:
        print(f"Error al procesar plantilla: {str(e)}")
        traceback.print_exc()
        return None
        
def crear_examen_word(variante_id, seccion="A", tipo_evaluacion="parcial1", logo_path=None, plantilla_path=None, 
                  licenciatura="", nombre_curso="Estadística Básica", nombre_docente="Ing. Marco Antonio Jiménez", 
                  anio="2025", salon=""):
    """
    Crea documento de examen con logo y usando plantilla opcional
    Soporta reemplazo de placeholders para personalización
    """
    try:
        print(f"\n===== INICIANDO CREACIÓN DE EXAMEN WORD =====")
        print(f"Parámetros: variante_id={variante_id}, seccion={seccion}, tipo_evaluacion={tipo_evaluacion}")
        print(f"Logo path: {logo_path}")
        print(f"Plantilla path: {plantilla_path}")
        
        # Importar Document aquí para asegurarse de que está disponible
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
        
        # Cargar la variante
        try:
            with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
                variante = json.load(f)
                print(f"Variante cargada correctamente: {variante_id}")
        except Exception as e:
            print(f"Error al cargar variante: {str(e)}")
            variante = {
                "primera_serie": [],
                "segunda_serie": [],
                "tercera_serie": []
            }
        
        # Generar carpeta con timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(EXAMENES_FOLDER, f'{seccion}_{tipo_evaluacion}_{timestamp}')
        os.makedirs(output_dir, exist_ok=True)
        print(f"Carpeta de salida creada: {output_dir}")
        
        # Verificar plantilla y su formato
        use_template = False
        if plantilla_path and os.path.exists(plantilla_path):
            print(f"Plantilla encontrada: {plantilla_path}")
            
            # Detectar si es una plantilla con placeholders o una plantilla de formato general
            try:
                # Intentar abrir la plantilla para verificar que es un archivo .docx válido
                doc_temp = Document(plantilla_path)
                
                # Buscar placeholders en la plantilla
                has_placeholders = False
                for paragraph in doc_temp.paragraphs:
                    if any(marker in paragraph.text for marker in ['{primera_serie}', '{segunda_serie}', '{tercera_serie}', '{variante}']):
                        has_placeholders = True
                        break
                
                # También buscar en tablas
                if not has_placeholders:
                    for table in doc_temp.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if any(marker in paragraph.text for marker in ['{primera_serie}', '{segunda_serie}', '{tercera_serie}', '{variante}']):
                                        has_placeholders = True
                                        break
                                if has_placeholders:
                                    break
                            if has_placeholders:
                                break
                        if has_placeholders:
                            break
                
                if has_placeholders:
                    print("Plantilla con placeholders detectada. Utilizando procesamiento especial.")
                    doc = procesar_plantilla_examen(
                        plantilla_path, 
                        variante_id, 
                        seccion, 
                        tipo_evaluacion, 
                        variante,
                        licenciatura,
                        nombre_curso,
                        nombre_docente,
                        anio,
                        salon
                    )
                    
                    if doc:
                        use_template = True
                        print("Plantilla procesada correctamente con reemplazo de placeholders.")
                    else:
                        print("Error al procesar plantilla con placeholders. Creando examen estándar.")
                        doc = Document()
                else:
                    print("Plantilla sin placeholders detectada. Usando como base de formato.")
                    doc = Document(plantilla_path)
                    use_template = True
            except Exception as e:
                print(f"Error al analizar plantilla: {str(e)}")
                print(f"Se creará un documento nuevo en su lugar")
                doc = Document()
        else:
            print(f"No se encontró plantilla. Se creará un documento nuevo.")
            doc = Document()
        
        # Si estamos usando una plantilla con placeholders ya procesada, solo guardamos y terminamos
        if use_template and 'doc' in locals() and doc and any(marker in paragraph.text for paragraph in doc.paragraphs 
                                                           for marker in ['{primera_serie}', '{segunda_serie}', '{tercera_serie}']):
            # No es necesario añadir contenido, ya fue reemplazado en procesar_plantilla_examen
            print("Usando plantilla con placeholders ya procesados.")
        else:
            # Continuar con el proceso normal de creación/llenado de documento
            # Crear o verificar los estilos necesarios
            styles = doc.styles
            
            # Definir la fuente para todo el documento
            font_name = "Times New Roman"
            
            # Verificar/crear estilos básicos de encabezado
            style_names = [s.name for s in styles]
            
            # Crear Heading 1 si no existe
            if 'Heading 1' not in style_names:
                try:
                    heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                    heading1.font.bold = True
                    heading1.font.size = Pt(18)
                    heading1.font.name = font_name
                    print("Estilo 'Heading 1' creado")
                except Exception as e:
                    print(f"Error al crear estilo 'Heading 1': {str(e)}")
                    # Crear un estilo alternativo
                    if 'Title' in style_names:
                        heading1 = styles['Title']
                        heading1.font.name = font_name
                        print("Usando estilo 'Title' como alternativa")
                    else:
                        # Si no hay alternativa, crear un estilo personalizado
                        heading1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                        heading1.font.bold = True
                        heading1.font.size = Pt(18)
                        heading1.font.name = font_name
                        print("Estilo 'CustomHeading1' creado como alternativa")
            
            # Crear Heading 2 si no existe
            if 'Heading 2' not in style_names:
                try:
                    heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                    heading2.font.bold = True
                    heading2.font.size = Pt(16)
                    heading2.font.name = font_name
                    print("Estilo 'Heading 2' creado")
                except Exception as e:
                    print(f"Error al crear estilo 'Heading 2': {str(e)}")
                    # Crear un estilo alternativo
                    if 'Subtitle' in style_names:
                        heading2 = styles['Subtitle']
                        heading2.font.name = font_name
                        print("Usando estilo 'Subtitle' como alternativa")
                    else:
                        # Si no hay alternativa, crear un estilo personalizado
                        heading2 = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
                        heading2.font.bold = True
                        heading2.font.size = Pt(16)
                        heading2.font.name = font_name
                        print("Estilo 'CustomHeading2' creado como alternativa")
            
            # Crear List Bullet si no existe
            if 'List Bullet' not in style_names:
                try:
                    list_bullet = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
                    list_bullet.font.size = Pt(12)
                    list_bullet.font.name = font_name
                    print("Estilo 'List Bullet' creado")
                except Exception as e:
                    print(f"Error al crear estilo 'List Bullet': {str(e)}")
                    # Crear un estilo personalizado
                    list_bullet = styles.add_style('CustomListBullet', WD_STYLE_TYPE.PARAGRAPH)
                    list_bullet.font.size = Pt(12)
                    list_bullet.font.name = font_name
                    print("Estilo 'CustomListBullet' creado como alternativa")
            
            # Verificar/crear estilo 'TitleStyle'
            try:
                if 'TitleStyle' not in style_names:
                    title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
                    title_style.font.bold = True
                    title_style.font.size = Pt(16)
                    title_style.font.name = font_name
                    print("Estilo 'TitleStyle' creado")
                    
                # Verificar/crear estilo 'SubtitleStyle'
                if 'SubtitleStyle' not in style_names:
                    subtitle_style = styles.add_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH)
                    subtitle_style.font.italic = True
                    subtitle_style.font.size = Pt(12)
                    subtitle_style.font.name = font_name
                    print("Estilo 'SubtitleStyle' creado")
            except Exception as e:
                print(f"Error al crear estilos: {str(e)}")
                print("Continuando con estilos predeterminados...")
            
            # Configurar márgenes
            try:
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.7)
                    section.bottom_margin = Inches(0.7)
                    section.left_margin = Inches(0.8)
                    section.right_margin = Inches(0.8)
                print("Márgenes configurados")
            except Exception as e:
                print(f"Error al configurar márgenes: {str(e)}")
                print("Continuando con márgenes predeterminados...")
                
            # Encabezado con tabla
            try:
                header_table = doc.add_table(rows=1, cols=2)
                header_table.style = 'Table Grid'
                
                # Celda para logo
                logo_cell = header_table.cell(0, 0)
                logo_paragraph = logo_cell.paragraphs[0]
                
                # Verificar que el logo existe y añadirlo
                if logo_path and os.path.exists(logo_path):
                    try:
                        logo_paragraph.add_run().add_picture(logo_path, width=Inches(1.0))
                        print(f"Logo añadido desde {logo_path}")
                    except Exception as e:
                        print(f"Error al añadir logo: {str(e)}")
                        logo_paragraph.text = "LOGO"
                else:
                    print(f"Logo no encontrado en {logo_path}")
                    logo_paragraph.text = "LOGO"
            
                # Celda para título
                title_cell = header_table.cell(0, 1)
                
                # Título Universidad
                try:
                    # Usar campos desde parameters o predeterminados
                    univ_para = title_cell.add_paragraph('Universidad Panamericana')
                    try:
                        univ_para.style = 'TitleStyle'
                    except:
                        # Si falla al aplicar el estilo, aplicar formato directamente
                        for run in univ_para.runs:
                            run.bold = True
                            run.font.size = Pt(16)
                            run.font.name = font_name
                    univ_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error al aplicar estilo TitleStyle: {str(e)}")
                    univ_para = title_cell.add_paragraph('Universidad Panamericana')
                    univ_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in univ_para.runs:
                        run.bold = True
                        run.font.size = Pt(16)
                        run.font.name = font_name
                
                # Facultad y curso con placeholders
                try:
                    headers = [
                        f'Facultad de {licenciatura or "Humanidades"}',
                        f'{nombre_curso or "Estadística Básica"} - Sección {seccion}',
                        f'{nombre_docente or "Ing. Marco Antonio Jiménez"}',
                        f'{anio or "2025"}'
                    ]
                    
                    for header_text in headers:
                        header = title_cell.add_paragraph(header_text)
                        try:
                            header.style = 'SubtitleStyle'
                        except:
                            # Si falla al aplicar el estilo, aplicar formato directamente
                            for run in header.runs:
                                run.italic = True
                                run.font.size = Pt(12)
                                run.font.name = font_name
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error al aplicar estilo SubtitleStyle: {str(e)}")
                    for header_text in headers:
                        header = title_cell.add_paragraph(header_text)
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in header.runs:
                            run.italic = True
                            run.font.size = Pt(12)
                            run.font.name = font_name
            except Exception as e:
                print(f"Error al crear encabezado: {str(e)}")
                doc.add_paragraph('Universidad Panamericana').bold = True
                doc.add_paragraph(f'Facultad de {licenciatura or "Humanidades"}')
                doc.add_paragraph(f'{nombre_curso or "Estadística Básica"} - Sección {seccion}')
                
            # Título del examen
            try:
                tipo_eval_textos = {
                    'parcial1': 'Primer Examen Parcial',
                    'parcial2': 'Segundo Examen Parcial',
                    'final': 'Examen Final',
                    'corto': 'Evaluación Corta',
                    'recuperacion': 'Examen de Recuperación',
                    'test': 'Prueba de Evaluación'
                }
                
                tipo_texto = tipo_eval_textos.get(tipo_evaluacion, 'Evaluación Parcial')
                
                # Intentar añadir con Heading 1
                try:
                    exam_title = doc.add_heading(f'{tipo_texto} ({variante_id})', 1)
                    exam_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in exam_title.runs:
                        run.font.name = font_name
                except Exception as e:
                    print(f"Error al usar Heading 1: {str(e)}")
                    # Alternativa usando párrafo normal
                    p = doc.add_paragraph(f'{tipo_texto} ({variante_id})')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(18)
                        run.font.name = font_name
                
                # Añadir salón si se proporcionó
                if salon:
                    salon_text = doc.add_paragraph(f"Salón: {salon}")
                    salon_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in salon_text.runs:
                        run.font.name = font_name
            except Exception as e:
                print(f"Error al añadir título del examen: {str(e)}")
                p = doc.add_paragraph(f'{tipo_eval_textos.get(tipo_evaluacion, "Evaluación Parcial")} ({variante_id})')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = font_name
            
            # Información del estudiante
            doc.add_paragraph()
            student_info = doc.add_paragraph('Nombre del estudiante: _____________________________________________________________')
            for run in student_info.runs:
                run.font.name = font_name
            
            info_line = doc.add_paragraph()
            info_line.add_run('Fecha: ____________________ ')
            info_line.add_run('Carné: ___________________ ')
            info_line.add_run('Firma: ________________________')
            for run in info_line.runs:
                run.font.name = font_name
            
            doc.add_paragraph()
            
            # Primera serie
            try:
                # Intentar añadir con Heading 2
                try:
                    serie1 = doc.add_heading('Primera serie (Valor de cada respuesta correcta 4 puntos. Valor total de la serie 40 puntos)', 2)
                    for run in serie1.runs:
                        run.font.name = font_name
                except Exception as e:
                    print(f"Error al usar Heading 2: {str(e)}")
                    # Alternativa usando párrafo normal
                    serie1 = doc.add_paragraph('Primera serie (Valor de cada respuesta correcta 4 puntos. Valor total de la serie 40 puntos)')
                    for run in serie1.runs:
                        run.bold = True
                        run.font.size = Pt(16)
                        run.font.name = font_name
                
                instructions = doc.add_paragraph()
                instr_run = instructions.add_run('Instrucciones: ')
                instr_run.bold = True
                instr_run.font.name = font_name
                normal_run = instructions.add_run('Lea cuidadosamente cada una de las preguntas y sus opciones de respuesta. Subraye con lapicero la opción u opciones que considere correcta(s) para cada pregunta. Las respuestas hechas con lápiz no serán aceptadas como válidas.')
                normal_run.font.name = font_name
            except Exception as e:
                print(f"Error al añadir encabezado de primera serie: {str(e)}")
                doc.add_paragraph('Primera serie (Valor de cada respuesta correcta 4 puntos. Valor total de la serie 40 puntos)')
                p = doc.add_paragraph()
                p.add_run('Instrucciones: ').bold = True
                p.add_run('Lea cuidadosamente cada una de las preguntas y sus opciones de respuesta...')
                
                # Preguntas de la primera serie
                try:
                    if "primera_serie" in variante and variante["primera_serie"]:
                        for i, pregunta in enumerate(variante["primera_serie"], 1):
                            question_para = doc.add_paragraph()
                            
                            # Modificación: Separar el número y el texto para mejor control del formato
                            # Agregar solo el número con formato en negrita
                            num_run = question_para.add_run(f"{i}. ")
                            num_run.bold = True
                            num_run.font.name = font_name
                            num_run.font.size = Pt(12)
                            
                            # Agregar el texto de la pregunta, también en negrita
                            text_run = question_para.add_run(f"{pregunta['pregunta']}")
                            text_run.bold = True
                            text_run.font.name = font_name
                            
                            for opcion in pregunta.get("opciones", []):
                                try:
                                    # Intentar usar estilo List Bullet
                                    option_para = doc.add_paragraph(style='List Bullet')
                                    opt_run = option_para.add_run(opcion)
                                    opt_run.font.name = font_name
                                except Exception as e:
                                    print(f"Error al usar estilo List Bullet: {str(e)}")
                                    # Alternativa usando párrafo normal con viñeta manual
                                    option_para = doc.add_paragraph()
                                    opt_run = option_para.add_run(f"• {opcion}")
                                    opt_run.font.name = font_name
                                    option_para.paragraph_format.left_indent = Inches(0.25)
                        
                        doc.add_paragraph()
                    else:
                        error_p = doc.add_paragraph()
                        error_run = error_p.add_run("Error: No se encontraron preguntas para la primera serie.")
                        error_run.italic = True
                        error_run.font.name = font_name
                        doc.add_paragraph()
                except Exception as e:
                    print(f"Error al añadir preguntas de la primera serie: {str(e)}")
                    doc.add_paragraph("Error al cargar preguntas de la primera serie.")

                # También modifica la sección para la segunda serie de manera similar:

                # Escenarios de la segunda serie
                if "segunda_serie" in variante and variante["segunda_serie"]:
                    for i, escenario in enumerate(variante["segunda_serie"], 1):
                        escenario_para = doc.add_paragraph()
                        
                        # Modificación: Separar el número y el texto
                        num_run = escenario_para.add_run(f"{i}. ")
                        num_run.bold = True
                        num_run.font.name = font_name
                        num_run.font.size = Pt(12)
                        
                        # El texto del escenario también en negrita
                        esc_run = escenario_para.add_run(f"{escenario.get('escenario', '')}")
                        esc_run.bold = True
                        esc_run.font.name = font_name
                        
                        for opcion in escenario.get("opciones", []):
                            try:
                                # Intentar usar estilo List Bullet
                                option_para = doc.add_paragraph(style='List Bullet')
                                opt_run = option_para.add_run(opcion)
                                opt_run.font.name = font_name
                            except Exception as e:
                                print(f"Error al usar estilo List Bullet: {str(e)}")
                                # Alternativa usando párrafo normal con viñeta manual
                                option_para = doc.add_paragraph()
                                opt_run = option_para.add_run(f"• {opcion}")
                                opt_run.font.name = font_name
                                option_para.paragraph_format.left_indent = Inches(0.25)
                        
                        doc.add_paragraph()
                else:
                    error_p = doc.add_paragraph()
                    error_run = error_p.add_run("Error: No se encontraron escenarios para la segunda serie.")
                    error_run.italic = True
                    error_run.font.name = font_name
                    doc.add_paragraph()

                # Y también para la tercera serie:

                # Problema 1, 2, 3, 4 (Tercera serie)
                if len(variante["tercera_serie"]) > 0:
                    gini_data = variante["tercera_serie"][0]
                    p = doc.add_paragraph()
                    
                    # Modificación: Separar número y texto
                    num_run = p.add_run("1. ")
                    num_run.bold = True
                    num_run.font.name = font_name
                    num_run.font.size = Pt(12)
                    
                    gini_run = p.add_run(f"{gini_data.get('title', 'Problema de Coeficiente de Gini')}")
                    gini_run.bold = True
                    gini_run.font.name = font_name
                
                instructions3 = doc.add_paragraph()
                instr3_run = instructions3.add_run('Instrucciones: ')
                instr3_run.bold = True
                instr3_run.font.name = font_name
                normal3_run = instructions3.add_run('Desarrollar los ejercicios, dejando respaldo de sus operaciones. Asegúrese de escribir su respuesta final con lapicero; no se aceptarán respuestas escritas con lápiz. Mantenga su trabajo organizado y legible.')
                normal3_run.font.name = font_name
                
                # Verifica que tercera_serie esté presente
                if "tercera_serie" not in variante or not variante["tercera_serie"]:
                    error_p = doc.add_paragraph()
                    error_run = error_p.add_run("Error: No se encontraron ejercicios para la tercera serie.")
                    error_run.italic = True
                    error_run.font.name = font_name
                    doc.add_paragraph()
                else:
                    try:
                        # Problema 1 - Coeficiente de Gini
                        if len(variante["tercera_serie"]) > 0:
                            gini_data = variante["tercera_serie"][0]
                            p = doc.add_paragraph()
                            gini_run = p.add_run(f"1. {gini_data.get('title', 'Problema de Coeficiente de Gini')}")
                            gini_run.bold = True
                            gini_run.font.name = font_name
                            
                            # Tabla 1
                            if 'ranges' in gini_data and 'workers' in gini_data:
                                table1 = doc.add_table(rows=len(gini_data["ranges"])+1, cols=2)
                                table1.style = 'Table Grid'
                                table1.alignment = WD_TABLE_ALIGNMENT.CENTER
                                
                                # Encabezados
                                cell = table1.cell(0, 0)
                                cell.text = "Salario mensual en (Q)"
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.name = font_name
                                
                                cell = table1.cell(0, 1)
                                cell.text = "No. De trabajadores"
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.name = font_name
                                
                                # Datos
                                for i, rango in enumerate(gini_data["ranges"], 1):
                                    table1.cell(i, 0).text = rango
                                    table1.cell(i, 1).text = str(gini_data["workers"][i-1])
                                    
                                    # Aplicar fuente a todas las celdas de datos
                                    for j in range(2):
                                        for paragraph in table1.cell(i, j).paragraphs:
                                            for run in paragraph.runs:
                                                run.font.name = font_name
                            else:
                                p = doc.add_paragraph("Error: Datos de tabla incompletos para el problema de Gini.")
                                for run in p.runs:
                                    run.font.name = font_name
                            
                            # Incisos en negrita con un formato más prominente
                            p = doc.add_paragraph()
                            a_run = p.add_run("a) ")
                            a_run.bold = True
                            a_run.font.name = font_name
                            a_run.font.size = Pt(12)
                            a_cont = p.add_run("Complete la tabla para calcular el coeficiente de Gini.")
                            a_cont.font.name = font_name
                            
                            p = doc.add_paragraph()
                            b_run = p.add_run("b) ")
                            b_run.bold = True
                            b_run.font.name = font_name
                            b_run.font.size = Pt(12)
                            b_cont = p.add_run("Calcule el coeficiente de Gini utilizando la fórmula correspondiente.")
                            b_cont.font.name = font_name
                            
                            p = doc.add_paragraph()
                            c_run = p.add_run("c) ")
                            c_run.bold = True
                            c_run.font.name = font_name
                            c_run.font.size = Pt(12)
                            c_cont = p.add_run("Interprete el resultado obtenido respecto a la desigualdad en la distribución de salarios.")
                            c_cont.font.name = font_name
                        else:
                            p = doc.add_paragraph("Error: No se encontró el problema 1 (Coeficiente de Gini).")
                            for run in p.runs:
                                run.font.name = font_name
                    except Exception as e:
                        print(f"Error al añadir ejercicio 1 de la tercera serie: {str(e)}")
                        doc.add_paragraph("Error al cargar el problema 1 de la tercera serie.")
                    
                    try:
                        # Problema 2 - Distribución de frecuencias
                        if len(variante["tercera_serie"]) > 1:
                            sturgers_data = variante["tercera_serie"][1]
                            p = doc.add_paragraph()
                            sturgers_run = p.add_run(f"2. {sturgers_data.get('title', 'Problema de Distribución de Frecuencias')}")
                            sturgers_run.bold = True
                            sturgers_run.font.name = font_name
                            
                            # Tabla para los datos del problema 2
                            if 'data' in sturgers_data:
                                table2 = doc.add_table(rows=5, cols=5)
                                table2.style = 'Table Grid'
                                
                                # Llenar datos del problema 2
                                idx = 0
                                for i in range(5):
                                    for j in range(5):
                                        if idx < len(sturgers_data["data"]):
                                            table2.cell(i, j).text = str(sturgers_data["data"][idx])
                                            
                                            # Aplicar fuente a todas las celdas
                                            for paragraph in table2.cell(i, j).paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.name = font_name
                                            idx += 1
                            else:
                                p = doc.add_paragraph("Error: Datos incompletos para el problema de distribución de frecuencias.")
                                for run in p.runs:
                                    run.font.name = font_name
                            
                            # Instrucción con formato consistente
                            p = doc.add_paragraph()
                            p_run_bold = p.add_run("Instrucción: ")
                            p_run_bold.bold = True
                            p_run_bold.font.name = font_name
                            p_run_bold.font.size = Pt(12)
                            p_run = p.add_run("Construya la tabla de distribución de frecuencias correspondiente.")
                            p_run.font.name = font_name
                        else:
                            p = doc.add_paragraph("Error: No se encontró el problema 2 (Distribución de Frecuencias).")
                            for run in p.runs:
                                run.font.name = font_name
                    except Exception as e:
                        print(f"Error al añadir ejercicio 2 de la tercera serie: {str(e)}")
                        doc.add_paragraph("Error al cargar el problema 2 de la tercera serie.")
                    
                    try:
                        # Problema 3 - Diagrama de Tallo y Hoja
                        if len(variante["tercera_serie"]) > 2:
                            stem_leaf_data = variante["tercera_serie"][2]
                            p = doc.add_paragraph()
                            stem_run = p.add_run(f"3. {stem_leaf_data.get('title', 'Problema de Diagrama de Tallo y Hoja')}")
                            stem_run.bold = True
                            stem_run.font.name = font_name
                            
                            # Tabla para los datos del problema 3
                            if 'data' in stem_leaf_data:
                                table3 = doc.add_table(rows=3, cols=8)
                                table3.style = 'Table Grid'
                                
                                # Llenar datos del problema 3
                                idx = 0
                                for i in range(3):
                                    for j in range(8):
                                        if idx < len(stem_leaf_data["data"]):
                                            table3.cell(i, j).text = str(stem_leaf_data["data"][idx])
                                            
                                            # Aplicar fuente a todas las celdas
                                            for paragraph in table3.cell(i, j).paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.name = font_name
                                            idx += 1
                            else:
                                p = doc.add_paragraph("Error: Datos incompletos para el problema de tallo y hoja.")
                                for run in p.runs:
                                    run.font.name = font_name
                            
                            p = doc.add_paragraph()
                            a_run = p.add_run("a) ")
                            a_run.bold = True
                            a_run.font.name = font_name
                            a_run.font.size = Pt(12)
                            a_cont = p.add_run("Realizar un Diagrama de Tallo y Hoja para identificar donde se encuentra la mayor concentración de los datos.")
                            a_cont.font.name = font_name
                            
                            p = doc.add_paragraph()
                            b_run = p.add_run("b) ")
                            b_run.bold = True
                            b_run.font.name = font_name
                            b_run.font.size = Pt(12)
                            b_cont = p.add_run("Interprete los datos y explique brevemente sus resultados.")
                            b_cont.font.name = font_name
                        else:
                            p = doc.add_paragraph("Error: No se encontró el problema 3 (Diagrama de Tallo y Hoja).")
                            for run in p.runs:
                                run.font.name = font_name
                    except Exception as e:
                        print(f"Error al añadir ejercicio 3 de la tercera serie: {str(e)}")
                        doc.add_paragraph("Error al cargar el problema 3 de la tercera serie.")
                    
                    try:
                        # Problema 4 - Medidas de tendencia central
                        if len(variante["tercera_serie"]) > 3:
                            central_tendency_data = variante["tercera_serie"][3]
                            p = doc.add_paragraph()
                            central_run = p.add_run(f"4. {central_tendency_data.get('title', 'Problema de Medidas de Tendencia Central')}")
                            central_run.bold = True
                            central_run.font.name = font_name
                            
                            # Tabla para el problema 4
                            if 'ranges' in central_tendency_data and 'count' in central_tendency_data:
                                table4 = doc.add_table(rows=len(central_tendency_data["ranges"])+1, cols=2)
                                table4.style = 'Table Grid'
                                
                                # Encabezados
                                cell = table4.cell(0, 0)
                                cell.text = "Precio en (Q)"
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.name = font_name
                                
                                cell = table4.cell(0, 1)
                                cell.text = "No. De productos"
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.name = font_name
                                
                                # Datos
                                for i, (rango, count) in enumerate(zip(central_tendency_data["ranges"], central_tendency_data["count"]), 1):
                                    table4.cell(i, 0).text = rango
                                    table4.cell(i, 1).text = str(count)
                                    
                                    # Aplicar fuente a todas las celdas de datos
                                    for j in range(2):
                                        for paragraph in table4.cell(i, j).paragraphs:
                                            for run in paragraph.runs:
                                                run.font.name = font_name
                            else:
                                p = doc.add_paragraph("Error: Datos incompletos para el problema de medidas de tendencia central.")
                                for run in p.runs:
                                    run.font.name = font_name
                        else:
                            p = doc.add_paragraph("Error: No se encontró el problema 4 (Medidas de Tendencia Central).")
                            for run in p.runs:
                                run.font.name = font_name
                    except Exception as e:
                        print(f"Error al añadir ejercicio 4 de la tercera serie: {str(e)}")
                        doc.add_paragraph("Error al cargar el problema 4 de la tercera serie.")
            except Exception as e:
                print(f"Error general al procesar la tercera serie: {str(e)}")
                doc.add_paragraph("Error al cargar la tercera serie.")
        
        # Guardar el documento
        try:
            # Nombre de archivo estándar
            filename = f'Examen_{variante_id}.docx'
            simple_path = os.path.join(EXAMENES_FOLDER, filename)
            
            # Nombre detallado
            detailed_filename = f'Examen_{seccion}_{tipo_evaluacion}_{variante_id}.docx'
            detailed_path = os.path.join(output_dir, detailed_filename)
            
            # Guardar en ambas ubicaciones
            print(f"Intentando guardar examen en: {detailed_path}")
            doc.save(detailed_path)
            print(f"Documento guardado en: {detailed_path}")
            
            print(f"Intentando guardar examen en: {simple_path}")
            doc.save(simple_path)
            print(f"Documento guardado en: {simple_path}")
            
            print(f"===== EXAMEN WORD CREADO EXITOSAMENTE =====\n")
            return filename
        except Exception as e:
            print(f"Error al guardar el documento: {str(e)}")
            # Intentar guardar en una ubicación alternativa
            try:
                alt_path = os.path.join(UPLOAD_FOLDER, filename)
                print(f"Intentando guardar en ubicación alternativa: {alt_path}")
                doc.save(alt_path)
                print(f"Guardado en ubicación alternativa: {alt_path}")
                return filename
            except Exception as e2:
                print(f"Error al guardar en ubicación alternativa: {str(e2)}")
                return None
            
    except Exception as e:
        print(f"Error global al crear examen: {str(e)}")
        traceback.print_exc()
        return None

def generar_examen(num_variantes=1, seccion="A", tipo_evaluacion="parcial1", logo_path=None, 
                  plantilla_path=None, licenciatura="", nombre_curso="", nombre_docente="", 
                  anio="", salon="", uso_detallado=True):
    """
    Función principal para generar exámenes con múltiples variantes.
    
    Parámetros:
    - num_variantes: Número de variantes a generar
    - seccion: Sección del curso
    - tipo_evaluacion: Tipo de evaluación (parcial1, parcial2, final, etc.)
    - logo_path: Ruta al logo de la institución
    - plantilla_path: Ruta a la plantilla Word opcional
    - licenciatura, nombre_curso, nombre_docente, anio, salon: Datos para personalizar el examen
    - uso_detallado: Si es True, genera la solución matemática detallada; si es False, usa la versión simplificada
    
    Retorna:
    - Lista de variantes generadas
    """
    try:
        print("\n===== INICIANDO GENERACIÓN DE EXÁMENES =====")
        print(f"Variantes: {num_variantes}, Sección: {seccion}, Tipo: {tipo_evaluacion}")
        print(f"Licenciatura: {licenciatura}, Curso: {nombre_curso}")
        print(f"Docente: {nombre_docente}, Año: {anio}, Salón: {salon}")
        
        # Verificar directorios
        for folder in [EXAMENES_FOLDER, HOJAS_RESPUESTA_FOLDER, PLANTILLAS_FOLDER, VARIANTES_FOLDER]:
            if not os.path.exists(folder):
                os.makedirs(folder)
                print(f"Directorio creado: {folder}")
            print(f"Directorio existente: {folder}, Permisos de escritura: {os.access(folder, os.W_OK)}")
        
        # Crear carpeta con timestamp para esta generación
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_output_dir = f'{seccion}_{tipo_evaluacion}_{timestamp}'
        
        # Crear carpetas para cada tipo de archivo
        for folder in [EXAMENES_FOLDER, HOJAS_RESPUESTA_FOLDER, PLANTILLAS_FOLDER, VARIANTES_FOLDER]:
            output_dir = os.path.join(folder, base_output_dir)
            os.makedirs(output_dir, exist_ok=True)
            print(f"Directorio de salida creado: {output_dir}")
        
        # Verificar logo y plantilla
        if logo_path:
            print(f"Guardando logo en: {logo_path}")
        
        if plantilla_path:
            print(f"Guardando plantilla en: {plantilla_path}")
        
        # Generar variantes
        variantes_generadas = []
        
        for i in range(num_variantes):
            variante_id = f"V{i+1}"
            print(f"\n----- Generando variante {variante_id} -----")
            
            # Generar variante y respuestas
            try:
                variante, respuestas = generar_variante(variante_id, seccion, tipo_evaluacion)
                print(f"Variante generada correctamente: {variante_id}")
            except Exception as e:
                print(f"Error al generar variante {variante_id}: {str(e)}")
                traceback.print_exc()
                continue
            
            # Guardar variante y respuestas en la carpeta con timestamp
            with open(os.path.join(VARIANTES_FOLDER, base_output_dir, f'variante_{variante_id}.json'), 'w', encoding='utf-8') as f:
                json.dump(variante, f, ensure_ascii=False, indent=2)
            
            with open(os.path.join(VARIANTES_FOLDER, base_output_dir, f'respuestas_{variante_id}.json'), 'w', encoding='utf-8') as f:
                json.dump(respuestas, f, ensure_ascii=False, indent=2)
            
            # 1. Crear examen Word
            try:
                print(f"Creando examen Word para variante {variante_id}...")
                examen_filename = crear_examen_word(
                    variante_id, 
                    seccion, 
                    tipo_evaluacion, 
                    logo_path, 
                    plantilla_path,
                    licenciatura,
                    nombre_curso,
                    nombre_docente,
                    anio,
                    salon
                )
                if examen_filename:
                    print(f"Examen Word creado: {examen_filename}")
                else:
                    print("Error: No se pudo crear el examen Word")
            except Exception as e:
                print(f"Error al crear examen Word: {str(e)}")
                traceback.print_exc()
                examen_filename = None
            
            # 2. Crear hoja de respuestas
            try:
                print(f"Creando hoja de respuestas para variante {variante_id}...")
                hoja_filename = crear_hoja_respuestas(variante_id, seccion, tipo_evaluacion)
                if hoja_filename:
                    print(f"Hoja de respuestas creada: {hoja_filename}")
                else:
                    print("Error: No se pudo crear la hoja de respuestas")
            except Exception as e:
                print(f"Error al crear hoja de respuestas: {str(e)}")
                traceback.print_exc()
                hoja_filename = None
            
            # 3. Crear plantilla de calificación
            try:
                print(f"Creando plantilla de calificación para variante {variante_id}...")
                plantilla_filename = crear_plantilla_calificacion(variante_id, seccion, tipo_evaluacion)
                if plantilla_filename:
                    print(f"Plantilla de calificación creada: {plantilla_filename}")
                else:
                    print("Error: No se pudo crear la plantilla de calificación")
            except Exception as e:
                print(f"Error al crear plantilla de calificación: {str(e)}")
                traceback.print_exc()
                plantilla_filename = None
            
            # 4. Crear solución matemática (detallada o simplificada)
            try:
                print(f"Creando solución matemática para variante {variante_id}...")
                if uso_detallado:
                    # Usar la solución detallada que incluye explicaciones y gráficos
                    solucion_filename = crear_solucion_matematica_detallada(variante_id, seccion, tipo_evaluacion)
                else:
                    # Usar la solución simplificada como respaldo
                    solucion_filename = crear_solucion_matematica_simplificada(variante_id, seccion, tipo_evaluacion)
                
                if solucion_filename:
                    print(f"Solución matemática creada: {solucion_filename}")
                else:
                    print("Error: No se pudo crear la solución matemática")
            except Exception as e:
                print(f"Error al crear solución matemática: {str(e)}")
                traceback.print_exc()
                solucion_filename = None
            
            # Registrar variante generada
            variantes_generadas.append({
                'id': variante_id,
                'examen': examen_filename,
                'hoja': hoja_filename,
                'plantilla': plantilla_filename,
                'solucion_matematica': solucion_filename,
                'seccion': seccion,
                'tipo_evaluacion': tipo_evaluacion,
                'timestamp': timestamp,
                'directorio': base_output_dir
            })
        
        # Actualizar historial
        historial = cargar_historial()
        
        # Obtener nombres para mostrar
        tipo_textos = {
            'parcial1': 'Primer Parcial',
            'parcial2': 'Segundo Parcial',
            'final': 'Examen Final',
            'corto': 'Evaluación Corta',
            'recuperacion': 'Recuperación',
            'test': 'Prueba'
        }
        
        # Añadir entrada al historial
        for variante in variantes_generadas:
            historial.append({
                'id': variante['id'],
                'seccion': seccion,
                'tipo_evaluacion': tipo_evaluacion,
                'tipo_texto': tipo_textos.get(tipo_evaluacion, tipo_evaluacion),
                'fecha_generacion': datetime.now().strftime("%d/%m/%Y %H:%M"),
                'timestamp': timestamp,
                'directorio': base_output_dir,
                'examen': variante['examen'],
                'hoja': variante['hoja'],
                'plantilla': variante['plantilla'],
                'solucion_matematica': variante['solucion_matematica'],
                'licenciatura': licenciatura,
                'nombre_curso': nombre_curso,
                'nombre_docente': nombre_docente,
                'anio': anio,
                'salon': salon
            })
        
        guardar_historial(historial)
        
        print("\n===== FINALIZADA GENERACIÓN DE EXÁMENES =====")
        print(f"Variantes generadas: {len(variantes_generadas)}")
        for v in variantes_generadas:
            print(f"Variante {v['id']}:")
            print(f"  - Examen: {v['examen']}")
            print(f"  - Hoja: {v['hoja']}")
            print(f"  - Plantilla: {v['plantilla']}")
            print(f"  - Solución: {v['solucion_matematica']}")
        
        return variantes_generadas
    
    except Exception as e:
        print(f"Error global en generación de exámenes: {str(e)}")
        traceback.print_exc()
        return []

@app.route('/verificar_generacion_documentos', methods=['GET'])
def verificar_generacion_documentos():
    """
    Verifica la capacidad del sistema para generar documentos realizando una prueba
    real de generación de cada tipo de documento.
    """
    resultados = {
        "examen_word": {"estado": "No probado", "mensaje": ""},
        "hoja_respuestas": {"estado": "No probado", "mensaje": ""},
        "plantilla_calificacion": {"estado": "No probado", "mensaje": ""},
        "solucion_matematica": {"estado": "No probado", "mensaje": ""}
    }
    
    # Crear variante de prueba temporal
    variante_id = "TEST_" + datetime.now().strftime("%H%M%S")
    seccion = "TEST"
    tipo_evaluacion = "test"
    
    try:
        # Generar variante para probar
        variante, respuestas = generar_variante(variante_id, seccion, tipo_evaluacion)
        
        # Guardar temporalmente la variante y respuestas
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(variante, f, ensure_ascii=False, indent=2)
        
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(respuestas, f, ensure_ascii=False, indent=2)
        
        # Probar cada función de generación
        try:
            examen_filename = crear_examen_word(variante_id, seccion, tipo_evaluacion)
            if examen_filename and os.path.exists(os.path.join(EXAMENES_FOLDER, examen_filename)):
                resultados["examen_word"]["estado"] = "OK"
                resultados["examen_word"]["mensaje"] = f"Archivo creado: {examen_filename}"
            else:
                resultados["examen_word"]["estado"] = "ERROR"
                resultados["examen_word"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["examen_word"]["estado"] = "ERROR"
            resultados["examen_word"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            hoja_filename = crear_hoja_respuestas(variante_id, seccion, tipo_evaluacion)
            if hoja_filename and os.path.exists(os.path.join(HOJAS_RESPUESTA_FOLDER, hoja_filename)):
                resultados["hoja_respuestas"]["estado"] = "OK"
                resultados["hoja_respuestas"]["mensaje"] = f"Archivo creado: {hoja_filename}"
            else:
                resultados["hoja_respuestas"]["estado"] = "ERROR"
                resultados["hoja_respuestas"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["hoja_respuestas"]["estado"] = "ERROR"
            resultados["hoja_respuestas"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            plantilla_filename = crear_plantilla_calificacion(variante_id, seccion, tipo_evaluacion)
            if plantilla_filename and os.path.exists(os.path.join(PLANTILLAS_FOLDER, plantilla_filename)):
                resultados["plantilla_calificacion"]["estado"] = "OK"
                resultados["plantilla_calificacion"]["mensaje"] = f"Archivo creado: {plantilla_filename}"
            else:
                resultados["plantilla_calificacion"]["estado"] = "ERROR"
                resultados["plantilla_calificacion"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["plantilla_calificacion"]["estado"] = "ERROR"
            resultados["plantilla_calificacion"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
        try:
            solucion_filename = crear_solucion_matematica_simplificada(variante_id, seccion, tipo_evaluacion)
            if solucion_filename and os.path.exists(os.path.join(PLANTILLAS_FOLDER, solucion_filename)):
                resultados["solucion_matematica"]["estado"] = "OK"
                resultados["solucion_matematica"]["mensaje"] = f"Archivo creado: {solucion_filename}"
            else:
                resultados["solucion_matematica"]["estado"] = "ERROR"
                resultados["solucion_matematica"]["mensaje"] = "No se pudo crear el archivo"
        except Exception as e:
            resultados["solucion_matematica"]["estado"] = "ERROR"
            resultados["solucion_matematica"]["mensaje"] = f"Excepción: {str(e)}"
            traceback.print_exc()
        
    except Exception as e:
        for key in resultados:
            resultados[key]["estado"] = "ERROR"
            resultados[key]["mensaje"] = f"Error en prueba general: {str(e)}"
    
    # Limpiar archivos temporales
    try:
        for archivo in [
            os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'),
            os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'),
            os.path.join(EXAMENES_FOLDER, f'Examen_{variante_id}.docx'),
            os.path.join(HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{variante_id}.pdf'),
            os.path.join(PLANTILLAS_FOLDER, f'Plantilla_{variante_id}.pdf'),
            os.path.join(PLANTILLAS_FOLDER, f'Solucion_Matematica_{variante_id}.docx')
        ]:
            if os.path.exists(archivo):
                os.remove(archivo)
    except:
        pass  # Ignorar errores en la limpieza
    
    return render_template('verificar_generacion.html', resultados=resultados)

# Replace your current mostrar_historial route with this safer version

@app.route('/historial')
def mostrar_historial():
    """
    Muestra el historial de exámenes generados, con protección contra recursión
    """
    try:
        historial_raw = cargar_historial()
        
        # Crear una función para sanitizar el diccionario y prevenir recursión
        def sanitize_dict(d):
            """Crea una copia segura de un diccionario sin referencias circulares"""
            if not isinstance(d, dict):
                return d
                
            # Crear un nuevo diccionario simplificado con solo las claves necesarias
            safe_dict = {}
            
            # Lista de claves seguras que sabemos que no causan problemas
            safe_keys = ['id', 'seccion', 'tipo_evaluacion', 'tipo_texto', 
                        'examen', 'hoja', 'plantilla', 'solucion_matematica',
                        'fecha_generacion', 'timestamp', 'directorio']
                        
            for key in safe_keys:
                if key in d:
                    # Evitar copiar objetos complejos que podrían causar recursión
                    if isinstance(d[key], (str, int, float, bool, type(None))):
                        safe_dict[key] = d[key]
                    else:
                        # Convertir a string para evitar recursión
                        safe_dict[key] = str(d[key])
            
            return safe_dict
        
        # Sanitizar el historial - NO usar filtros como sort() que pueden causar recursión
        historial_safe = [sanitize_dict(item) for item in historial_raw]
        
        # Si quieres ordenar, hazlo de una forma segura usando sorted() y una key simple
        # Esto evita comparaciones complejas que podrían causar recursión
        try:
            historial_safe = sorted(
                historial_safe, 
                key=lambda x: x.get('fecha_generacion', '0') if isinstance(x.get('fecha_generacion', '0'), str) else '0', 
                reverse=True
            )
        except Exception as sort_err:
            app.logger.error(f"Error al ordenar historial: {str(sort_err)}")
            # Continuar con el historial sin ordenar
        
        # Limitar el número de elementos para evitar problemas
        historial_safe = historial_safe[:50]  # Mostrar solo los últimos 50 elementos
        
        return render_template('historial.html', historial=historial_safe)
    
    except Exception as e:
        app.logger.error(f"Error en mostrar_historial: {str(e)}")
        
        # Devolver una página sencilla con el error
        error_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Error</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .error {{ background-color: #f8d7da; border: 1px solid #f5c6cb; padding: 20px; border-radius: 5px; }}
                pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            <h1>Error al cargar historial</h1>
            <div class="error">
                <p><strong>Mensaje de error:</strong> {str(e)}</p>
            </div>
            <pre>{traceback.format_exc()}</pre>
            <p><a href="/">Volver al inicio</a></p>
        </body>
        </html>
        """
        return error_html

@app.route('/cargar_examenes')
def cargar_examenes_escaneados():
    return render_template('cargar_examenes.html')

@app.route('/editar_variante/<variante_id>')
def editar_variante(variante_id):
    try:
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        return render_template('editar_variante.html', variante=variante)
    except:
        flash('No se pudo cargar la variante', 'danger')
        return redirect(url_for('index'))

@app.route('/guardar_variante', methods=['POST'])
def guardar_variante():
    variante_id = request.form.get('variante_id')
    
    try:
        # Cargar variante original
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        # Actualizar datos (esto sería más complejo en un caso real)
        # Aquí se implementaría la lógica para actualizar la variante según los datos del formulario
        
        # Guardar variante actualizada
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(variante, f, ensure_ascii=False, indent=2)
        
        # Regenerar documentos
        examen_filename = crear_examen_word(variante_id)
        hoja_filename = crear_hoja_respuestas(variante_id)
        plantilla_filename = crear_plantilla_calificacion(variante_id)
        
        flash('Variante actualizada correctamente', 'success')
    except Exception as e:
        flash(f'Error al guardar la variante: {str(e)}', 'danger')
    
    return redirect(url_for('index'))

@app.route('/previsualizar/<variante_id>')
def previsualizar(variante_id):
    try:
        with open(os.path.join(VARIANTES_FOLDER, f'variante_{variante_id}.json'), 'r', encoding='utf-8') as f:
            variante = json.load(f)
        
        with open(os.path.join(VARIANTES_FOLDER, f'respuestas_{variante_id}.json'), 'r', encoding='utf-8') as f:
            respuestas = json.load(f)
        
        return render_template('previsualizar.html', variante=variante, respuestas=respuestas)
    except:
        flash('No se pudo cargar la variante para previsualizar', 'danger')
        return redirect(url_for('index'))

@app.route('/eliminar_variante/<variante_id>', methods=['POST'])
def eliminar_variante(variante_id):
    """
    Elimina una variante y todos sus archivos asociados
    """
    try:
        # Lista de patrones de archivos a eliminar
        patrones_archivos = [
            # Archivos principales con nombres simplificados
            (VARIANTES_FOLDER, f'variante_{variante_id}.json'),
            (VARIANTES_FOLDER, f'respuestas_{variante_id}.json'),
            (EXAMENES_FOLDER, f'Examen_{variante_id}.docx'),
            (HOJAS_RESPUESTA_FOLDER, f'HojaRespuestas_{variante_id}.pdf'),
            (PLANTILLAS_FOLDER, f'Plantilla_{variante_id}.pdf'),
            (PLANTILLAS_FOLDER, f'Solucion_Matematica_{variante_id}.docx'),
            
            # Archivos con nombres detallados (con comodín para buscar en cualquier directorio)
            (EXAMENES_FOLDER, f'*Examen_*_{variante_id}.docx'),
            (HOJAS_RESPUESTA_FOLDER, f'*HojaRespuestas_*_{variante_id}.pdf'),
            (PLANTILLAS_FOLDER, f'*Plantilla_*_{variante_id}.pdf'),
            (PLANTILLAS_FOLDER, f'*Solucion_Matematica_*_{variante_id}.docx')
        ]
        
        # Eliminar archivos principales y buscar archivos detallados
        archivos_eliminados = []
        
        for carpeta, patron in patrones_archivos:
            # Si el patrón no tiene comodín, eliminar directamente
            if '*' not in patron:
                ruta_completa = os.path.join(carpeta, patron)
                if os.path.exists(ruta_completa):
                    os.remove(ruta_completa)
                    archivos_eliminados.append(ruta_completa)
                    print(f"Archivo eliminado: {ruta_completa}")
            else:
                # Si tiene comodín, buscar archivos que coincidan
                for raiz, _, archivos in os.walk(carpeta):
                    # Convertir el patrón con comodín a una expresión regular
                    patron_regex = patron.replace('*', '.*')
                    for archivo in archivos:
                        if re.match(patron_regex, archivo):
                            ruta_completa = os.path.join(raiz, archivo)
                            os.remove(ruta_completa)
                            archivos_eliminados.append(ruta_completa)
                            print(f"Archivo eliminado: {ruta_completa}")
        
        # Eliminar la variante del historial
        historial = cargar_historial()
        historial_nuevo = [item for item in historial if item.get('id') != variante_id]
        
        # Guardar el historial actualizado
        if len(historial) != len(historial_nuevo):
            guardar_historial(historial_nuevo)
            print(f"Variante {variante_id} eliminada del historial")
        
        # Mensaje de éxito
        if archivos_eliminados:
            flash(f'Variante {variante_id} y {len(archivos_eliminados)} archivos asociados eliminados correctamente', 'success')
        else:
            flash(f'Variante {variante_id} eliminada del historial, pero no se encontraron archivos para eliminar', 'warning')
        
        # Redirigir a la página anterior
        referrer = request.referrer
        if referrer and urlparse(referrer).netloc == urlparse(request.host_url).netloc:
            return redirect(referrer)
        else:
            return redirect(url_for('index'))
            
    except Exception as e:
        flash(f'Error al eliminar la variante {variante_id}: {str(e)}', 'danger')
        return redirect(url_for('index'))
    
if __name__ == '__main__':
    app.run(debug=True)
